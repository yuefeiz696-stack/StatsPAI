"""
Unified results class for all econometric models
"""

from html import escape as _html_escape
from typing import Dict, Any, Optional, List, Union
import pandas as pd
import numpy as np


def _scipy_stats():
    """Lazily import ``scipy.stats`` for result-time inference only."""
    from scipy import stats as _stats
    return _stats


class SummaryText(str):
    """``str`` subclass that renders formatted summaries in Jupyter / IPython.

    ``CausalResult.summary()`` and ``EconometricResults.summary()`` produce
    multi-line text. Returning a plain ``str`` as the last expression of a
    notebook cell triggers Python's ``repr()`` — the user sees an escaped
    one-liner with literal ``\\n``. Wrapping the same text in
    :class:`SummaryText` keeps every ``str`` operation intact (slicing,
    ``.split``, ``in`` membership, ``isinstance(x, str)``) and adds the
    IPython display hooks so the cell renders the formatted block directly,
    without requiring ``print()``.
    """

    __slots__ = ()

    def _repr_pretty_(self, p, cycle):  # IPython terminal & notebook
        p.text(str(self))

    def _repr_html_(self) -> str:  # Jupyter HTML rendering
        return (
            "<pre style=\"font-family: 'SFMono-Regular', Menlo, Consolas, "
            "monospace; line-height: 1.35; white-space: pre;\">"
            f"{_html_escape(str(self))}"
            "</pre>"
        )


# ----------------------------------------------------------------------
# JSON-safe coercion used by to_dict() / for_agent() on the result
# classes below.  Agents consume these methods; the output MUST round-trip
# through json.dumps without raising — numpy scalars, pandas objects,
# tuples, and numpy arrays all need primitive-form conversion.
# ----------------------------------------------------------------------

def _to_jsonable(value: Any, *, _depth: int = 0, _max_depth: int = 6) -> Any:
    """Return a JSON-safe representation of ``value``.

    Handles numpy scalars/arrays, pandas Series/DataFrame, tuples/lists,
    dicts, NaN/Inf (converted to None) and bails out gracefully on
    anything exotic by falling back to ``str(value)``.
    """
    if _depth > _max_depth:
        return str(value)
    if value is None:
        return None
    if isinstance(value, bool):
        return bool(value)
    if isinstance(value, (int, np.integer)):
        return int(value)
    if isinstance(value, (float, np.floating)):
        v = float(value)
        if np.isnan(v) or np.isinf(v):
            return None
        return v
    if isinstance(value, str):
        return value
    if isinstance(value, (list, tuple)):
        return [_to_jsonable(v, _depth=_depth + 1) for v in value]
    if isinstance(value, dict):
        return {str(k): _to_jsonable(v, _depth=_depth + 1)
                for k, v in value.items()}
    if isinstance(value, np.ndarray):
        return _to_jsonable(value.tolist(), _depth=_depth + 1)
    if isinstance(value, pd.Series):
        return _to_jsonable(value.to_dict(), _depth=_depth + 1)
    if isinstance(value, pd.DataFrame):
        # Keep the payload bounded — agents don't need full detail tables.
        head = value.head(20)
        return _to_jsonable(head.to_dict(orient='records'),
                            _depth=_depth + 1)
    if isinstance(value, pd.Timestamp):
        return value.isoformat()
    # Last resort: stringify so json.dumps doesn't blow up.
    try:
        return str(value)
    except Exception:
        return None


def _filter_jsonable_scalars(d: Dict[str, Any]) -> Dict[str, Any]:
    """Return the subset of ``d`` whose values round-trip as JSON scalars.

    Used for ``diagnostics`` payloads where agents want simple key/value
    pairs — not nested frames.
    """
    out: Dict[str, Any] = {}
    for k, v in (d or {}).items():
        j = _to_jsonable(v)
        if isinstance(j, (int, float, str, bool)) or j is None:
            out[str(k)] = j
    return out


class EconometricResults:
    """
    Unified results class for econometric models
    
    This class provides a consistent interface for accessing results
    from different econometric estimators, similar to R's broom package.
    """
    
    def __init__(
        self,
        params: pd.Series,
        std_errors: pd.Series,
        model_info: Dict[str, Any],
        data_info: Optional[Dict[str, Any]] = None,
        diagnostics: Optional[Dict[str, Any]] = None
    ):
        """
        Initialize results object
        
        Parameters
        ----------
        params : pd.Series
            Parameter estimates with variable names as index
        std_errors : pd.Series
            Standard errors with variable names as index
        model_info : Dict[str, Any]
            Model metadata (model type, estimation method, etc.)
        data_info : Dict[str, Any], optional
            Data metadata (sample size, variable names, etc.)
        diagnostics : Dict[str, Any], optional
            Model diagnostics (R-squared, F-statistics, etc.)
        """
        self.params = params
        self.std_errors = std_errors
        self.model_info = model_info
        self.data_info = data_info or {}
        self.diagnostics = diagnostics or {}
        
        # Compute derived statistics
        self._compute_statistics()
    
    def _compute_statistics(self):
        """Compute t-statistics, p-values, and confidence intervals"""
        stats = _scipy_stats()
        self.tvalues = self.params / self.std_errors
        self.pvalues = 2 * (1 - stats.t.cdf(np.abs(self.tvalues), 
                                           self.data_info.get('df_resid', np.inf)))
        
        # 95% confidence intervals by default
        alpha = 0.05
        t_crit = stats.t.ppf(1 - alpha/2, self.data_info.get('df_resid', np.inf))
        self.conf_int_lower = self.params - t_crit * self.std_errors
        self.conf_int_upper = self.params + t_crit * self.std_errors
    
    def summary(self, alpha: float = 0.05) -> str:
        """
        Generate a summary table of results
        
        Parameters
        ----------
        alpha : float, default 0.05
            Significance level for confidence intervals
            
        Returns
        -------
        str
            Formatted summary table
        """
        # Create coefficients table
        coef_table = pd.DataFrame({
            'Coefficient': self.params,
            'Std. Error': self.std_errors,
            't-statistic': self.tvalues,
            'P>|t|': self.pvalues,
            f'[{alpha/2:.3f}': self.conf_int_lower,
            f'{1-alpha/2:.3f}]': self.conf_int_upper
        })
        
        # Format the output
        output = []
        output.append("=" * 80)
        output.append(f"Model: {self.model_info.get('model_type', 'Unknown')}")
        output.append(f"Method: {self.model_info.get('method', 'Unknown')}")
        if 'dependent_var' in self.data_info:
            output.append(f"Dependent Variable: {self.data_info['dependent_var']}")
        output.append("=" * 80)
        
        # Add coefficient table
        output.append(coef_table.to_string(float_format='%.4f'))
        
        # Add model diagnostics
        if self.diagnostics:
            output.append("")
            output.append("Model Diagnostics:")
            output.append("-" * 20)
            for key, value in self.diagnostics.items():
                if isinstance(value, (int, float)):
                    output.append(f"{key:20s}: {value:.4f}")
                else:
                    output.append(f"{key:20s}: {value}")
        
        output.append("=" * 80)
        return SummaryText("\n".join(output))

    def conf_int(self, alpha: float = 0.05) -> pd.DataFrame:
        """
        Return confidence intervals for parameters
        
        Parameters
        ----------
        alpha : float, default 0.05
            Significance level
            
        Returns
        -------
        pd.DataFrame
            Confidence intervals
        """
        stats = _scipy_stats()
        t_crit = stats.t.ppf(1 - alpha/2, self.data_info.get('df_resid', np.inf))
        lower = self.params - t_crit * self.std_errors
        upper = self.params + t_crit * self.std_errors
        
        return pd.DataFrame({
            f'{alpha/2:.3f}': lower,
            f'{1-alpha/2:.3f}': upper
        }, index=self.params.index)
    
    # ------------------------------------------------------------------
    # NOTE on cross-estimator tidy surface
    # ------------------------------------------------------------------
    # We intentionally do NOT add ``.estimate`` / ``.se`` / ``.pvalue`` /
    # ``.ci`` aliases to ``EconometricResults``: those names mean "the
    # single treatment effect" on :class:`CausalResult` (scalar) but would
    # have to mean "the coefficient vector" on multi-coef estimators
    # (Series) — and downstream code like ``if hasattr(r, 'estimate'):
    # est = r.estimate`` (workflow dispatch, agent serialisers) relies on
    # their absence to distinguish the two result types.  Adding the
    # aliases broke that dispatch in round-1 and is not worth the
    # ambiguity.
    #
    # The correct cross-estimator tidy surface is :meth:`tidy` (defined
    # below), which returns a long-format DataFrame on BOTH result types.
    # ------------------------------------------------------------------

    # ------------------------------------------------------------------
    # Broom-style tidy interface (EconometricResults)
    # ------------------------------------------------------------------

    def tidy(self, conf_level: float = 0.95) -> pd.DataFrame:
        """Return a long-format DataFrame of coefficients, broom-style.

        Columns
        -------
        term : str
            Variable / coefficient name.
        estimate : float
        std_error : float
        statistic : float
            t-statistic.
        p_value : float
        conf_low, conf_high : float
            Two-sided conf_level CI.

        Examples
        --------
        >>> result = sp.regress("y ~ x1 + x2", data=df)
        >>> result.tidy()
           term  estimate  std_error  statistic  p_value  conf_low  conf_high
        0  Intercept     ...

        See Also
        --------
        glance : 1-row model-level summary (R^2, F, N, AIC, BIC).
        """
        alpha = 1 - conf_level
        df_resid = self.data_info.get('df_resid', np.inf)
        stats = _scipy_stats()
        t_crit = stats.t.ppf(1 - alpha/2, df_resid)
        lo = self.params - t_crit * self.std_errors
        hi = self.params + t_crit * self.std_errors

        def _arr(x):
            return x.values if hasattr(x, 'values') else np.asarray(x)

        return pd.DataFrame({
            'term': list(self.params.index),
            'estimate': _arr(self.params),
            'std_error': _arr(self.std_errors),
            'statistic': _arr(self.tvalues),
            'p_value': _arr(self.pvalues),
            'conf_low': _arr(lo),
            'conf_high': _arr(hi),
        })

    def glance(self) -> pd.DataFrame:
        """Return a 1-row DataFrame of model-level statistics, broom-style.

        Columns (present subset depends on the model type)
        --------------------------------------------------
        nobs : int
        r_squared : float
        adj_r_squared : float
        f_statistic : float
        f_p_value : float
        aic, bic : float
        df_resid, df_model : int
        method : str
            Estimation method label.

        See Also
        --------
        tidy : long-format coefficient table.
        """
        g: Dict[str, Any] = {}
        g['method'] = self.model_info.get('method', self.model_info.get('model_type', ''))
        if 'nobs' in self.data_info:
            g['nobs'] = int(self.data_info['nobs'])
        key_map = {
            'R-squared': 'r_squared',
            'Adj R-squared': 'adj_r_squared',
            'F-statistic': 'f_statistic',
            'F p-value': 'f_p_value',
            'Prob (F-statistic)': 'f_p_value',
            'AIC': 'aic',
            'BIC': 'bic',
            'Log-Likelihood': 'log_likelihood',
        }
        for src, dst in key_map.items():
            if src in self.diagnostics and dst not in g:
                v = self.diagnostics[src]
                if isinstance(v, (int, float, np.integer, np.floating)):
                    g[dst] = float(v)
        if 'df_resid' in self.data_info:
            g['df_resid'] = int(self.data_info['df_resid'])
        if 'df_model' in self.data_info:
            g['df_model'] = int(self.data_info['df_model'])
        return pd.DataFrame([g])

    def predict(self, data: Optional[pd.DataFrame] = None) -> np.ndarray:
        """
        Generate predictions from the fitted model.

        Parameters
        ----------
        data : pd.DataFrame, optional
            New data for out-of-sample prediction. If None, returns
            in-sample fitted values.

        Returns
        -------
        np.ndarray
            Predicted values.
        """
        # In-sample: return fitted values if available
        if data is None:
            fv = self.data_info.get('fitted_values')
            if fv is not None:
                return fv
            raise NotImplementedError(
                "In-sample fitted values not stored. "
                "Pass data= for out-of-sample prediction."
            )

        # Out-of-sample: X @ params (works for simple linear models only)
        if self.params is not None and isinstance(self.params, pd.Series):
            var_names = list(self.params.index)
            has_intercept = 'Intercept' in var_names

            # Identify derived terms (interactions, categoricals, transforms)
            X_cols = [v for v in var_names if v != 'Intercept']
            missing = [c for c in X_cols if c not in data.columns]
            if missing:
                # Check if these are formula-derived terms
                import re
                derived = [c for c in missing
                           if ':' in c or '[' in c
                           or re.search(r'[()]', c)]
                if derived:
                    raise ValueError(
                        f"Out-of-sample prediction is not supported for "
                        f"models with formula transforms (found: "
                        f"{derived[:3]}{'...' if len(derived) > 3 else ''}). "
                        f"Re-fit using statsmodels directly, or use "
                        f"in-sample prediction with result.predict() "
                        f"(no arguments)."
                    )
                raise ValueError(
                    f"Prediction data missing column(s): {missing}"
                )
            X = data[X_cols].values.astype(float)
            if has_intercept:
                ones = np.ones((X.shape[0], 1))
                X = np.column_stack([ones, X])
                ordered = ['Intercept'] + X_cols
            else:
                ordered = X_cols
            beta = np.array([self.params[v] for v in ordered])
            return X @ beta

        raise NotImplementedError(
            "Prediction not available for this model type."
        )
    
    def residuals(self) -> Optional[np.ndarray]:
        """
        Return model residuals if available
        
        Returns
        -------
        np.ndarray or None
            Residuals
        """
        return self.data_info.get('residuals')
    
    def fitted_values(self) -> Optional[np.ndarray]:
        """
        Return fitted values if available
        
        Returns
        -------
        np.ndarray or None
            Fitted values
        """
        return self.data_info.get('fitted_values')

    def next_steps(self, print_result: bool = True) -> List[Dict[str, str]]:
        """
        Agent-native workflow guidance: what to do after fitting this model.

        Returns a list of recommended next steps — diagnostics, robustness
        checks, sensitivity analysis, and export options — tailored to the
        model type (OLS, IV, panel, etc.).

        Parameters
        ----------
        print_result : bool, default True
            Print formatted recommendations to stdout.

        Returns
        -------
        list of dict
            Each dict has keys: ``action``, ``reason``, ``priority``, ``category``.

        Examples
        --------
        >>> result = sp.regress("y ~ x1 + x2", data=df)
        >>> result.next_steps()
        """
        from .next_steps import econometric_next_steps, _format_steps
        steps = econometric_next_steps(self)
        if print_result:
            print(_format_steps(steps))
        return [s.to_dict() for s in steps]

    def _next_steps_html(self) -> str:
        from .next_steps import econometric_next_steps, _steps_repr_html
        return _steps_repr_html(econometric_next_steps(self))

    def violations(self) -> List[Dict[str, Any]]:
        """
        Agent-native structured list of assumption / diagnostic issues.

        Inspects stored diagnostics (first-stage F, standard error
        finiteness, …) and returns any flagged concerns as dicts with
        keys ``kind``, ``severity``, ``test``, ``value``, ``threshold``,
        ``message``, ``recovery_hint``, ``alternatives``.

        Returns
        -------
        list of dict
            Empty list if nothing flagged.

        Examples
        --------
        >>> result = sp.iv("y ~ (x ~ z) + c", data=df)
        >>> for v in result.violations():
        ...     if v['severity'] == 'error':
        ...         print(v['recovery_hint'])
        """
        from ._agent_summary import econometric_violations
        return econometric_violations(self)

    def to_agent_summary(self) -> Dict[str, Any]:
        """
        JSON-ready *nested* summary for agent consumption.

        Unlike ``summary()`` (prose for humans) and ``tidy()`` (long-form
        DataFrame), this returns a plain ``dict`` with coefficients,
        scalar diagnostics, violations, and recommended next steps —
        suitable for feeding into an LLM tool loop or logging.

        Returns
        -------
        dict
            Keys: ``kind``, ``model_type``, ``robust``, ``n_obs``,
            ``df_resid``, ``dependent_var``, ``coefficients``,
            ``diagnostics``, ``violations``, ``next_steps``.

        See Also
        --------
        to_dict :
            Canonical *flat* agent payload — prefer
            ``to_dict(detail="agent")`` for new code.
            ``to_agent_summary`` is kept because it surfaces a richer
            ``kind`` / ``model_type`` / ``robust`` triplet that ``to_dict``
            collapses into a single ``method`` field; two methods, two
            intentionally different shapes.

        Examples
        --------
        >>> result = sp.regress("y ~ x", data=df)
        >>> import json
        >>> agent_payload = json.dumps(result.to_agent_summary())
        """
        from ._agent_summary import econometric_agent_summary
        return econometric_agent_summary(self)

    def to_docx(self, filename: str, title: Optional[str] = None):
        """
        Export results to a Word (.docx) document.

        Parameters
        ----------
        filename : str
            Output path (.docx).
        title : str, optional
            Table title. Defaults to model type.
        """
        _result_to_docx(self, filename, title)

    # ------------------------------------------------------------------
    # Agent-native serialisation
    # ------------------------------------------------------------------

    def to_dict(self, *, detail: str = "standard") -> Dict[str, Any]:
        """Return a JSON-safe dict representation of the regression result.

        Parameters
        ----------
        detail : {"minimal", "standard", "agent"}, default ``"standard"``
            Payload depth, bounded by approximate token budget:

            - ``"minimal"`` (~ < 600 chars / < 150 tokens) — identity
              only: ``method``, ``model_type``, ``dependent_var``,
              ``n_obs``, plus ``fit_stats`` (R², F, AIC, BIC) when
              available.  No coefficient table.
            - ``"standard"`` (variable, ~ 50 chars × n_terms) — full
              coefficient table + diagnostics + glance row.  Matches
              the legacy ``to_dict()`` shape.
            - ``"agent"`` — standard + ``violations`` + ``warnings`` +
              ``next_steps`` + ``suggested_functions``.  Equivalent to
              legacy :meth:`for_agent` and the form returned by
              ``sp.agent.execute_tool`` and the MCP server.

        Returns
        -------
        dict
            JSON-safe and bounded — round-trips through ``json.dumps``.

        Notes
        -----
        Used by ``sp.agent.execute_tool`` to send results back to an
        LLM, and useful for caching / pickling-free persistence.
        """
        if detail not in ("minimal", "standard", "agent"):
            raise ValueError(
                "detail must be 'minimal', 'standard', or 'agent'; "
                f"got {detail!r}"
            )

        try:
            glance_row = self.glance().iloc[0].to_dict()
        except Exception:
            glance_row = {}

        base: Dict[str, Any] = {
            'method': str(self.model_info.get(
                'method', self.model_info.get('model_type', ''))),
            'model_type': str(self.model_info.get('model_type', '')),
            'dependent_var': _to_jsonable(
                self.data_info.get('dependent_var')),
            'n_obs': _to_jsonable(self.data_info.get('nobs')),
        }

        if detail == "minimal":
            # Compact subset of glance fit stats so agents can decide
            # whether to drill in (low R² → call diagnostics).
            fit_keys = (
                'r_squared', 'r.squared', 'r2',
                'adj_r_squared', 'adj.r.squared',
                'f_statistic', 'f.statistic',
                'aic', 'AIC', 'bic', 'BIC',
                'log_likelihood', 'logLik',
            )
            fit_stats: Dict[str, Any] = {}
            for k in fit_keys:
                if k in glance_row:
                    fit_stats[k] = _to_jsonable(glance_row[k])
            if fit_stats:
                base['fit_stats'] = fit_stats
            return base

        # standard: + full coefficient table + diagnostics + glance
        coefs: Dict[str, Dict[str, Any]] = {}
        try:
            terms = list(self.params.index)

            def _iget(x: Any, i: int) -> Any:
                """Positional get for Series or numpy array."""
                if x is None:
                    return None
                if isinstance(x, pd.Series):
                    return x.iloc[i]
                try:
                    return x[i]
                except (TypeError, IndexError, KeyError):
                    return None

            for i, term in enumerate(terms):
                coefs[str(term)] = {
                    'estimate': _to_jsonable(self.params.iloc[i]),
                    'std_error': _to_jsonable(self.std_errors.iloc[i]),
                    't_statistic': _to_jsonable(_iget(self.tvalues, i)),
                    'p_value': _to_jsonable(_iget(self.pvalues, i)),
                    'conf_low': _to_jsonable(self.conf_int_lower.iloc[i]),
                    'conf_high': _to_jsonable(self.conf_int_upper.iloc[i]),
                }
        except Exception:
            coefs = {}

        base.update({
            'coefficients': coefs,
            'diagnostics': _filter_jsonable_scalars(self.diagnostics),
            'glance': _to_jsonable(glance_row),
        })

        if detail == "standard":
            return base

        # agent: + violations + warnings + next_steps + suggested_functions
        try:
            viols = self.violations() or []
        except Exception:
            viols = []
        warns: List[str] = [
            v.get('message', '') for v in viols if v.get('message')
        ]

        try:
            steps = self.next_steps(print_result=False) or []
        except Exception:
            steps = []

        suggested: List[str] = []
        for s in steps:
            fn = s.get('suggest_function') or s.get('function')
            if fn and fn not in suggested:
                suggested.append(fn)
        for v in viols:
            for alt in v.get('alternatives', []) or []:
                if alt and alt not in suggested:
                    suggested.append(alt)

        base.update({
            'violations': _to_jsonable(viols),
            'warnings': warns,
            'next_steps': steps[:8],
            'suggested_functions': suggested,
        })
        return base

    def for_agent(self) -> Dict[str, Any]:
        """Agent-ready payload — alias for ``to_dict(detail="agent")``.

        Kept for backward compatibility with code written before the
        unified ``detail`` parameter.  New code should prefer
        ``to_dict(detail="agent")`` for explicit semantics.
        """
        return self.to_dict(detail="agent")

    def brief(self) -> str:
        """One-line dashboard status string (≤ ~120 chars).

        Surfaces the most-significant non-intercept coefficient so
        agents scanning a list of regressions can spot the active
        finding without paying a full ``to_dict`` round-trip.
        """
        from ..smart.brief import brief as _brief
        return _brief(self)

    def to_json(self, indent: Optional[int] = None) -> str:
        """Serialise :meth:`to_dict` via ``json.dumps``."""
        import json
        return json.dumps(self.to_dict(), indent=indent,
                          default=_to_jsonable)

    def _repr_html_(self) -> str:
        """Rich HTML display for Jupyter notebooks."""
        model_type = self.model_info.get('model_type', 'Unknown')
        method = self.model_info.get('method', '')
        dep_var = self.data_info.get('dependent_var', '')
        n_obs = self.data_info.get('nobs', '?')
        r2 = self.diagnostics.get('R-squared', None)
        f_stat = self.diagnostics.get('F-statistic', None)
        f_pv = self.diagnostics.get('F p-value', self.diagnostics.get('Prob (F-statistic)', None))

        def _safe(v: Any) -> str:
            return _html_escape(str(v))

        def _fmt(v: Any, spec: str = "") -> str:
            if isinstance(v, (int, float, np.integer, np.floating)) and not pd.isna(v):
                return format(v, spec)
            return _safe(v)

        def _s(pv):
            if pd.isna(pv): return ''
            if pv < 0.01: return '<span style="color:#E74C3C;">***</span>'
            if pv < 0.05: return '<span style="color:#E67E22;">**</span>'
            if pv < 0.1: return '<span style="color:#F39C12;">*</span>'
            return ''

        def _val(v):
            if isinstance(v, (int, float, np.integer, np.floating)) and not pd.isna(v):
                return f'{v:.4f}'
            return _safe(v)

        # CSS
        S = ('<style scoped>'
             '.sp-box{font-family:"Helvetica Neue",Arial,sans-serif;max-width:720px;border:1px solid #E5E7EB;border-radius:8px;overflow:hidden;margin:6px 0}'
             '.sp-hdr{background:linear-gradient(135deg,#1a1a2e 0%,#16213e 100%);color:#fff;padding:12px 16px}'
             '.sp-hdr h3{margin:0;font-size:15px;font-weight:600;letter-spacing:0.3px}'
             '.sp-hdr .sp-sub{font-size:11px;color:#94A3B8;margin-top:2px}'
             '.sp-metrics{display:flex;gap:0;border-bottom:1px solid #E5E7EB}'
             '.sp-metric{flex:1;padding:10px 14px;text-align:center;border-right:1px solid #E5E7EB}'
             '.sp-metric:last-child{border-right:none}'
             '.sp-metric .sp-val{font-size:18px;font-weight:700;color:#1a1a2e}'
             '.sp-metric .sp-lab{font-size:10px;color:#94A3B8;text-transform:uppercase;letter-spacing:0.5px;margin-top:2px}'
             'table.sp-coef{width:100%;border-collapse:collapse;font-size:12px}'
             'table.sp-coef th{padding:6px 10px;text-align:right;font-weight:600;color:#64748B;border-bottom:2px solid #E5E7EB;font-size:11px}'
             'table.sp-coef th:first-child{text-align:left}'
             'table.sp-coef td{padding:5px 10px;text-align:right;border-bottom:1px solid #F1F5F9}'
             'table.sp-coef td:first-child{text-align:left;font-weight:500;color:#1a1a2e}'
             'table.sp-coef tr:hover{background:#F8FAFC}'
             '.sp-diag{display:grid;grid-template-columns:1fr 1fr;gap:0;border-top:1px solid #E5E7EB;font-size:11px}'
             '.sp-diag-item{padding:4px 14px;display:flex;justify-content:space-between;border-bottom:1px solid #F8FAFC}'
             '.sp-diag-item:nth-child(odd){border-right:1px solid #F1F5F9}'
             '.sp-diag-k{color:#94A3B8}.sp-diag-v{color:#334155;font-weight:500}'
             '.sp-foot{padding:6px 14px;font-size:10px;color:#94A3B8;border-top:1px solid #E5E7EB;display:flex;justify-content:space-between}'
             '</style>')

        h = [S, '<div class="sp-box">']

        # --- Header ---
        sub_parts = []
        if dep_var:
            sub_parts.append(f'Y = {_safe(dep_var)}')
        if method:
            sub_parts.append(_safe(method))
        h.append(f'<div class="sp-hdr"><h3>{_safe(model_type)}</h3>')
        if sub_parts:
            h.append(f'<div class="sp-sub">{" · ".join(sub_parts)}</div>')
        h.append('</div>')

        # --- Key Metrics Bar ---
        h.append('<div class="sp-metrics">')
        if r2 is not None:
            h.append(f'<div class="sp-metric"><div class="sp-val">{_fmt(r2, ".4f")}</div><div class="sp-lab">R-squared</div></div>')
        if f_stat is not None:
            h.append(f'<div class="sp-metric"><div class="sp-val">{_fmt(f_stat, ".1f")}</div><div class="sp-lab">F-statistic</div></div>')
        h.append(f'<div class="sp-metric"><div class="sp-val">{_fmt(n_obs, ",")}</div><div class="sp-lab">Observations</div></div>')
        h.append(f'<div class="sp-metric"><div class="sp-val">{len(self.params)}</div><div class="sp-lab">Parameters</div></div>')
        h.append('</div>')

        # --- Coefficient Table ---
        h.append('<table class="sp-coef"><tr>')
        for col in ['', 'Coefficient', 'Std. Error', 't-stat', 'P&gt;|t|', '95% CI']:
            h.append(f'<th>{col}</th>')
        h.append('</tr>')

        for i, var in enumerate(self.params.index):
            coef = self.params.iloc[i]
            se = self.std_errors.iloc[i]
            t = self.tvalues.iloc[i] if isinstance(self.tvalues, pd.Series) else self.tvalues[i]
            pv = self.pvalues.iloc[i] if isinstance(self.pvalues, pd.Series) else self.pvalues[i]
            lo = self.conf_int_lower.iloc[i] if isinstance(self.conf_int_lower, pd.Series) else self.conf_int_lower[i]
            hi = self.conf_int_upper.iloc[i] if isinstance(self.conf_int_upper, pd.Series) else self.conf_int_upper[i]
            pv_color = '#DC2626' if pv < 0.01 else ('#EA580C' if pv < 0.05 else ('#D97706' if pv < 0.1 else '#64748B'))
            h.append(f'<tr><td>{_safe(var)}</td>')
            h.append(f'<td>{coef:.4f} {_s(pv)}</td>')
            h.append(f'<td style="color:#94A3B8;">({se:.4f})</td>')
            h.append(f'<td>{t:.2f}</td>')
            h.append(f'<td style="color:{pv_color};font-weight:600;">{pv:.4f}</td>')
            h.append(f'<td style="color:#94A3B8;">[{lo:.3f}, {hi:.3f}]</td></tr>')
        h.append('</table>')

        # --- Diagnostics Grid ---
        diag_items = [(k, v) for k, v in self.diagnostics.items()
                      if isinstance(v, (int, float, str)) and k not in ('R-squared', 'F-statistic', 'F p-value', 'Prob (F-statistic)')]
        if diag_items:
            h.append('<div class="sp-diag">')
            for k, v in diag_items:
                h.append(f'<div class="sp-diag-item"><span class="sp-diag-k">{_safe(k)}</span><span class="sp-diag-v">{_val(v)}</span></div>')
            h.append('</div>')

        # --- IV-specific: First-stage diagnostics ---
        iv_keys = [k for k in self.diagnostics if 'First-stage' in k or 'Hausman' in k or 'Partial' in k or 'Sargan' in k]
        if iv_keys:
            h.append('<details open style="border-top:1px solid #E5E7EB;"><summary style="padding:6px 14px;font-size:12px;'
                     'font-weight:600;color:#1a1a2e;cursor:pointer;">IV Diagnostics</summary>')
            h.append('<div class="sp-diag">')
            for k in iv_keys:
                v = self.diagnostics[k]
                h.append(f'<div class="sp-diag-item"><span class="sp-diag-k">{_safe(k)}</span><span class="sp-diag-v">{_val(v)}</span></div>')
            h.append('</div></details>')

        # --- Footer ---
        h.append(f'<div class="sp-foot"><span>N = {_fmt(n_obs, ",")}</span><span>* p&lt;0.1 &nbsp; ** p&lt;0.05 &nbsp; *** p&lt;0.01</span></div>')
        h.append('</div>')
        return '\n'.join(h)

    def __repr__(self) -> str:
        """String representation of results"""
        model_type = self.model_info.get('model_type', 'Unknown')
        n_params = len(self.params)
        n_obs = self.data_info.get('nobs', 'Unknown')
        return f"<EconometricResults: {model_type}, {n_params} parameters, {n_obs} observations>"

    def sensitivity(self, **kwargs):
        """Run the unified sensitivity dashboard on this result.

        See :func:`statspai.robustness.unified_sensitivity`.
        """
        from ..robustness.unified_sensitivity import unified_sensitivity
        # Expose a 1-entry "estimate" view for compatibility
        class _View:
            pass
        view = _View()
        view.estimate = float(self.params.iloc[0])
        view.se = float(self.std_errors.iloc[0])
        view.ci = (float(self.conf_int_lower.iloc[0]),
                   float(self.conf_int_upper.iloc[0]))
        view.params = self.params
        view.std_errors = self.std_errors
        return unified_sensitivity(view, **kwargs)


class CausalResult:
    """
    Unified result object for all causal inference methods in StatsPAI.

    All causal inference estimators (DID, RD, SCM, matching, etc.) return
    this object, providing a consistent interface for summaries, plots,
    and publication-quality output.

    Parameters
    ----------
    method : str
        Name of the estimation method (displayed in summary).
    estimand : str
        What is being estimated ('ATT', 'ATE', 'LATE').
    estimate : float
        Point estimate of the main treatment effect.
    se : float
        Standard error.
    pvalue : float
        Two-sided p-value for H0: effect = 0.
    ci : tuple of (float, float)
        Confidence interval (lower, upper).
    alpha : float
        Significance level used for CI.
    n_obs : int
        Number of observations.
    detail : pd.DataFrame, optional
        Detailed estimates (e.g., group-time ATTs).
    model_info : dict, optional
        Model metadata and aggregated results.
    _influence_funcs : np.ndarray, optional
        Influence function matrix (n_units, n_estimates) for joint inference.
    _citation_key : str, optional
        Key into the citation registry.
    """

    _CITATIONS: Dict[str, str] = {
        'did_2x2': (
            "@book{angrist2009mostly,\n"
            "  title={Mostly Harmless Econometrics: An Empiricist's Companion},\n"
            "  author={Angrist, Joshua D and Pischke, J{\\\"o}rn-Steffen},\n"
            "  year={2009},\n"
            "  publisher={Princeton University Press}\n"
            "}"
        ),
        'callaway_santanna': (
            "@article{callaway2021difference,\n"
            "  title={Difference-in-differences with multiple time periods},\n"
            "  author={Callaway, Brantly and Sant'Anna, Pedro H.C.},\n"
            "  journal={Journal of Econometrics},\n"
            "  volume={225},\n"
            "  number={2},\n"
            "  pages={200--230},\n"
            "  year={2021},\n"
            "  publisher={Elsevier}\n"
            "}"
        ),
        'sun_abraham': (
            "@article{sun2021estimating,\n"
            "  title={Estimating dynamic treatment effects in event studies "
            "with heterogeneous treatment effects},\n"
            "  author={Sun, Liyang and Abraham, Sarah},\n"
            "  journal={Journal of Econometrics},\n"
            "  volume={225},\n"
            "  number={2},\n"
            "  pages={175--199},\n"
            "  year={2021},\n"
            "  publisher={Elsevier}\n"
            "}"
        ),
        'rdrobust': (
            "@article{calonico2014robust,\n"
            "  title={Robust nonparametric confidence intervals for "
            "regression-discontinuity designs},\n"
            "  author={Calonico, Sebastian and Cattaneo, Matias D "
            "and Titiunik, Rocio},\n"
            "  journal={Econometrica},\n"
            "  volume={82},\n"
            "  number={6},\n"
            "  pages={2295--2326},\n"
            "  year={2014},\n"
            "  publisher={Wiley}\n"
            "}"
        ),
        'zubizarreta_2015_sbw': (
            "@article{zubizarreta2015stable,\n"
            "  title={Stable weights that balance covariates for "
            "estimation with incomplete outcome data},\n"
            "  author={Zubizarreta, Jos{\\'e} R},\n"
            "  journal={Journal of the American Statistical "
            "Association},\n"
            "  volume={110},\n"
            "  number={511},\n"
            "  pages={910--922},\n"
            "  year={2015},\n"
            "  publisher={Taylor \\& Francis}\n"
            "}"
        ),
        'bacon_decomposition': (
            "@article{goodmanbacon2021difference,\n"
            "  title={Difference-in-differences with variation in treatment timing},\n"
            "  author={Goodman-Bacon, Andrew},\n"
            "  journal={Journal of Econometrics},\n"
            "  volume={225},\n"
            "  number={2},\n"
            "  pages={254--277},\n"
            "  year={2021},\n"
            "  publisher={Elsevier}\n"
            "}"
        ),
        'did_multiplegt': (
            "@article{dechaisemartin2020two,\n"
            "  title={Two-Way Fixed Effects Estimators with "
            "Heterogeneous Treatment Effects},\n"
            "  author={de Chaisemartin, Cl{\\'e}ment and "
            "D'Haultf{\\oe}uille, Xavier},\n"
            "  journal={American Economic Review},\n"
            "  volume={110},\n"
            "  number={9},\n"
            "  pages={2964--2996},\n"
            "  year={2020}\n"
            "}"
        ),
        'stacked_did': (
            "@article{cengiz2019effect,\n"
            "  title={The Effect of Minimum Wages on Low-Wage Jobs},\n"
            "  author={Cengiz, Doruk and Dube, Arindrajit and "
            "Lindner, Attila and Zipperer, Ben},\n"
            "  journal={Quarterly Journal of Economics},\n"
            "  volume={134},\n"
            "  number={3},\n"
            "  pages={1405--1454},\n"
            "  year={2019},\n"
            "  publisher={Oxford University Press}\n"
            "}"
        ),
        'wooldridge_twfe': (
            "@unpublished{wooldridge2021two,\n"
            "  title={Two-Way Fixed Effects, the Two-Way Mundlak Regression, "
            "and Difference-in-Differences Estimators},\n"
            "  author={Wooldridge, Jeffrey M.},\n"
            "  year={2021},\n"
            "  note={Working paper, Michigan State University}\n"
            "}"
        ),
        'drdid': (
            "@article{santanna2020doubly,\n"
            "  title={Doubly Robust Difference-in-Differences Estimators},\n"
            "  author={Sant'Anna, Pedro H.C. and Zhao, Jun},\n"
            "  journal={Journal of Econometrics},\n"
            "  volume={219},\n"
            "  number={1},\n"
            "  pages={101--122},\n"
            "  year={2020},\n"
            "  publisher={Elsevier}\n"
            "}"
        ),
        'twfe_decomposition': (
            "@article{goodmanbacon2021difference,\n"
            "  title={Difference-in-differences with variation in treatment timing},\n"
            "  author={Goodman-Bacon, Andrew},\n"
            "  journal={Journal of Econometrics},\n"
            "  volume={225},\n"
            "  number={2},\n"
            "  pages={254--277},\n"
            "  year={2021},\n"
            "  publisher={Elsevier}\n"
            "}\n"
            "@article{dechaisemartin2020two,\n"
            "  title={Two-Way Fixed Effects Estimators with "
            "Heterogeneous Treatment Effects},\n"
            "  author={de Chaisemartin, Cl{\\'e}ment and "
            "D'Haultf{\\oe}uille, Xavier},\n"
            "  journal={American Economic Review},\n"
            "  volume={110},\n"
            "  number={9},\n"
            "  pages={2964--2996},\n"
            "  year={2020}\n"
            "}"
        ),
    }

    def __init__(
        self,
        method: str,
        estimand: str,
        estimate: float,
        se: float,
        pvalue: float,
        ci: tuple,
        alpha: float,
        n_obs: int,
        detail: Optional[pd.DataFrame] = None,
        model_info: Optional[Dict[str, Any]] = None,
        _influence_funcs: Optional[np.ndarray] = None,
        _citation_key: Optional[str] = None,
    ):
        self.method = method
        self.estimand = estimand
        self.estimate = estimate
        self.se = se
        self.pvalue = pvalue
        self.ci = ci
        self.alpha = alpha
        self.n_obs = n_obs
        self.detail = detail
        self.model_info = model_info or {}
        self._influence_funcs = _influence_funcs
        self._citation_key = _citation_key

    # ------------------------------------------------------------------
    # Backward compatibility with EconometricResults
    # ------------------------------------------------------------------

    @property
    def params(self) -> pd.Series:
        """Treatment effect as a params Series (for outreg2 compatibility)."""
        return pd.Series({self.estimand: self.estimate})

    @property
    def std_errors(self) -> pd.Series:
        return pd.Series({self.estimand: self.se})

    @property
    def tvalues(self) -> pd.Series:
        t = self.estimate / self.se if self.se > 0 else np.nan
        return pd.Series({self.estimand: t})

    @property
    def pvalues(self) -> pd.Series:
        return pd.Series({self.estimand: self.pvalue})

    @property
    def diagnostics(self) -> Dict[str, Any]:
        return self.model_info

    @property
    def data_info(self) -> Dict[str, Any]:
        return {'nobs': self.n_obs}

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _stars(pvalue: float) -> str:
        """Significance stars."""
        if pd.isna(pvalue):
            return ""
        if pvalue < 0.01:
            return "***"
        if pvalue < 0.05:
            return "**"
        if pvalue < 0.1:
            return "*"
        return ""

    # ------------------------------------------------------------------
    # Summary
    # ------------------------------------------------------------------

    def summary(self, alpha: Optional[float] = None) -> str:
        """
        Generate a formatted text summary of the causal estimation results.

        Parameters
        ----------
        alpha : float, optional
            Override significance level for display.

        Returns
        -------
        str
        """
        alpha = alpha or self.alpha
        lines: List[str] = []

        lines.append("=" * 78)
        lines.append(f"  {self.method}")
        lines.append("=" * 78)
        lines.append("")

        stars = self._stars(self.pvalue)
        lines.append(f"  {self.estimand}:      {self.estimate: .6f} {stars}")
        lines.append(f"  Std. Error:  ({self.se:.6f})")
        pct = int(100 * (1 - alpha))
        lines.append(f"  [{pct}% CI]:    [{self.ci[0]:.6f},  {self.ci[1]:.6f}]")
        lines.append(f"  P-value:     {self.pvalue:.4f}")
        lines.append("")

        # Event study coefficients — accept either the (relative_time, att)
        # convention emitted by did_imputation / callaway_santanna or the
        # (rel_time, estimate) convention emitted by wooldridge_did /
        # etwfe. Skip the block silently when neither pair is present so a
        # bare-minimum result object still summarises cleanly.
        if 'event_study' in self.model_info:
            es = self.model_info['event_study']
            time_col = (
                'relative_time' if 'relative_time' in es.columns
                else ('rel_time' if 'rel_time' in es.columns else None)
            )
            est_col = (
                'att' if 'att' in es.columns
                else ('estimate' if 'estimate' in es.columns else None)
            )
            if time_col and est_col and 'se' in es.columns:
                lines.append("-" * 78)
                lines.append("  Event Study Coefficients")
                lines.append("-" * 78)
                for _, row in es.iterrows():
                    e = int(row[time_col])
                    att = row[est_col]
                    se_v = row['se']
                    pv = row.get('pvalue', np.nan)
                    s = self._stars(pv)
                    lines.append(
                        f"  e = {e:>3d}  |  {att:>10.4f}  ({se_v:.4f}) {s}"
                    )
                lines.append("")

        # Detailed estimates
        if self.detail is not None and len(self.detail) > 0:
            if 'att' in self.detail.columns:
                # Causal inference format (group-time ATTs)
                cols = [c for c in ['group', 'time', 'att', 'se',
                                    'ci_lower', 'ci_upper', 'pvalue']
                        if c in self.detail.columns]
                title_str = "Group-Time ATT Estimates"
            elif 'method' in self.detail.columns and 'estimate' in self.detail.columns:
                # RD-style inference table
                cols = [c for c in ['method', 'estimate', 'se', 'z',
                                    'pvalue', 'ci_lower', 'ci_upper']
                        if c in self.detail.columns]
                title_str = "Inference"
            elif 'coefficient' in self.detail.columns:
                # Regression format (variable / coefficient / se)
                cols = [c for c in ['variable', 'coefficient', 'se',
                                    'tstat', 'pvalue']
                        if c in self.detail.columns]
                title_str = "Regression Coefficients"
            else:
                cols = list(self.detail.columns)
                title_str = "Detailed Estimates"
            lines.append("-" * 78)
            lines.append(f"  {title_str}")
            lines.append("-" * 78)
            lines.append(
                self.detail[cols].to_string(index=False, float_format='%.4f')
            )
            lines.append("")

        # Pre-trend test
        if 'pretrend_test' in self.model_info:
            pt = self.model_info['pretrend_test']
            lines.append("-" * 78)
            lines.append(
                f"  Pre-trend Test: chi2({pt['df']}) = {pt['statistic']:.4f}, "
                f"p-value = {pt['pvalue']:.4f}"
            )
            lines.append("")

        # Model info footer
        lines.append("-" * 78)
        lines.append(f"  Observations:    {self.n_obs:,}")
        _skip = {'event_study', 'pretrend_test', 'aggregations',
                 'cohort_sizes', 'influence_funcs_matrix',
                 'conventional', 'robust'}
        for key, val in self.model_info.items():
            if key in _skip or isinstance(val, (pd.DataFrame, np.ndarray,
                                                dict, list)):
                continue
            label = key.replace('_', ' ').title()
            lines.append(f"  {label}:    {val}")
        lines.append("=" * 78)
        lines.append("  * p<0.1, ** p<0.05, *** p<0.01")

        return SummaryText("\n".join(lines))

    # ------------------------------------------------------------------
    # Broom-style tidy interface (CausalResult)
    # ------------------------------------------------------------------

    def tidy(self, conf_level: Optional[float] = None) -> pd.DataFrame:
        """Long-format coefficient table, broom-compatible.

        Primary row is the overall estimand (ATT / ATE / LATE).  If the
        result carries ``detail`` with group-time ATTs (CS/SA) or event
        study coefficients, those rows are appended with a ``type``
        column distinguishing ``'main'``, ``'group_time'``, or
        ``'event_study'``.

        Columns
        -------
        term, estimate, std_error, statistic, p_value,
        conf_low, conf_high, type

        Examples
        --------
        >>> r = sp.callaway_santanna(df, y='y', g='g', t='t', i='i')
        >>> r.tidy()                  # includes group-time ATTs
        >>> r.tidy().query("type=='main'")
        """
        # Main row
        alpha = 1 - conf_level if conf_level is not None else self.alpha
        if conf_level is not None and abs(conf_level - (1 - self.alpha)) > 1e-9:
            # Recompute CI at requested conf_level. Prefer t-distribution
            # when df_resid is recorded (small-sample correctness);
            # fall back to normal only when no df is available (e.g. CS /
            # synth results where influence-function SEs are asymptotic).
            df_resid = self.model_info.get('df_resid', None)
            stats = _scipy_stats()
            if df_resid is not None and np.isfinite(df_resid) and df_resid > 0:
                crit = stats.t.ppf(1 - alpha/2, df_resid)
            else:
                crit = stats.norm.ppf(1 - alpha/2)
            lo, hi = (self.estimate - crit * self.se,
                      self.estimate + crit * self.se)
        else:
            lo, hi = self.ci
        main_row = {
            'term': self.estimand,
            'estimate': self.estimate,
            'std_error': self.se,
            'statistic': self.estimate / self.se if self.se > 0 else np.nan,
            'p_value': self.pvalue,
            'conf_low': lo,
            'conf_high': hi,
            'type': 'main',
        }
        rows = [main_row]

        # Event study coefficients
        es = self.model_info.get('event_study')
        if isinstance(es, pd.DataFrame) and len(es) > 0:
            for _, r in es.iterrows():
                e = r.get('relative_time')
                att = r.get('att')
                se = r.get('se')
                pv = r.get('pvalue', np.nan)
                lo_r = r.get('ci_lower', att - 1.96 * se if pd.notna(att) else np.nan)
                hi_r = r.get('ci_upper', att + 1.96 * se if pd.notna(att) else np.nan)
                rows.append({
                    'term': f'event_{int(e):+d}' if pd.notna(e) else 'event',
                    'estimate': att, 'std_error': se,
                    'statistic': att / se if pd.notna(se) and se > 0 else np.nan,
                    'p_value': pv,
                    'conf_low': lo_r, 'conf_high': hi_r,
                    'type': 'event_study',
                })

        # Group-time ATTs
        if self.detail is not None and 'att' in getattr(self.detail, 'columns', []):
            for _, r in self.detail.iterrows():
                g = r.get('group', '')
                t = r.get('time', '')
                rows.append({
                    'term': f'att(g={g},t={t})',
                    'estimate': r.get('att'),
                    'std_error': r.get('se'),
                    'statistic': (r.get('att') / r.get('se')
                                  if pd.notna(r.get('se')) and r.get('se', 0) > 0
                                  else np.nan),
                    'p_value': r.get('pvalue', np.nan),
                    'conf_low': r.get('ci_lower', np.nan),
                    'conf_high': r.get('ci_upper', np.nan),
                    'type': 'group_time',
                })

        return pd.DataFrame(rows)

    def glance(self) -> pd.DataFrame:
        """1-row model-level summary, broom-compatible.

        Columns
        -------
        method, estimand, estimate, std_error, p_value, nobs,
        conf_low, conf_high, alpha, n_groups (if applicable),
        n_periods (if applicable), pretrend_pvalue (if applicable).
        """
        g: Dict[str, Any] = {
            'method': self.method,
            'estimand': self.estimand,
            'estimate': self.estimate,
            'std_error': self.se,
            'p_value': self.pvalue,
            'conf_low': self.ci[0],
            'conf_high': self.ci[1],
            'nobs': int(self.n_obs) if pd.notna(self.n_obs) else np.nan,
            'alpha': self.alpha,
        }
        # Method-specific extras
        if 'n_groups' in self.model_info:
            g['n_groups'] = int(self.model_info['n_groups'])
        if 'n_periods' in self.model_info:
            g['n_periods'] = int(self.model_info['n_periods'])
        pt = self.model_info.get('pretrend_test')
        if isinstance(pt, dict) and 'pvalue' in pt:
            g['pretrend_pvalue'] = float(pt['pvalue'])
        # SC-specific
        if 'pre_treatment_rmse' in self.model_info:
            g['pre_treatment_rmse'] = float(self.model_info['pre_treatment_rmse'])
        # RD-specific
        if 'bandwidth_h' in self.model_info:
            try:
                g['bandwidth'] = float(self.model_info['bandwidth_h'])
            except (TypeError, ValueError):
                pass
        return pd.DataFrame([g])

    # ------------------------------------------------------------------
    # Plotting
    # ------------------------------------------------------------------

    def event_study_plot(
        self,
        ax=None,
        title: Optional[str] = None,
        color: str = '#2C3E50',
        ci_alpha: float = 0.15,
        figsize: tuple = (10, 6),
        **kwargs,
    ):
        """
        Plot event study coefficients with confidence intervals.

        Parameters
        ----------
        ax : matplotlib.axes.Axes, optional
        title : str, optional
        color : str
        ci_alpha : float
        figsize : tuple

        Returns
        -------
        (fig, ax)
        """
        try:
            import matplotlib.pyplot as plt
        except ImportError:
            raise ImportError(
                "matplotlib required for plotting. "
                "Install: pip install matplotlib"
            )

        if 'event_study' not in self.model_info:
            raise ValueError(
                "No event study estimates. Use a staggered DID method."
            )

        es = self.model_info['event_study'].copy()

        if ax is None:
            fig, ax = plt.subplots(figsize=figsize)
        else:
            fig = ax.get_figure()

        e = es['relative_time'].values
        att = es['att'].values
        lo = es['ci_lower'].values
        hi = es['ci_upper'].values

        ax.fill_between(e, lo, hi, alpha=ci_alpha, color=color)
        ax.scatter(e, att, color=color, s=40, zorder=5)
        ax.plot(e, att, color=color, linewidth=1, alpha=0.7, zorder=4)
        ax.errorbar(
            e, att,
            yerr=[att - lo, hi - att],
            fmt='none', color=color, capsize=3, linewidth=1, zorder=3,
        )

        ax.axhline(y=0, color='gray', linestyle='--', linewidth=0.8)
        ax.axvline(
            x=-0.5, color='#E74C3C', linestyle=':',
            linewidth=1, alpha=0.5, label='Treatment onset',
        )

        ax.set_xlabel('Periods Relative to Treatment', fontsize=11)
        ax.set_ylabel('Estimated Effect', fontsize=11)
        ax.set_title(title or f'Event Study: {self.method}', fontsize=13)
        ax.tick_params(labelsize=10)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.legend(fontsize=9, frameon=False)
        fig.tight_layout()
        return fig, ax

    def plot(self, type: str = 'auto', **kwargs):
        """
        Generate appropriate visualisation based on model type.

        Parameters
        ----------
        type : str
            'auto' (recommended), 'event_study', 'coefplot',
            'trajectory', 'gap', 'both', 'weights', 'placebo',
            'placebo_gap', 'conformal', 'staggered', 'factors',
            'balance', 'density', 'original', 'pointwise',
            'cumulative', 'all'.

        Returns
        -------
        (fig, ax) or (fig, axes)

        Notes
        -----
        In 'auto' mode, the plot type is selected by method:

        - DID/event study → event study plot
        - Synthetic Control (all variants) → trajectory + gap (both)
        - Conformal SCM → conformal period-level CI plot
        - Staggered SCM → cohort-level ATT comparison
        - Causal Impact → 3-panel (original + pointwise + cumulative)
        - Matching → Love plot (covariate balance)
        - Neural causal estimators → CATE distribution
        - RD → coefplot (use ``rdplot()`` for binned scatter)
        - Other → coefplot
        """
        method_lower = self.method.lower()

        # All synth-related plot types handled by unified synthplot
        _synth_types = {
            'trajectory', 'gap', 'both', 'weights', 'placebo',
            'placebo_gap', 'placebo_dist', 'conformal', 'staggered',
            'factors', 'loadings', 'compare',
        }

        if type == 'auto':
            # Event study (DID)
            if 'event_study' in self.model_info:
                return self.event_study_plot(**kwargs)

            # Synthetic Control — detect ALL variants
            if self._is_synth_result():
                from ..synth.plots import synthplot
                # Pick best auto type based on variant
                if 'period_results' in self.model_info:
                    return synthplot(self, type='conformal', **kwargs)
                if 'cohort_effects' in self.model_info:
                    return synthplot(self, type='staggered', **kwargs)
                return synthplot(self, type='both', **kwargs)

            # Causal Impact
            if 'causal impact' in method_lower or \
               'intervention_time' in self.model_info:
                from ..causal_impact.impact import impactplot
                return impactplot(self, type='all', **kwargs)

            # Matching (has balance/SMD table)
            if self.detail is not None and 'smd' in getattr(
                    self.detail, 'columns', []):
                from ..matching.match import balanceplot
                return balanceplot(self, **kwargs)

            if self.model_info.get('neural_causal'):
                from ..neural_causal.plots import neural_causal_plot
                return neural_causal_plot(self, type='cate', **kwargs)

            return self._coefplot(**kwargs)

        # Explicit type overrides
        if type == 'event_study':
            return self.event_study_plot(**kwargs)
        if type in _synth_types:
            from ..synth.plots import synthplot
            return synthplot(self, type=type, **kwargs)
        if type in ('original', 'pointwise', 'cumulative', 'all'):
            from ..causal_impact.impact import impactplot
            return impactplot(self, type=type, **kwargs)
        if type == 'balance':
            from ..matching.match import balanceplot
            return balanceplot(self, **kwargs)
        if type in ('cate', 'ite', 'effects', 'propensity', 'overlap', 'loss'):
            from ..neural_causal.plots import neural_causal_plot
            return neural_causal_plot(self, type=type, **kwargs)
        return self._coefplot(**kwargs)

    def _is_synth_result(self) -> bool:
        """Check if this result is from any synthetic control variant."""
        mi = self.model_info
        # Direct markers from various synth variants
        synth_keys = {
            'gap_table',        # classic, demeaned, robust
            'Y_obs',            # sdid
            'trajectory',       # gsynth
            'factors_pre',      # gsynth
            'cohort_effects',   # staggered
            'period_results',   # conformal
        }
        if synth_keys & set(mi.keys()):
            return True
        # Augmented SCM
        if mi.get('model_type', '').startswith('Synthetic'):
            return True
        # Method name check
        m = self.method.lower()
        return any(kw in m for kw in (
            'synthetic', 'synth', 'sdid', 'gsynth', 'staggered',
            'conformal', 'augmented', 'ascm', 'demeaned', 'de-meaned',
            'de-trended', 'unconstrained', 'factor',
        ))

    def _coefplot(self, ax=None, figsize=(8, 5), **kwargs):
        try:
            import matplotlib.pyplot as plt
        except ImportError:
            raise ImportError("matplotlib required for plotting")

        if ax is None:
            fig, ax = plt.subplots(figsize=figsize)
        else:
            fig = ax.get_figure()

        ax.errorbar(
            0, self.estimate,
            yerr=[[self.estimate - self.ci[0]], [self.ci[1] - self.estimate]],
            fmt='o', color='#2C3E50', capsize=5, markersize=8,
        )
        ax.axhline(y=0, color='gray', linestyle='--', linewidth=0.8)
        ax.set_xlim(-1, 1)
        ax.set_xticks([0])
        ax.set_xticklabels([self.estimand])
        ax.set_ylabel('Estimated Effect')
        ax.set_title(f'{self.method}: {self.estimand}')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        fig.tight_layout()
        return fig, ax

    # ------------------------------------------------------------------
    # Export
    # ------------------------------------------------------------------

    def to_markdown(self, path: Optional[str] = None,
                    digits: int = 4) -> str:
        """Render the causal result as Markdown.

        Neural causal results delegate to the richer neural exporter,
        which includes unit-level effects and training diagnostics.
        Other causal results render the broom-style tidy table plus
        scalar diagnostics.
        """
        if self.model_info.get('neural_causal'):
            from ..neural_causal.exports import neural_causal_to_markdown
            return neural_causal_to_markdown(self, path=path, digits=digits)

        tidy = self.tidy().round(digits)
        glance = self.glance().round(digits)
        parts = [
            f"# {self.method}",
            "",
            "## Estimates",
            tidy.to_markdown(index=False),
            "",
            "## Summary",
            glance.to_markdown(index=False),
        ]
        if self.detail is not None and len(self.detail) > 0:
            parts.extend([
                "",
                "## Detail",
                self.detail.round(digits).to_markdown(index=False),
            ])
        text = "\n".join(parts) + "\n"
        if path is not None:
            from pathlib import Path
            Path(path).write_text(text, encoding="utf-8")
        return text

    def to_html(self, path: Optional[str] = None,
                digits: int = 4) -> str:
        """Render the causal result as an HTML report."""
        if self.model_info.get('neural_causal'):
            from ..neural_causal.exports import neural_causal_to_html
            return neural_causal_to_html(self, path=path, digits=digits)

        tidy = self.tidy().round(digits)
        glance = self.glance().round(digits)
        blocks = [
            "<html><body>",
            f"<h1>{_html_escape(self.method)}</h1>",
            "<h2>Estimates</h2>",
            tidy.to_html(index=False),
            "<h2>Summary</h2>",
            glance.to_html(index=False),
        ]
        if self.detail is not None and len(self.detail) > 0:
            blocks.extend([
                "<h2>Detail</h2>",
                self.detail.round(digits).to_html(index=False),
            ])
        blocks.append("</body></html>")
        html = "\n".join(blocks)
        if path is not None:
            from pathlib import Path
            Path(path).write_text(html, encoding="utf-8")
        return html

    def to_excel(self, path: str, digits: int = 6) -> str:
        """Write a multi-sheet Excel workbook for the causal result."""
        if self.model_info.get('neural_causal'):
            from ..neural_causal.exports import neural_causal_to_excel
            return neural_causal_to_excel(self, path, digits=digits)

        with pd.ExcelWriter(path) as writer:
            self.tidy().round(digits).to_excel(
                writer, sheet_name="Estimates", index=False
            )
            self.glance().round(digits).to_excel(
                writer, sheet_name="Summary", index=False
            )
            if self.detail is not None and len(self.detail) > 0:
                self.detail.round(digits).to_excel(
                    writer, sheet_name="Detail", index=False
                )
            diagnostics = _filter_jsonable_scalars(self.model_info)
            if diagnostics:
                pd.DataFrame([diagnostics]).to_excel(
                    writer, sheet_name="Diagnostics", index=False
                )
        return path

    def to_word(
        self,
        path: str,
        digits: int = 4,
        caption: Optional[str] = None,
    ) -> str:
        """Write a Word (``.docx``) report for the causal result.

        Produces a publication-style three-block document:

        1. Title (``caption`` or ``"<method> Results"``).
        2. Estimates table (variables × coefficient/SE/t/p/CI), formatted
           with stars (``* p<0.1, ** p<0.05, *** p<0.01``) and Times New
           Roman 9-10pt typography matching :file:`output/_aer_style.py`.
        3. Detail table (group/time/ATT or estimator-specific) if
           ``self.detail`` is non-empty.
        4. Trailing notes paragraph: ``"Standard errors in parentheses.
           Observations: N. Method: <method>."``.

        Requires ``python-docx``.
        """
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError as e:  # pragma: no cover
            raise ImportError(
                "python-docx required for to_word(). "
                "Install: pip install python-docx"
            ) from e
        from ..output._aer_style import (
            apply_word_booktab_rules,
            style_word_table_typography,
            add_word_notes_paragraph,
        )

        doc = Document()
        title = caption or f"{self.method} — {self.estimand} Estimates"
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _write_df_to_table(df: pd.DataFrame) -> None:
            n_rows = len(df) + 1
            n_cols = len(df.columns) + 1
            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            table.rows[0].cells[0].text = ""
            for j, col in enumerate(df.columns, 1):
                table.rows[0].cells[j].text = str(col)
            for i, (idx, row) in enumerate(df.iterrows()):
                table.rows[i + 1].cells[0].text = str(idx)
                for j, val in enumerate(row, 1):
                    if isinstance(val, float):
                        cell_text = f"{val:.{digits}f}" if np.isfinite(val) else ""
                    elif pd.isna(val):
                        cell_text = ""
                    else:
                        cell_text = str(val)
                    table.rows[i + 1].cells[j].text = cell_text
            style_word_table_typography(
                table, header_rows=(0,),
                header_pt=10, body_pt=9,
                align_first_col="left", align_data_cols="center",
            )
            apply_word_booktab_rules(table, header_top_idx=0, header_bot_idx=0)

        # Estimates block — always present
        tidy = self.tidy().round(digits).set_index(
            self.tidy().columns[0]
        ) if "term" in self.tidy().columns or "estimand" in self.tidy().columns else \
            self.tidy().round(digits)
        _write_df_to_table(tidy)

        # Detail block — group/time ATTs, etc.
        if self.detail is not None and len(self.detail) > 0:
            doc.add_paragraph().add_run("Detail").bold = True
            detail_view = self.detail.round(digits)
            if not detail_view.index.is_unique or detail_view.index.equals(
                pd.RangeIndex(len(detail_view))
            ):
                detail_view = detail_view.reset_index(drop=True)
            _write_df_to_table(detail_view)

        notes = (
            "Standard errors in parentheses. "
            "* p<0.1, ** p<0.05, *** p<0.01. "
            f"Observations: {self.n_obs:,}. Method: {self.method}."
        )
        add_word_notes_paragraph(doc, notes)
        doc.save(path)
        return path

    def to_latex(self, caption: Optional[str] = None,
                 label: Optional[str] = None) -> str:
        """Generate a LaTeX table of the results."""
        caption = caption or f'{self.method} Results'
        label = label or 'tab:causal_result'

        lines = [
            '\\begin{table}[htbp]',
            '\\centering',
            f'\\caption{{{caption}}}',
            f'\\label{{{label}}}',
        ]

        if self.detail is not None and len(self.detail) > 0:
            # Detect table format: causal (group/time/att) vs regression (variable/coefficient)
            if 'att' in self.detail.columns:
                cols = [c for c in ['group', 'time', 'att', 'se', 'pvalue']
                        if c in self.detail.columns]
                coef_col, star_col = 'att', 'att'
            elif 'method' in self.detail.columns and 'estimate' in self.detail.columns:
                cols = [c for c in ['method', 'estimate', 'se', 'pvalue']
                        if c in self.detail.columns]
                coef_col, star_col = 'estimate', 'estimate'
            elif 'coefficient' in self.detail.columns:
                cols = [c for c in ['variable', 'coefficient', 'se', 'pvalue']
                        if c in self.detail.columns]
                coef_col, star_col = 'coefficient', 'coefficient'
            else:
                cols = list(self.detail.columns)
                coef_col, star_col = None, None

            n_cols = len(cols)
            spec = 'l' + 'c' * (n_cols - 1)
            lines.append(f'\\begin{{tabular}}{{{spec}}}')
            lines.append('\\hline\\hline')
            hdr = {'group': 'Group', 'time': 'Time', 'att': 'ATT',
                   'se': 'Std.\\ Error', 'pvalue': 'P-value',
                   'variable': 'Variable', 'coefficient': 'Coefficient',
                   'tstat': 't-stat'}
            lines.append(' & '.join(hdr.get(c, c) for c in cols) + ' \\\\')
            lines.append('\\hline')
            for _, row in self.detail.iterrows():
                vals = []
                for c in cols:
                    v = row[c]
                    if isinstance(v, float):
                        s = self._stars(row.get('pvalue', np.nan)) if c == star_col else ''
                        vals.append(f'{v:.4f}{s}')
                    else:
                        vals.append(str(int(v)) if isinstance(v, (int, np.integer)) else str(v))
                lines.append(' & '.join(vals) + ' \\\\')
        else:
            lines.append('\\begin{tabular}{lc}')
            lines.append('\\hline\\hline')
            lines.append(
                f'{self.estimand} & '
                f'{self.estimate:.4f}{self._stars(self.pvalue)} \\\\'
            )
            lines.append(f'& ({self.se:.4f}) \\\\')

        lines += [
            '\\hline',
            f'Observations & {self.n_obs:,} \\\\',
            '\\hline\\hline',
            '\\end{tabular}',
            '\\begin{tablenotes}',
            '\\footnotesize',
            '\\item Standard errors in parentheses.',
            '\\item * p<0.1, ** p<0.05, *** p<0.01',
            '\\end{tablenotes}',
            '\\end{table}',
        ]
        return '\n'.join(lines)

    def cite(self, format: str = "bibtex") -> Any:
        """Return the canonical citation for this estimator.

        Parameters
        ----------
        format : {"bibtex", "apa", "json"}, default ``"bibtex"``
            - ``"bibtex"`` — full ``@article{...}`` / ``@book{...}``
              entry (the form used in :file:`paper.bib`).
            - ``"apa"`` — APA-style prose for inline use:
              ``"Callaway, B., & Sant'Anna, P. H. C. (2021). Title.
              Journal, vol(num), pages."``.
            - ``"json"`` — structured payload with parsed
              ``authors`` / ``year`` / ``title`` / ``journal`` / ...
              for agent consumption.

        Returns
        -------
        str | dict
            ``"bibtex"`` / ``"apa"`` → ``str``; ``"json"`` → ``dict``.

        Notes
        -----
        Citations are zero-hallucination per the project's policy
        (CLAUDE.md §10): every entry comes from the hand-curated
        ``_CITATIONS`` table on the result class, which mirrors
        :file:`paper.bib`. APA / JSON forms are derived by parsing
        that single source — never by generating new bibliographic
        facts.
        """
        key = self._citation_key or self.method.lower().replace(' ', '_')
        bibtex: Optional[str] = None
        if key in self._CITATIONS:
            bibtex = self._CITATIONS[key]
        else:
            for k, v in self._CITATIONS.items():
                if k in key or key in k:
                    bibtex = v
                    break
        if bibtex is None:
            placeholder = (
                f"% No citation registered for method: {self.method}"
            )
            if format == "bibtex":
                return placeholder
            if format == "apa":
                return placeholder
            if format == "json":
                return {"type": None, "key": None, "authors": [],
                        "fields": {}, "raw": placeholder,
                        "note": "no citation registered"}
            raise ValueError(
                f"format must be 'bibtex', 'apa' or 'json'; "
                f"got {format!r}"
            )
        from ..smart.citations import render_citation
        return render_citation(bibtex, fmt=format)

    def pretrend_test(self) -> Dict[str, Any]:
        """Return pre-trend test results (DID methods)."""
        if 'pretrend_test' not in self.model_info:
            raise ValueError("Pre-trend test not available for this method.")
        return self.model_info['pretrend_test']

    def next_steps(self, print_result: bool = True) -> List[Dict[str, str]]:
        """
        Agent-native workflow guidance: what to do after this causal analysis.

        Returns method-specific recommendations — pre-trend tests for DID,
        McCrary test for RD, balance checks for matching, etc.

        Parameters
        ----------
        print_result : bool, default True
            Print formatted recommendations to stdout.

        Returns
        -------
        list of dict
            Each dict has keys: ``action``, ``reason``, ``priority``, ``category``.

        Examples
        --------
        >>> result = sp.did(df, y='wage', treat='treated', time='post')
        >>> result.next_steps()
        """
        from .next_steps import causal_next_steps, _format_steps
        steps = causal_next_steps(self)
        if print_result:
            print(_format_steps(steps))
        return [s.to_dict() for s in steps]

    def _next_steps_html(self) -> str:
        from .next_steps import causal_next_steps, _steps_repr_html
        return _steps_repr_html(causal_next_steps(self))

    def violations(self) -> List[Dict[str, Any]]:
        """
        Agent-native structured list of assumption / diagnostic issues.

        Inspects the diagnostics the estimator already stored on
        ``model_info`` (pre-trend p-value, first-stage F, McCrary,
        rhat / ESS / divergences, overlap, balance, …) and returns
        flagged items as dicts with ``kind`` / ``severity`` / ``test``
        / ``value`` / ``threshold`` / ``message`` / ``recovery_hint``
        / ``alternatives``.

        Returns
        -------
        list of dict
            Empty list if nothing flagged. ``severity`` is one of
            ``"error"`` / ``"warning"`` / ``"info"``.

        Examples
        --------
        >>> r = sp.did(df, y='wage', treat='treated', time='post')
        >>> [v['test'] for v in r.violations() if v['severity'] == 'error']
        ['pretrend']
        """
        from ._agent_summary import causal_violations
        return causal_violations(self)

    def to_agent_summary(self) -> Dict[str, Any]:
        """
        JSON-ready *nested* summary for agent consumption.

        Returns a plain dict with point estimate, CI, scalar diagnostics,
        violations (via :meth:`violations`), and recommended next steps
        (via :meth:`next_steps`) — everything an agent needs to decide
        the next action without re-reading the prose summary.

        Returns
        -------
        dict
            Keys: ``kind``, ``method``, ``method_family``, ``estimand``,
            ``point`` (``estimate``/``se``/``pvalue``/``ci``/``alpha``),
            ``n_obs``, ``diagnostics``, ``violations``, ``next_steps``,
            ``citation_key``.

        See Also
        --------
        to_dict :
            Canonical *flat* agent payload — prefer
            ``to_dict(detail="agent")`` for new code.
            ``to_agent_summary`` is kept because it groups the point
            estimate under a ``point`` sub-dict (handy for templating
            tables) and surfaces a ``method_family`` field that ``to_dict``
            does not. Two methods, two intentionally different shapes.

        Examples
        --------
        >>> r = sp.did(df, y='wage', treat='treated', time='post')
        >>> import json
        >>> print(json.dumps(r.to_agent_summary(), indent=2, default=str))
        """
        from ._agent_summary import causal_agent_summary
        return causal_agent_summary(self)

    def to_dict(self, *, detail_head: int = 5,
                detail: str = "standard") -> Dict[str, Any]:
        """Return a JSON-safe flat dict representation of the causal result.

        Parameters
        ----------
        detail_head : int, default 5
            Rows of ``self.detail`` to include in the ``"standard"`` and
            ``"agent"`` levels (0 to omit). Ignored when
            ``detail="minimal"``.
        detail : {"minimal", "standard", "agent"}, default ``"standard"``
            Payload depth, bounded by approximate token budget so agents
            running on token-metered APIs can pick the right level:

            - ``"minimal"`` (~ < 600 chars / < 150 tokens) — bare answer:
              ``method``, ``estimand``, ``estimate``, ``se``, ``pvalue``,
              ``ci``, ``alpha``, ``n_obs``, ``citation_key``.  No
              diagnostics, no detail rows.  For sub-step calls where
              the agent only needs the point estimate to decide what
              to do next.
            - ``"standard"`` (~ < 4 000 chars / < 1 000 tokens) — the
              legacy default: minimal + scalar diagnostics +
              ``detail_head`` rows.
            - ``"agent"`` (~ < 8 000 chars / < 2 000 tokens) — standard
              + ``violations`` + ``warnings`` + ``next_steps`` +
              ``suggested_functions``.  This is what
              ``sp.agent.execute_tool`` and the MCP server emit by
              default so a tool-using agent gets violations + workflow
              hints in one round-trip.  Equivalent to legacy
              :meth:`for_agent`.

        Returns
        -------
        dict
            JSON-safe and bounded — round-trips through ``json.dumps``.

        Notes
        -----
        :meth:`to_agent_summary` returns a *nested* payload (``point`` /
        ``violations`` / …).  ``to_dict`` is the flat form used by the
        agent / MCP layer.
        """
        if detail not in ("minimal", "standard", "agent"):
            raise ValueError(
                "detail must be 'minimal', 'standard', or 'agent'; "
                f"got {detail!r}"
            )

        ci = self.ci
        if (ci is not None
                and not isinstance(ci, (pd.Series, pd.DataFrame))
                and hasattr(ci, '__len__') and len(ci) == 2):
            ci_out: Optional[List[Optional[float]]] = [
                _to_jsonable(ci[0]),
                _to_jsonable(ci[1]),
            ]
        else:
            ci_out = None

        out: Dict[str, Any] = {
            'method': str(self.method),
            'estimand': str(self.estimand),
            'estimate': _to_jsonable(self.estimate),
            'se': _to_jsonable(self.se),
            'pvalue': _to_jsonable(self.pvalue),
            'ci': ci_out,
            'alpha': _to_jsonable(self.alpha),
            'n_obs': _to_jsonable(self.n_obs),
            'citation_key': _to_jsonable(self._citation_key),
        }

        if detail == "minimal":
            return out

        # standard: + scalar diagnostics + detail_head rows
        out['diagnostics'] = _filter_jsonable_scalars(self.model_info)
        if detail_head and self.detail is not None:
            try:
                head = self.detail.head(int(detail_head))
                out['detail_head'] = _to_jsonable(
                    head.to_dict(orient='records'))
            except Exception:
                pass

        if detail == "standard":
            return out

        # agent: + violations + warnings + next_steps + suggested_functions
        try:
            viols = self.violations() or []
        except Exception:
            viols = []
        warns: List[str] = [
            v.get('message', '') for v in viols if v.get('message')
        ]

        try:
            steps = self.next_steps(print_result=False) or []
        except Exception:
            steps = []

        suggested: List[str] = []
        for s in steps:
            fn = s.get('suggest_function') or s.get('function')
            if fn and fn not in suggested:
                suggested.append(fn)
        for v in viols:
            for alt in v.get('alternatives', []) or []:
                if alt and alt not in suggested:
                    suggested.append(alt)

        out.update({
            'violations': _to_jsonable(viols),
            'warnings': warns,
            'next_steps': steps[:8],
            'suggested_functions': suggested,
        })
        return out

    def for_agent(self, detail_head: int = 5) -> Dict[str, Any]:
        """Agent-ready payload — alias for ``to_dict(detail="agent")``.

        Kept for backward compatibility with code written before the
        unified ``detail`` parameter.  New code should prefer
        ``to_dict(detail="agent")`` for explicit semantics.
        """
        return self.to_dict(detail_head=detail_head, detail="agent")

    def brief(self) -> str:
        """One-line dashboard status string (≤ ~120 chars).

        Cheaper than ``summary()`` (multi-line prose) and
        ``to_dict(detail="minimal")`` (JSON ~ 300 chars). Use in
        agent dashboards or ``for r in results: print(r.brief())``.
        """
        from ..smart.brief import brief as _brief
        return _brief(self)

    def to_json(self, indent: Optional[int] = None,
                 detail_head: int = 5) -> str:
        """Serialise :meth:`to_dict` via ``json.dumps``."""
        import json
        return json.dumps(
            self.to_dict(detail_head=detail_head),
            indent=indent,
            default=_to_jsonable,
        )

    def to_docx(self, filename: str, title: Optional[str] = None):
        """
        Export results to a Word (.docx) document.

        Parameters
        ----------
        filename : str
            Output path (.docx).
        title : str, optional
            Table title. Defaults to method name.
        """
        _result_to_docx(self, filename, title)

    def _repr_html_(self) -> str:
        """Rich HTML display for Jupyter notebooks — model-specific layouts."""
        mi = self.model_info
        pct = int(100 * (1 - self.alpha))
        stars_raw = self._stars(self.pvalue)

        def _safe(v: Any) -> str:
            return _html_escape(str(v))

        def _fmt(v: Any, spec: str = "") -> str:
            if isinstance(v, (int, float, np.integer, np.floating)) and not pd.isna(v):
                return format(v, spec)
            return _safe(v)

        def _td(v: Any) -> str:
            if isinstance(v, (int, float, np.integer, np.floating)) and not pd.isna(v):
                return f'<td>{v:.4f}</td>'
            return f'<td>{_safe(v)}</td>'

        def _s(pv):
            if pd.isna(pv): return ''
            if pv < 0.01: return '<span style="color:#DC2626">***</span>'
            if pv < 0.05: return '<span style="color:#EA580C">**</span>'
            if pv < 0.1: return '<span style="color:#D97706">*</span>'
            return ''

        # Significance-based accent color
        if self.pvalue < 0.01:
            accent, accent_bg = '#059669', '#ECFDF5'  # green
        elif self.pvalue < 0.05:
            accent, accent_bg = '#2563EB', '#EFF6FF'  # blue
        elif self.pvalue < 0.1:
            accent, accent_bg = '#D97706', '#FFFBEB'  # amber
        else:
            accent, accent_bg = '#64748B', '#F8FAFC'  # gray

        # Shared CSS (scoped)
        S = ('<style scoped>'
             '.sp-box{font-family:"Helvetica Neue",Arial,sans-serif;max-width:720px;'
             'border:1px solid #E5E7EB;border-radius:8px;overflow:hidden;margin:6px 0}'
             '.sp-hdr{background:linear-gradient(135deg,#1a1a2e 0%,#16213e 100%);color:#fff;padding:12px 16px}'
             '.sp-hdr h3{margin:0;font-size:15px;font-weight:600;letter-spacing:0.3px}'
             '.sp-hdr .sp-sub{font-size:11px;color:#94A3B8;margin-top:2px}'
             '.sp-effect{display:flex;align-items:center;gap:16px;padding:14px 16px;border-bottom:1px solid #E5E7EB}'
             '.sp-effect-num{font-size:28px;font-weight:700;letter-spacing:-0.5px}'
             '.sp-effect-meta{font-size:12px;color:#64748B;line-height:1.6}'
             '.sp-effect-badge{display:inline-block;padding:2px 8px;border-radius:10px;font-size:10px;font-weight:600;letter-spacing:0.3px}'
             '.sp-metrics{display:flex;gap:0;border-bottom:1px solid #E5E7EB}'
             '.sp-metric{flex:1;padding:8px 12px;text-align:center;border-right:1px solid #E5E7EB}'
             '.sp-metric:last-child{border-right:none}'
             '.sp-metric .sp-val{font-size:15px;font-weight:700;color:#1a1a2e}'
             '.sp-metric .sp-lab{font-size:9px;color:#94A3B8;text-transform:uppercase;letter-spacing:0.5px;margin-top:1px}'
             '.sp-section{border-top:1px solid #E5E7EB}'
             '.sp-section summary{padding:8px 14px;font-size:12px;font-weight:600;color:#1a1a2e;cursor:pointer;'
             'list-style:none;display:flex;align-items:center;gap:6px}'
             '.sp-section summary::before{content:"\\25B6";font-size:8px;color:#94A3B8;transition:transform 0.2s}'
             '.sp-section[open] summary::before{transform:rotate(90deg)}'
             'table.sp-tbl{width:100%;border-collapse:collapse;font-size:11px}'
             'table.sp-tbl th{padding:4px 10px;text-align:right;font-weight:600;color:#64748B;border-bottom:1px solid #E5E7EB;font-size:10px}'
             'table.sp-tbl th:first-child{text-align:left}'
             'table.sp-tbl td{padding:4px 10px;text-align:right;border-bottom:1px solid #F8FAFC}'
             'table.sp-tbl td:first-child{text-align:left;font-weight:500;color:#334155}'
             'table.sp-tbl tr:hover{background:#F8FAFC}'
             '.sp-grid{display:grid;grid-template-columns:1fr 1fr;gap:0;font-size:11px}'
             '.sp-grid-item{padding:4px 14px;display:flex;justify-content:space-between;border-bottom:1px solid #F8FAFC}'
             '.sp-grid-item:nth-child(odd){border-right:1px solid #F1F5F9}'
             '.sp-gk{color:#94A3B8}.sp-gv{color:#334155;font-weight:500}'
             '.sp-bar{height:6px;border-radius:3px;margin-top:2px}'
             '.sp-foot{padding:6px 14px;font-size:10px;color:#94A3B8;border-top:1px solid #E5E7EB;display:flex;justify-content:space-between}'
             '</style>')

        h = [S, '<div class="sp-box">']

        # ── Header ──
        h.append(f'<div class="sp-hdr"><h3>{_safe(self.method)}</h3>')
        sub_parts = [self.estimand]
        if mi.get('rd_type'):
            sub_parts.append(f'{mi["rd_type"]} RD')
        elif mi.get('distance'):
            sub_parts.append(f'{mi["distance"]} {mi.get("method", "")}')
        sub = _safe(" · ".join(str(p) for p in sub_parts if p is not None))
        h.append(f'<div class="sp-sub">{sub}</div></div>')

        # ── Main Effect Card ──
        sig_label = '< 0.01' if self.pvalue < 0.01 else ('< 0.05' if self.pvalue < 0.05 else ('< 0.10' if self.pvalue < 0.1 else f'= {self.pvalue:.3f}'))
        h.append(f'<div class="sp-effect" style="background:{accent_bg};">')
        h.append(f'<div class="sp-effect-num" style="color:{accent};">{self.estimate:.4f}</div>')
        h.append(f'<div class="sp-effect-meta">')
        h.append(f'<span class="sp-effect-badge" style="background:{accent};color:white;">{stars_raw or "n.s."}</span> &nbsp; p {sig_label}<br>')
        h.append(f'SE = {self.se:.4f} &nbsp;&nbsp; {pct}% CI [{self.ci[0]:.4f}, {self.ci[1]:.4f}]')
        h.append(f'</div></div>')

        # ── Model-Specific Metric Bars ──
        h.append('<div class="sp-metrics">')
        h.append(f'<div class="sp-metric"><div class="sp-val">{_fmt(self.n_obs, ",")}</div><div class="sp-lab">Observations</div></div>')

        if self._is_synth_result():
            # SC metrics
            for key, label in [('n_donors', 'Donors'), ('n_pre_periods', 'Pre-periods'), ('n_post_periods', 'Post-periods')]:
                if key in mi:
                    h.append(f'<div class="sp-metric"><div class="sp-val">{_fmt(mi[key])}</div><div class="sp-lab">{label}</div></div>')
            if 'pre_treatment_rmse' in mi:
                h.append(f'<div class="sp-metric"><div class="sp-val">{_fmt(mi["pre_treatment_rmse"], ".3f")}</div><div class="sp-lab">Pre-RMSE</div></div>')
        elif mi.get('rd_type') is not None:
            # RD metrics
            for key, label in [('n_effective_left', 'N Left (eff.)'), ('n_effective_right', 'N Right (eff.)'),
                               ('bandwidth_h', 'Bandwidth')]:
                if key in mi:
                    v = mi[key]
                    vs = _fmt(v, ".3f")
                    h.append(f'<div class="sp-metric"><div class="sp-val">{vs}</div><div class="sp-lab">{label}</div></div>')
        elif self.detail is not None and 'smd' in getattr(self.detail, 'columns', []):
            # Matching metrics
            for key, label in [('n_treated', 'Treated'), ('n_control', 'Control'), ('n_matches', 'Matches')]:
                if key in mi:
                    h.append(f'<div class="sp-metric"><div class="sp-val">{_fmt(mi[key])}</div><div class="sp-lab">{label}</div></div>')
        h.append('</div>')

        # ── SC: Donor Weights ──
        if self._is_synth_result() and self.detail is not None and 'weight' in getattr(self.detail, 'columns', []):
            weights_df = self.detail[self.detail['weight'] > 0.001].sort_values('weight', ascending=False)
            if len(weights_df) > 0:
                max_w = weights_df['weight'].max()
                h.append('<details class="sp-section" open><summary>Donor Weights</summary>')
                h.append('<div style="padding:4px 14px 8px;">')
                for _, row in weights_df.iterrows():
                    unit_name = row.get('unit', row.iloc[0])
                    w = row['weight']
                    pct_w = (w / max_w) * 100
                    h.append(f'<div style="display:flex;align-items:center;gap:8px;margin:3px 0;font-size:11px;">'
                             f'<span style="width:60px;color:#334155;font-weight:500;">Unit {_safe(unit_name)}</span>'
                             f'<div style="flex:1;background:#F1F5F9;border-radius:3px;height:8px;">'
                             f'<div class="sp-bar" style="width:{pct_w:.0f}%;background:{accent};"></div></div>'
                             f'<span style="width:50px;text-align:right;color:#64748B;">{w:.3f}</span></div>')
                h.append('</div></details>')

        # ── SC: Gap Table ──
        if 'gap_table' in mi and isinstance(mi['gap_table'], pd.DataFrame):
            gap = mi['gap_table']
            h.append('<details class="sp-section"><summary>Period-by-Period Effects</summary>')
            h.append('<table class="sp-tbl"><tr>')
            for c in gap.columns:
                h.append(f'<th>{_safe(c)}</th>')
            h.append('</tr>')
            for _, row in gap.iterrows():
                h.append('<tr>')
                for c in gap.columns:
                    h.append(_td(row[c]))
                h.append('</tr>')
            h.append('</table></details>')

        # ── RD: Conventional vs Robust Inference ──
        if self.detail is not None and 'method' in getattr(self.detail, 'columns', []) and 'estimate' in getattr(self.detail, 'columns', []):
            h.append('<details class="sp-section" open><summary>Inference Comparison</summary>')
            h.append('<table class="sp-tbl"><tr>')
            for col in ['Method', 'Estimate', 'Std. Err.', 'z', 'p-value', 'CI']:
                h.append(f'<th>{col}</th>')
            h.append('</tr>')
            for _, row in self.detail.iterrows():
                meth = row.get('method', '')
                est = row.get('estimate', np.nan)
                se_v = row.get('se', np.nan)
                z_v = row.get('z', np.nan)
                pv = row.get('pvalue', np.nan)
                lo = row.get('ci_lower', np.nan)
                hi = row.get('ci_upper', np.nan)
                pvc = '#DC2626' if pv < 0.01 else ('#EA580C' if pv < 0.05 else ('#D97706' if pv < 0.1 else '#64748B'))
                bold = 'font-weight:600;' if 'Robust' in str(meth) else ''
                h.append(f'<tr style="{bold}">')
                h.append(f'<td>{_safe(meth)}</td><td>{est:.4f} {_s(pv)}</td>')
                h.append(f'<td style="color:#94A3B8;">({se_v:.4f})</td><td>{z_v:.2f}</td>')
                h.append(f'<td style="color:{pvc};font-weight:600;">{pv:.4f}</td>')
                h.append(f'<td style="color:#94A3B8;">[{lo:.4f}, {hi:.4f}]</td></tr>')
            h.append('</table></details>')

        # ── RD: Design Parameters ──
        if mi.get('rd_type') is not None:
            h.append('<details class="sp-section"><summary>Design Parameters</summary><div class="sp-grid">')
            rd_params = [('cutoff', 'Cutoff'), ('polynomial_p', 'Poly Order (p)'), ('polynomial_q', 'Bias Poly (q)'),
                         ('kernel', 'Kernel'), ('bwselect', 'BW Selection'), ('bandwidth_h', 'Bandwidth (h)'),
                         ('bandwidth_b', 'Bias BW (b)'), ('n_left', 'N Left'), ('n_right', 'N Right'),
                         ('n_effective_left', 'N Eff. Left'), ('n_effective_right', 'N Eff. Right')]
            for key, label in rd_params:
                if key in mi:
                    v = mi[key]
                    vs = _fmt(v, ".4f")
                    h.append(f'<div class="sp-grid-item"><span class="sp-gk">{label}</span><span class="sp-gv">{vs}</span></div>')
            h.append('</div></details>')

        # ── Matching: Covariate Balance ──
        if self.detail is not None and 'smd' in getattr(self.detail, 'columns', []):
            h.append('<details class="sp-section" open><summary>Covariate Balance</summary>')
            h.append('<table class="sp-tbl"><tr>')
            for col in ['Variable', 'Treated', 'Control', 'SMD', '']:
                h.append(f'<th>{col}</th>')
            h.append('</tr>')
            for _, row in self.detail.iterrows():
                var = row.get('variable', row.iloc[0])
                mt = row.get('mean_treated', np.nan)
                mc = row.get('mean_control', np.nan)
                smd = row.get('smd', np.nan)
                smd_abs = abs(smd) if not pd.isna(smd) else 0
                bar_color = '#059669' if smd_abs < 0.1 else ('#D97706' if smd_abs < 0.25 else '#DC2626')
                bar_w = min(smd_abs / 0.5 * 100, 100)
                h.append(f'<tr><td>{_safe(var)}</td>')
                h.append(f'<td>{mt:.2f}</td><td>{mc:.2f}</td>')
                h.append(f'<td style="color:{bar_color};font-weight:600;">{smd:.3f}</td>')
                h.append(f'<td style="width:80px;"><div style="background:#F1F5F9;border-radius:3px;height:6px;">'
                         f'<div style="width:{bar_w:.0f}%;height:6px;border-radius:3px;background:{bar_color};"></div>'
                         f'</div></td></tr>')
            h.append('</table>')
            # Balance threshold annotation
            h.append('<div style="padding:4px 10px;font-size:10px;color:#94A3B8;">'
                     'SMD: <span style="color:#059669">|d|&lt;0.1 balanced</span> · '
                     '<span style="color:#D97706">0.1-0.25 borderline</span> · '
                     '<span style="color:#DC2626">&gt;0.25 imbalanced</span></div>')
            h.append('</details>')

        # ── Matching: Design info ──
        if mi.get('distance') and 'smd' in getattr(self.detail, 'columns', []) if self.detail is not None else False:
            h.append('<details class="sp-section"><summary>Matching Parameters</summary><div class="sp-grid">')
            match_params = [('distance', 'Distance'), ('method', 'Method'), ('estimand', 'Estimand'),
                            ('n_treated', 'N Treated'), ('n_control', 'N Control'), ('n_matches', 'Matches'),
                            ('caliper', 'Caliper'), ('replace', 'With Replacement'), ('bias_correction', 'Bias Correction')]
            for key, label in match_params:
                if key in mi:
                    v = mi[key]
                    vs = _fmt(v, ".4f")
                    h.append(f'<div class="sp-grid-item"><span class="sp-gk">{label}</span><span class="sp-gv">{vs}</span></div>')
            h.append('</div></details>')

        # ── Event study coefficients ──
        if 'event_study' in mi:
            es = mi['event_study']
            h.append('<details class="sp-section" open><summary>Event Study Coefficients</summary>')
            h.append('<table class="sp-tbl"><tr>')
            for col in ['Period', 'ATT', 'Std. Err.', 'CI', 'p-value']:
                h.append(f'<th>{col}</th>')
            h.append('</tr>')
            for _, row in es.iterrows():
                e = int(row['relative_time'])
                att = row['att']
                se_v = row['se']
                pv = row.get('pvalue', np.nan)
                lo = row.get('ci_lower', np.nan)
                hi = row.get('ci_upper', np.nan)
                pvc = '#DC2626' if pv < 0.01 else ('#EA580C' if pv < 0.05 else ('#D97706' if pv < 0.1 else '#64748B'))
                bg = 'background:#FFFBEB;' if e == 0 else ''
                h.append(f'<tr style="{bg}">')
                h.append(f'<td style="font-weight:600;">e = {e}</td>')
                h.append(f'<td>{att:.4f} {_s(pv)}</td>')
                h.append(f'<td style="color:#94A3B8;">({se_v:.4f})</td>')
                ci_str = f'[{lo:.4f}, {hi:.4f}]' if not pd.isna(lo) else ''
                h.append(f'<td style="color:#94A3B8;">{ci_str}</td>')
                pv_str = f'{pv:.4f}' if not pd.isna(pv) else ''
                h.append(f'<td style="color:{pvc};font-weight:600;">{pv_str}</td></tr>')
            h.append('</table></details>')

        # ── Generic detail table (group-time ATTs, etc.) ──
        if (self.detail is not None and len(self.detail) > 0
                and 'event_study' not in mi
                and 'smd' not in getattr(self.detail, 'columns', [])
                and not ('method' in getattr(self.detail, 'columns', []) and 'estimate' in getattr(self.detail, 'columns', []))
                and not ('weight' in getattr(self.detail, 'columns', []) and self._is_synth_result())):
            if 'att' in self.detail.columns:
                cols = [c for c in ['group', 'time', 'att', 'se', 'pvalue'] if c in self.detail.columns]
                title_str = 'Group-Time ATT Estimates'
            else:
                cols = list(self.detail.columns)
                title_str = 'Detail'
            h.append(f'<details class="sp-section"><summary>{_safe(title_str)}</summary>')
            h.append('<table class="sp-tbl"><tr>')
            for c in cols:
                h.append(f'<th>{_safe(c)}</th>')
            h.append('</tr>')
            max_rows = 20
            for _, row in self.detail[cols].head(max_rows).iterrows():
                h.append('<tr>')
                for c in cols:
                    h.append(_td(row[c]))
                h.append('</tr>')
            if len(self.detail) > max_rows:
                h.append(f'<tr><td colspan="{len(cols)}" style="text-align:center;color:#94A3B8;">... {len(self.detail)-max_rows} more</td></tr>')
            h.append('</table></details>')

        # ── SC: Model metadata ──
        if self._is_synth_result():
            h.append('<details class="sp-section"><summary>Model Parameters</summary><div class="sp-grid">')
            sc_params = [('treatment_time', 'Treatment Time'), ('treated_unit', 'Treated Unit'),
                         ('penalization', 'Penalization'), ('pre_treatment_mspe', 'Pre-MSPE'),
                         ('pre_treatment_rmse', 'Pre-RMSE'), ('n_placebos', 'Placebo Tests')]
            for key, label in sc_params:
                if key in mi:
                    v = mi[key]
                    vs = _fmt(v, ".4f")
                    h.append(f'<div class="sp-grid-item"><span class="sp-gk">{label}</span><span class="sp-gv">{vs}</span></div>')
            h.append('</div></details>')

        # ── Footer ──
        h.append(f'<div class="sp-foot"><span>N = {_fmt(self.n_obs, ",")}</span><span>* p&lt;0.1 &nbsp; ** p&lt;0.05 &nbsp; *** p&lt;0.01</span></div>')
        h.append('</div>')
        return '\n'.join(h)

    def __repr__(self) -> str:
        s = self._stars(self.pvalue)
        return (
            f"<CausalResult: {self.method}, {self.estimand} = "
            f"{self.estimate:.4f}{s}, SE = {self.se:.4f}, n = {self.n_obs:,}>"
        )

    def __str__(self) -> str:
        return self.summary()

    def sensitivity(self, **kwargs):
        """Run the unified sensitivity dashboard on this result.

        See :func:`statspai.robustness.unified_sensitivity`.
        """
        from ..robustness.unified_sensitivity import unified_sensitivity
        return unified_sensitivity(self, **kwargs)


# ======================================================================
# Shared Word export helper
# ======================================================================

def _result_to_docx(result, filename: str, title: Optional[str] = None):
    """
    Export a single EconometricResults or CausalResult to Word (.docx).

    Produces a one-model table with coefficients, SEs, stars, and
    diagnostics in APA format.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError(
            "python-docx required for Word export. "
            "Install: pip install python-docx"
        )

    if not filename.endswith('.docx'):
        filename += '.docx'

    doc = Document()

    # Title
    if title is None:
        title = getattr(result, 'method', None) or result.model_info.get('model_type', 'Results')
    p = doc.add_paragraph()
    run = p.add_run(str(title))
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Build rows from result
    params = result.params
    std_errors = result.std_errors
    pvalues = result.pvalues if hasattr(result, 'pvalues') else None

    def _stars(pv):
        if pv is None or (isinstance(pv, float) and np.isnan(pv)):
            return ''
        if pv < 0.01: return '***'
        if pv < 0.05: return '**'
        if pv < 0.1: return '*'
        return ''

    # Table: Variable | Coefficient | SE
    rows_data = []
    if isinstance(params, pd.Series):
        for var in params.index:
            coef = params[var]
            se = std_errors[var] if var in std_errors.index else np.nan
            if pvalues is not None and isinstance(pvalues, pd.Series) and var in pvalues.index:
                pv = float(pvalues[var])
            else:
                pv = np.nan
            rows_data.append((str(var), f'{coef:.4f}{_stars(pv)}', f'({se:.4f})'))

    n_rows = len(rows_data) * 2 + 1  # coef rows + SE rows + header
    # Actually: each variable gets 2 rows (coef, SE)
    table = doc.add_table(rows=1 + len(rows_data) * 2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    table.rows[0].cells[0].text = ''
    table.rows[0].cells[1].text = title
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data
    row_idx = 1
    for var, coef_str, se_str in rows_data:
        # Coefficient row
        table.rows[row_idx].cells[0].text = var
        table.rows[row_idx].cells[1].text = coef_str
        for cell in table.rows[row_idx].cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if cell == table.rows[row_idx].cells[1] else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
        row_idx += 1

        # SE row
        table.rows[row_idx].cells[0].text = ''
        table.rows[row_idx].cells[1].text = se_str
        for para in table.rows[row_idx].cells[1].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
        row_idx += 1

    # Diagnostics as additional paragraph
    diag = getattr(result, 'diagnostics', {})
    n_obs = None
    if hasattr(result, 'data_info') and isinstance(result.data_info, dict):
        n_obs = result.data_info.get('nobs')
    if hasattr(result, 'n_obs'):
        n_obs = result.n_obs

    diag_lines = []
    if n_obs is not None:
        diag_lines.append(f'Observations: {n_obs:,}')
    if isinstance(diag, dict):
        for k, v in diag.items():
            if isinstance(v, (int, float)) and not isinstance(v, bool):
                diag_lines.append(f'{k}: {v:.4f}' if isinstance(v, float) else f'{k}: {v:,}')

    if diag_lines:
        dp = doc.add_paragraph()
        for line in diag_lines:
            run = dp.add_run(line + '\n')
            run.font.size = Pt(8)

    # Significance note
    np_ = doc.add_paragraph()
    run = np_.add_run('* p<0.1, ** p<0.05, *** p<0.01')
    run.italic = True
    run.font.size = Pt(8)

    doc.save(filename)
