"""
Publication-quality regression tables and balance tables.

Provides ``regtable()`` for unified regression output across formats,
and ``mean_comparison()`` for balance / summary statistics tables.

Usage
-----
>>> import statspai as sp
>>> m1 = sp.regress("y ~ x1", data=df)
>>> m2 = sp.regress("y ~ x1 + x2", data=df)
>>> sp.regtable(m1, m2)
>>> sp.regtable(m1, m2, output="latex", filename="table1.tex")
>>>
>>> sp.mean_comparison(df, variables=["age", "income"], group="treated")
"""

from __future__ import annotations

import warnings
from collections import OrderedDict
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple, Union  # noqa: F401

import numpy as np
import pandas as pd
from scipy import stats as sp_stats


# ---------------------------------------------------------------------------
# Re-use extraction helpers from estimates module
# ---------------------------------------------------------------------------

from .estimates import (
    _ci_bounds,
    _extract_model_data,
    _ModelData,
    _format_stars,
    _fmt_val,
    _fmt_int,
    _latex_escape,
    _html_escape,
    _STAT_ALIASES,
    _STAT_DISPLAY,
)
from ._diagnostics import extract_diagnostic_rows
from ._journals import get_template, list_templates, star_note_for
from ._repro import build_repro_note


# Bracket styles cycled through when rendering ``multi_se`` extra SE rows.
# The four bracket pairs are deliberately Markdown-safe: a fourth ``||`` pair
# would collide with GFM pipe-table delimiters and break the row, so we use
# guillemets ``«»`` instead. The primary SE always uses parentheses.
_MULTI_SE_BRACKETS = (("[", "]"), ("{", "}"), ("⟨", "⟩"), ("«", "»"))


def _hc_recompute_se(model_data: _ModelData, raw_result: Any, vcov: str) -> Optional[Dict[str, Any]]:
    """Recompute HC0/HC1/HC2/HC3 SE/t/p/CI for an OLS-style result.

    Returns ``{'std_errors', 'tvalues', 'pvalues', 'conf_int_lower',
    'conf_int_upper'}`` keyed by variable name, or ``None`` when the
    raw result lacks the required ``data_info`` keys (X / residuals).

    Algorithm:

    - HC0: Omega = diag(e_i^2)
    - HC1: HC0 * n / (n - k)        (Stata ``robust`` default)
    - HC2: Omega = diag(e_i^2 / (1 - h_ii))
    - HC3: Omega = diag(e_i^2 / (1 - h_ii)^2)

    The leverage diagonal ``h_ii`` is computed from
    ``H_ij = X_ij · ((X'X)^{-1} X')_ji`` without forming the full
    n×n hat matrix.
    """
    dinfo = getattr(raw_result, "data_info", None)
    if not isinstance(dinfo, dict):
        return None
    X = dinfo.get("X")
    e = dinfo.get("residuals")
    if X is None or e is None:
        return None
    X_arr = np.asarray(X, dtype=float)
    e_arr = np.asarray(e, dtype=float).ravel()
    if X_arr.ndim != 2 or X_arr.shape[0] != e_arr.shape[0]:
        return None
    n, k = X_arr.shape
    if n <= k:
        return None
    XtX_inv = np.linalg.pinv(X_arr.T @ X_arr)
    if vcov.upper() in ("HC2", "HC3"):
        # h_ii = sum_j X_ij * (X · XtX_inv)_ij
        H_full = X_arr @ XtX_inv  # n×k
        h_ii = np.einsum("ij,ij->i", X_arr, H_full)
        h_ii = np.clip(h_ii, 0.0, 1.0 - 1e-10)
    if vcov.upper() == "HC0":
        omega = e_arr ** 2
    elif vcov.upper() in ("HC1", "ROBUST"):
        omega = (e_arr ** 2) * (n / (n - k))
    elif vcov.upper() == "HC2":
        omega = (e_arr ** 2) / (1.0 - h_ii)
    elif vcov.upper() == "HC3":
        omega = (e_arr ** 2) / ((1.0 - h_ii) ** 2)
    else:
        raise ValueError(
            f"vcov={vcov!r} not supported for OLS recompute. Choose "
            f"from 'HC0', 'HC1' (Stata robust), 'HC2', 'HC3', 'robust'."
        )
    XtOX = X_arr.T @ (omega[:, None] * X_arr)
    V = XtX_inv @ XtOX @ XtX_inv
    se_arr = np.sqrt(np.maximum(np.diag(V), 0.0))
    var_names = list(model_data.params.index)
    if len(var_names) != len(se_arr):
        # var_cov ordering vs params ordering mismatch; use data_info if present
        dvn = dinfo.get("var_names")
        if dvn is not None and len(dvn) == len(se_arr):
            var_names = list(dvn)
        else:
            return None
    se_series = pd.Series(se_arr, index=var_names)
    b = pd.Series(np.asarray(model_data.params, dtype=float), index=var_names)
    with np.errstate(divide="ignore", invalid="ignore"):
        t_arr = b.values / se_arr
    tvalues = pd.Series(t_arr, index=var_names)
    df_resid = getattr(model_data, "df_resid", None)
    if df_resid is None or not np.isfinite(df_resid) or df_resid <= 0:
        df_resid = float(n - k)
    pvalues = pd.Series(
        2.0 * (1.0 - sp_stats.t.cdf(np.abs(t_arr), df_resid)),
        index=var_names,
    )
    crit = sp_stats.t.ppf(0.975, df_resid)
    ci_lo = pd.Series(b.values - crit * se_arr, index=var_names)
    ci_hi = pd.Series(b.values + crit * se_arr, index=var_names)
    return {
        "std_errors": se_series,
        "tvalues": tvalues,
        "pvalues": pvalues,
        "conf_int_lower": ci_lo,
        "conf_int_upper": ci_hi,
    }


def _apply_vcov_to_panels(
    panels: List["_PanelData"],
    flat_results: List[Any],
    vcov: str,
) -> None:
    """Mutate each ``_ModelData`` in ``panels`` to substitute the recomputed
    HC SEs / t / p / CI bounds. No-op for results lacking the required
    ``data_info`` keys (other than emitting a UserWarning so users know
    the column was left unchanged).
    """
    flat_idx = 0
    skipped: List[int] = []
    for panel in panels:
        for md in panel.models:
            raw = flat_results[flat_idx]
            recompute = _hc_recompute_se(md, raw, vcov)
            if recompute is None:
                skipped.append(flat_idx + 1)
                flat_idx += 1
                continue
            object.__setattr__(md, "std_errors", recompute["std_errors"])
            object.__setattr__(md, "tvalues", recompute["tvalues"])
            object.__setattr__(md, "pvalues", recompute["pvalues"])
            object.__setattr__(md, "conf_int_lower", recompute["conf_int_lower"])
            object.__setattr__(md, "conf_int_upper", recompute["conf_int_upper"])
            flat_idx += 1
    if skipped:
        cols = ", ".join(f"({c})" for c in skipped)
        warnings.warn(
            f"vcov={vcov!r}: recompute skipped for column{'s' if len(skipped)!=1 else ''} "
            f"{cols} — the result lacks data_info['X']/['residuals'] needed for "
            f"HC reweighting. Those columns retain their fit-time SEs.",
            UserWarning,
            stacklevel=3,
        )


def _format_fe_label(token: str) -> str:
    """``"firm"`` → ``"# Firm"``; ``"firm^year"`` → ``"# Firm × Year"``."""
    parts = [p.strip() for p in token.split("^") if p.strip()]
    if not parts:
        return f"# {token}"
    pretty = " × ".join(p[:1].upper() + p[1:] for p in parts)
    return f"# {pretty}"


def _build_fe_size_rows(
    flat_results: Sequence[Any],
) -> "OrderedDict[str, List[str]]":
    """Build per-FE level-count rows for ``fixef_sizes=True``.

    Reads ``model_info['n_fe_levels']`` from each result (a dict of
    FE-name → int). Currently populated by count.py and the pyfixest
    adapter family; other estimators silently contribute empty cells.
    Rows that would be empty across every column are dropped so plain
    OLS bundles stay clean.
    """
    rows: "OrderedDict[str, List[str]]" = OrderedDict()
    union: List[str] = []
    per_col: List[Dict[str, int]] = []
    for r in flat_results:
        mi = getattr(r, "model_info", {}) or {}
        levels = mi.get("n_fe_levels") if isinstance(mi, dict) else None
        if levels is None or not isinstance(levels, dict):
            per_col.append({})
            continue
        coerced: Dict[str, int] = {}
        for k, v in levels.items():
            try:
                coerced[str(k)] = int(v)
            except (TypeError, ValueError):
                continue
        per_col.append(coerced)
        for k in coerced:
            if k not in union:
                union.append(k)
    if not union:
        return rows
    for tok in union:
        cells: List[str] = []
        for d in per_col:
            cells.append(_fmt_int(d.get(tok)) if tok in d else "")
        if any(c for c in cells):
            rows[_format_fe_label(tok)] = cells
    return rows


def _resolve_tests(
    tests: Dict[str, Sequence[Any]],
    n_models: int,
    *,
    fmt: str,
    stars: bool,
    star_levels: Optional[Tuple[float, ...]],
    notation: Union[str, Tuple[str, ...]],
) -> List[Tuple[str, List[str]]]:
    """Normalise ``tests=`` into ordered ``[(label, [cell_per_model, ...])]``.

    Each per-model entry can be:

    - ``(statistic, pvalue)`` tuple → ``"<stat>***"`` (stars from p)
    - bare ``pvalue`` float        → ``"<p>***"``
    - ``None`` / NaN               → empty cell
    - any string                   → passed through as-is

    Stars honour the configured ``notation`` family for cross-table
    consistency.
    """
    if star_levels is None:
        star_levels = (0.10, 0.05, 0.01)
    # Resolve notation to its symbol triple via the same helper as
    # ``RegtableResult._resolve_notation`` to guarantee consistency.
    if isinstance(notation, str):
        key = notation.lower()
        if key == "stars":
            ladder = ("*", "**", "***")
        elif key == "symbols":
            ladder = ("†", "‡", "§")
        else:
            ladder = ("*", "**", "***")
    else:
        try:
            tup = tuple(str(s) for s in notation)
        except TypeError:
            tup = ("*", "**", "***")
        ladder = tup if len(tup) == 3 else ("*", "**", "***")

    def _stars_for(p: float) -> str:
        if p is None or (isinstance(p, float) and np.isnan(p)):
            return ""
        out = ""
        for i, lev in enumerate(sorted(star_levels, reverse=True)):
            if p < lev:
                out = ladder[i] if i < len(ladder) else ladder[-1]
        return out

    out: List[Tuple[str, List[str]]] = []
    for label, seq in tests.items():
        if not isinstance(label, str) or not label:
            raise ValueError(f"tests keys must be non-empty strings, got {label!r}.")
        items = list(seq)
        if len(items) != n_models:
            raise ValueError(
                f"tests[{label!r}] has {len(items)} entries but there are "
                f"{n_models} models."
            )
        cells: List[str] = []
        for entry in items:
            if entry is None:
                cells.append("")
                continue
            if isinstance(entry, str):
                cells.append(entry)
                continue
            # (stat, pvalue) tuple
            if isinstance(entry, (tuple, list)) and len(entry) == 2:
                stat, p = entry
                try:
                    s_txt = _fmt_val(float(stat), fmt)
                except (TypeError, ValueError):
                    s_txt = str(stat) if stat is not None else ""
                marker = _stars_for(float(p)) if (stars and p is not None) else ""
                cells.append(f"{s_txt}{marker}")
                continue
            # Bare scalar — treat as p-value
            try:
                p = float(entry)
            except (TypeError, ValueError):
                cells.append(str(entry))
                continue
            if not np.isfinite(p):
                cells.append("")
                continue
            cells.append(f"{_fmt_val(p, fmt)}{_stars_for(p) if stars else ''}")
        out.append((label, cells))
    return out


def _resolve_multi_se(
    multi_se: Optional[Dict[str, Sequence[Any]]],
    n_models: int,
) -> List[Tuple[str, List[Dict[str, float]]]]:
    """Validate and normalise a ``multi_se`` argument.

    Returns a list of ``(label, [per-model dict-of-var->se, ...])`` tuples
    in user-supplied order. Each per-model entry is a plain ``dict``
    keyed by coefficient name; missing variables yield empty cells.
    """
    if not multi_se:
        return []
    out: List[Tuple[str, List[Dict[str, float]]]] = []
    for label, per_model in multi_se.items():
        if not isinstance(label, str) or not label:
            raise ValueError(f"multi_se keys must be non-empty strings, got {label!r}.")
        seq = list(per_model) if per_model is not None else []
        if len(seq) != n_models:
            raise ValueError(
                f"multi_se[{label!r}] has {len(seq)} entries but there are "
                f"{n_models} models."
            )
        normalized: List[Dict[str, float]] = []
        for entry in seq:
            if entry is None:
                normalized.append({})
                continue
            if isinstance(entry, pd.Series):
                normalized.append({str(k): float(v) for k, v in entry.items()
                                   if v is not None and not pd.isna(v)})
            elif isinstance(entry, dict):
                normalized.append({str(k): float(v) for k, v in entry.items()
                                   if v is not None and not pd.isna(v)})
            else:
                raise TypeError(
                    f"multi_se[{label!r}] entries must be pandas.Series or "
                    f"dict, got {type(entry).__name__}."
                )
        out.append((label, normalized))
    return out


# ═══════════════════════════════════════════════════════════════════════════
# RegtableResult
# ═══════════════════════════════════════════════════════════════════════════

class RegtableResult:
    """Rich result object for regression tables with multi-format export."""

    def __init__(
        self,
        panels: List["_PanelData"],
        *,
        panel_labels: Optional[List[str]],
        model_labels: List[str],
        dep_var_labels: Optional[List[str]],
        coef_labels: Optional[Dict[str, str]],
        keep: Optional[List[str]],
        drop: Optional[List[str]],
        order: Optional[List[str]],
        se_type: str,
        stars: bool,
        star_levels: Tuple[float, ...],
        fmt: str,
        title: Optional[str],
        notes: Optional[List[str]],
        add_rows: Optional[Dict[str, List[str]]],
        stats: Optional[List[str]],
        output: str = "text",
        alpha: float = 0.05,
        multi_se: Optional[List[Tuple[str, List[Dict[str, float]]]]] = None,
        se_label: Optional[str] = None,
        template: Optional[str] = None,
        quarto_label: Optional[str] = None,
        quarto_caption: Optional[str] = None,
        eform_flags: Optional[List[bool]] = None,
        column_spanners: Optional[List[Tuple[str, int]]] = None,
        estimate_template: Optional[str] = None,
        statistic_template: Optional[str] = None,
        notation: Union[str, Tuple[str, ...]] = "stars",
        apply_coef: Optional[Any] = None,
        apply_coef_deriv: Optional[Any] = None,
        escape: bool = True,
        tests_rows: Optional[List[Tuple[str, List[str]]]] = None,
        transpose: bool = False,
    ):
        if not 0.0 < alpha < 1.0:
            raise ValueError(f"alpha must be in (0, 1), got {alpha!r}")
        self.panels = panels
        self.panel_labels = panel_labels
        self.model_labels = model_labels
        self.dep_var_labels = dep_var_labels
        self.coef_labels = coef_labels or {}
        self.keep = keep
        self.drop = set(drop) if drop else set()
        self.order = order
        self.se_type = se_type
        self.show_stars = stars
        self.star_levels = star_levels
        self.fmt = fmt
        self.title = title
        self.notes = notes or []
        self.add_rows = add_rows or {}
        self.requested_stats = stats or ["N", "R2", "adj_R2", "F"]
        self.alpha = float(alpha)
        self.n_models = sum(len(p.models) for p in panels)
        self.multi_se = multi_se or []
        # Override the SE-row footer label (e.g. "Robust standard errors"
        # via QJE preset). When None, the label is derived from se_type.
        self._se_label_override = se_label
        # Journal template name (informational; resolution already happened
        # at the regtable() call site).
        self.template = template
        # Quarto cross-reference metadata. When ``quarto_label`` is set,
        # ``to_quarto()`` (and ``to_markdown(quarto=True)``) emit a Quarto
        # ``: caption {#tbl-<label>}`` line so the table can be referenced
        # via ``@tbl-<label>`` in the manuscript.
        self.quarto_label = quarto_label
        self.quarto_caption = quarto_caption
        # Controls which renderer __str__ uses. Jupyter still gets HTML via
        # _repr_html_ regardless, so output='latex' in a notebook still renders
        # pretty HTML — users who want the LaTeX source call to_latex() or
        # print(result).
        self._output = output

        # Per-model eform flag (length n_models). When True for a model,
        # the rendered point estimate becomes exp(b), the SE becomes
        # exp(b)·SE(b) (delta method), and CI bounds become (exp(lo), exp(hi))
        # of the original endpoints. t and p-values are unchanged because
        # H0: b=0 is equivalent to H0: exp(b)=1.
        if eform_flags is None:
            eform_flags = [False] * self.n_models
        if len(eform_flags) != self.n_models:
            raise ValueError(
                f"eform_flags has {len(eform_flags)} entries but "
                f"there are {self.n_models} models."
            )
        self.eform_flags = [bool(f) for f in eform_flags]

        # Column spanners: a list of ``(label, span)`` tuples where the
        # spans must sum to n_models. Renders as a multi-row header above
        # the model labels (LaTeX ``\multicolumn``, HTML ``colspan``,
        # text-mode centered ASCII).
        if column_spanners is not None:
            total_span = sum(int(s) for _, s in column_spanners)
            if total_span != self.n_models:
                raise ValueError(
                    f"column_spanners total span = {total_span} but there are "
                    f"{self.n_models} models. The spans must partition the "
                    f"columns exactly (consecutive grouping)."
                )
        self.column_spanners = (
            [(str(lbl), int(s)) for lbl, s in column_spanners]
            if column_spanners
            else None
        )

        # ── estimate / statistic templates ─────────────────────────────────
        # When set, ``_coef_cell`` and ``_se_cell`` build a context dict and
        # use ``str.format_map`` to produce the cell. Backward compat: when
        # both are None, the legacy "{val}{stars}" / "({se})" path runs.
        self.estimate_template = estimate_template
        self.statistic_template = statistic_template

        # ── notation (significance markers) ────────────────────────────────
        # ``"stars"``  → ("*", "**", "***") — backward compatible
        # ``"symbols"`` → ("†", "‡", "§")     — AER / JPE alternative
        # tuple        → custom 3-tuple of strings
        self._notation_symbols = self._resolve_notation(notation)

        # ── apply_coef ────────────────────────────────────────────────────
        # Generalised eform. Conflicts with eform_flags are caught up at the
        # public regtable() entry. Optional ``apply_coef_deriv`` enables
        # delta-method SE rescaling: SE → |f'(b)| · SE(b).
        if apply_coef is not None and not callable(apply_coef):
            raise TypeError(
                f"apply_coef must be callable, got {type(apply_coef).__name__}"
            )
        if apply_coef_deriv is not None and not callable(apply_coef_deriv):
            raise TypeError(
                f"apply_coef_deriv must be callable, got "
                f"{type(apply_coef_deriv).__name__}"
            )
        self.apply_coef = apply_coef
        self.apply_coef_deriv = apply_coef_deriv

        # ── escape ─────────────────────────────────────────────────────────
        # When False, the LaTeX/HTML renderers skip ``_latex_escape`` /
        # ``_html_escape`` on user-supplied label paths so users can pass
        # raw markup (e.g. ``"$\\beta_1$"``). Cell content (numeric
        # estimates, computed stat values) is still safe — it never
        # contains user-controlled metacharacters.
        self.escape = bool(escape)

        # ── tests rows (footer) ────────────────────────────────────────────
        # Pre-formatted hypothesis-test rows produced at the regtable()
        # entry. Stored as ordered ``[(label, [cell_per_model, ...])]`` so
        # rendering threads through the same code path as ``add_rows``.
        self.tests_rows = tests_rows or []

        # ── transpose ──────────────────────────────────────────────────────
        # Single-panel-only pivot: rows become models, columns become
        # variables. The validation lives at the regtable() entry; here
        # we just store the flag for renderers.
        self.transpose = bool(transpose)

        # Resolve stat keys once
        self._stat_keys = self._resolve_stat_keys()

    # --- helpers -----------------------------------------------------------

    def _resolve_stat_keys(self) -> List[str]:
        keys: List[str] = []
        for s in self.requested_stats:
            canonical = _STAT_ALIASES.get(s, s)
            if canonical not in keys:
                keys.append(canonical)
        return keys

    def _resolve_vars(self, models: List[_ModelData]) -> List[str]:
        seen: OrderedDict[str, None] = OrderedDict()
        for m in models:
            for v in m.params.index:
                seen[v] = None
        all_vars = list(seen)

        if self.keep is not None:
            keep_set = set(self.keep)
            all_vars = [v for v in all_vars if v in keep_set]
        if self.drop:
            all_vars = [v for v in all_vars if v not in self.drop]
        if self.order:
            ordered: List[str] = []
            remaining = list(all_vars)
            for v in self.order:
                if v in remaining:
                    ordered.append(v)
                    remaining.remove(v)
            ordered.extend(remaining)
            all_vars = ordered
        return all_vars

    def _model_eform(self, flat_idx: int) -> bool:
        """Return whether model at flat position ``flat_idx`` uses eform."""
        if not self.eform_flags:
            return False
        if 0 <= flat_idx < len(self.eform_flags):
            return self.eform_flags[flat_idx]
        return False

    @staticmethod
    def _resolve_notation(
        notation: Union[str, Tuple[str, ...]],
    ) -> Tuple[str, str, str]:
        """Map a notation spec to a 3-tuple ``(low, mid, high)`` symbols.

        Mirrors the canonical Top-5 ladder where the *high* symbol marks
        the strictest significance threshold. Three valid input shapes:

        - ``"stars"`` (default) → ``("*", "**", "***")``
        - ``"symbols"``        → ``("†", "‡", "§")`` (AER / JPE family)
        - 3-tuple of strings   → custom (must be ordered low → high)
        """
        if isinstance(notation, str):
            key = notation.lower()
            if key == "stars":
                return ("*", "**", "***")
            if key == "symbols":
                return ("†", "‡", "§")
            raise ValueError(
                f"notation={notation!r} unrecognised. Use 'stars', "
                f"'symbols', or pass a 3-tuple of custom strings."
            )
        try:
            tup = tuple(str(s) for s in notation)
        except TypeError:
            raise TypeError(
                f"notation must be a string or 3-tuple, got "
                f"{type(notation).__name__}"
            )
        if len(tup) != 3:
            raise ValueError(
                f"notation tuple must have exactly 3 entries (low, mid, "
                f"high), got {len(tup)}."
            )
        return tup  # type: ignore[return-value]

    def _format_marker(self, pvalue: float) -> str:
        """Significance marker honouring the configured notation family."""
        if pvalue is None or (isinstance(pvalue, float) and np.isnan(pvalue)):
            return ""
        # The shipping notation family is a 3-tuple ``(low, mid, high)``
        # where *high* corresponds to the strictest threshold. Building
        # via cumulative compare keeps semantics identical to the legacy
        # ``_format_stars``: e.g. p<0.01 yields "***" when stars / "§"
        # when symbols.
        sorted_levels = sorted(self.star_levels, reverse=True)
        out = ""
        low, mid, high = self._notation_symbols
        ladder = (low, mid, high)
        for i, lev in enumerate(sorted_levels):
            if pvalue < lev:
                out = ladder[i] if i < len(ladder) else ladder[-1]
        return out

    def _apply_coef_transform(
        self, b: float, se: Optional[float] = None,
        flat_idx: int = 0,
    ) -> Tuple[float, Optional[float]]:
        """Apply (eform OR apply_coef) to a (coef, SE) pair.

        Returns transformed ``(coef, se)``. eform and apply_coef are
        mutually exclusive; eform_flags is per-model whereas apply_coef
        is global. Delta method:

        - eform:  b' = exp(b), se' = exp(b) * se
        - apply_coef + apply_coef_deriv: b' = f(b), se' = |f'(b)| * se
        - apply_coef alone:              b' = f(b), se' = se (unchanged,
          user takes responsibility — flagged in footer note)
        """
        if self.apply_coef is not None:
            try:
                b_new = float(self.apply_coef(b))
            except Exception:
                b_new = b  # transform raised; leave unchanged
            if se is not None and self.apply_coef_deriv is not None:
                try:
                    deriv = float(self.apply_coef_deriv(b))
                    se_new: Optional[float] = abs(deriv) * float(se)
                    return b_new, se_new
                except Exception:
                    pass
            return b_new, se
        if self._model_eform(flat_idx) and np.isfinite(b):
            b_new = float(np.exp(b))
            if se is not None and not pd.isna(se):
                return b_new, float(np.exp(b)) * float(se)
            return b_new, se
        return b, se

    def _build_cell_context(
        self, model: _ModelData, var: str, flat_idx: int,
    ) -> Dict[str, str]:
        """Return the variable-substitution dict for template rendering.

        Keys exposed (matching R ``modelsummary`` conventions):

        - ``estimate``     — coefficient, after eform / apply_coef
        - ``std_error``    — primary SE, delta-method-rescaled
        - ``t_value``      — t-statistic (always on the original scale)
        - ``p_value``      — p-value (always on the original scale)
        - ``conf_low`` / ``conf_high`` — CI bounds, exp-rescaled under eform
        - ``stars``        — significance marker (notation-aware)

        We deliberately use underscore names (``std_error``, ``t_value``)
        because Python's ``str.format`` syntax disallows dots in field
        names; aliases ``std.error`` / ``t.statistic`` / ``conf.low`` /
        ``conf.high`` are added so prose lifted from R templates "just
        works".
        """
        if var not in model.params.index:
            return {}
        b = float(model.params[var])
        se = model.std_errors.get(var, np.nan)
        t_val = float(model.tvalues.get(var, np.nan))
        p_val = float(model.pvalues.get(var, np.nan))
        ci_lo, ci_hi = _ci_bounds(model, var, self.alpha)

        b_new, se_new = self._apply_coef_transform(b, se, flat_idx)
        if self._model_eform(flat_idx):
            if not (ci_lo is None or pd.isna(ci_lo)):
                ci_lo = float(np.exp(ci_lo))
            if not (ci_hi is None or pd.isna(ci_hi)):
                ci_hi = float(np.exp(ci_hi))
        elif self.apply_coef is not None:
            try:
                ci_lo = float(self.apply_coef(ci_lo)) if not pd.isna(ci_lo) else ci_lo
                ci_hi = float(self.apply_coef(ci_hi)) if not pd.isna(ci_hi) else ci_hi
            except Exception:
                pass

        marker = self._format_marker(p_val) if self.show_stars else ""
        ctx = {
            "estimate": _fmt_val(b_new, self.fmt),
            "std_error": _fmt_val(se_new if se_new is not None else np.nan, self.fmt),
            "t_value": _fmt_val(t_val, self.fmt),
            "p_value": _fmt_val(p_val, self.fmt),
            "conf_low": _fmt_val(ci_lo, self.fmt),
            "conf_high": _fmt_val(ci_hi, self.fmt),
            "stars": marker,
        }
        return ctx

    def _coef_cell(self, model: _ModelData, var: str, flat_idx: int = 0) -> str:
        if var not in model.params.index:
            return ""
        if self.estimate_template is not None or self.statistic_template is not None:
            ctx = self._build_cell_context(model, var, flat_idx)
            template = self.estimate_template or "{estimate}{stars}"
            return template.format_map(ctx)
        # Legacy path
        b = float(model.params[var])
        b_new, _ = self._apply_coef_transform(b, None, flat_idx)
        txt = _fmt_val(b_new, self.fmt)
        if self.show_stars and var in model.pvalues.index:
            txt += self._format_marker(model.pvalues[var])
        return txt

    def _se_cell(self, model: _ModelData, var: str, flat_idx: int = 0) -> str:
        if var not in model.params.index:
            return ""
        if self.statistic_template is not None:
            ctx = self._build_cell_context(model, var, flat_idx)
            return self.statistic_template.format_map(ctx)
        eform = self._model_eform(flat_idx)
        apply_active = self.apply_coef is not None
        if self.se_type == "ci":
            lo_v, hi_v = _ci_bounds(model, var, self.alpha)
            if eform:
                if not (lo_v is None or pd.isna(lo_v)):
                    lo_v = float(np.exp(lo_v))
                if not (hi_v is None or pd.isna(hi_v)):
                    hi_v = float(np.exp(hi_v))
            elif apply_active:
                try:
                    if not (lo_v is None or pd.isna(lo_v)):
                        lo_v = float(self.apply_coef(lo_v))
                    if not (hi_v is None or pd.isna(hi_v)):
                        hi_v = float(self.apply_coef(hi_v))
                except Exception:
                    pass
            lo = _fmt_val(lo_v, self.fmt)
            hi = _fmt_val(hi_v, self.fmt)
            return f"[{lo}, {hi}]"
        if self.se_type == "t":
            return f"({_fmt_val(model.tvalues.get(var, np.nan), self.fmt)})"
        if self.se_type == "p":
            return f"({_fmt_val(model.pvalues.get(var, np.nan), self.fmt)})"
        # default: standard error (delta method under eform / apply_coef_deriv)
        b = float(model.params[var])
        se_v = model.std_errors.get(var, np.nan)
        _, se_new = self._apply_coef_transform(b, se_v, flat_idx)
        return f"({_fmt_val(se_new if se_new is not None else se_v, self.fmt)})"

    def _has_any_eform(self) -> bool:
        return any(self.eform_flags)

    def _has_apply_coef_no_deriv(self) -> bool:
        """True when apply_coef is set without a matching derivative.

        Surface a footer note in this regime so reviewers know the SE
        column is **not** rescaled by the transform — only the point
        estimate is.
        """
        return self.apply_coef is not None and self.apply_coef_deriv is None

    def _esc_latex(self, s: str) -> str:
        """Escape ``s`` for LaTeX iff ``self.escape`` is True."""
        return _latex_escape(s) if self.escape else (s or "")

    def _esc_html(self, s: str) -> str:
        """Escape ``s`` for HTML iff ``self.escape`` is True."""
        return _html_escape(s) if self.escape else (s or "")

    def _star_note_text(self) -> str:
        """Render the notation-family footer line.

        Reuses :func:`star_note_for` for the legacy ``"stars"`` family so
        journal-template footnotes stay byte-aligned. For non-star
        notations we build our own line listing each ``symbol`` against
        its threshold.
        """
        low, mid, high = self._notation_symbols
        if (low, mid, high) == ("*", "**", "***"):
            return star_note_for(self.star_levels)
        sorted_levels = sorted(self.star_levels, reverse=True)
        ladder = [low, mid, high]
        parts: List[str] = []
        # Iterate in strict-first order to match _format_marker ordering
        for i, lev in enumerate(sorted_levels):
            sym = ladder[i] if i < len(ladder) else ladder[-1]
            parts.append(f"{sym} p<{lev:g}")
        return ", ".join(reversed(parts))

    def _se_label(self) -> str:
        if self._se_label_override is not None and self.se_type == "se":
            return self._se_label_override
        if self.se_type == "ci":
            level = (1.0 - self.alpha) * 100.0
            return f"{level:g}% CI"
        return {"t": "t-statistics", "p": "p-values"}.get(
            self.se_type, "Standard errors"
        )

    def _multi_se_cell(
        self,
        per_model: Dict[str, float],
        var: str,
        bracket_idx: int,
        model: Optional[_ModelData] = None,
        flat_idx: int = 0,
    ) -> str:
        """Render the bracket-wrapped extra-SE cell for one model column.

        When eform is active for the column, the extra SE (bootstrap /
        cluster / jackknife / etc.) is rescaled by ``exp(b)`` via the
        same delta-method approximation used for the primary SE — keeps
        the cells reading on a single scale.
        """
        if var not in per_model:
            return ""
        try:
            val = float(per_model[var])
        except (TypeError, ValueError):
            return ""
        if not np.isfinite(val):
            return ""
        if (
            self._model_eform(flat_idx)
            and model is not None
            and var in model.params.index
        ):
            b = float(model.params[var])
            if np.isfinite(b):
                val = float(np.exp(b)) * val
        lo, hi = _MULTI_SE_BRACKETS[bracket_idx % len(_MULTI_SE_BRACKETS)]
        return f"{lo}{_fmt_val(val, self.fmt)}{hi}"

    def _stat_cell(self, model: _ModelData, key: str) -> str:
        val = model.stats.get(key)
        if val is None:
            return ""
        if key == "N":
            return _fmt_int(val)
        return _fmt_val(float(val), "%.3f")

    def _star_note(self) -> str:
        # Delegate to ``_star_note_text``, which understands both the legacy
        # ``"stars"`` family (deferring to ``star_note_for`` so journal
        # presets stay byte-aligned) and the alternative notation tuples.
        return self._star_note_text()

    def _all_models_flat(self) -> List[_ModelData]:
        out: List[_ModelData] = []
        for p in self.panels:
            out.extend(p.models)
        return out

    # ═══════════════════════════════════════════════════════════════════════
    # TEXT
    # ═══════════════════════════════════════════════════════════════════════

    def _text_panel(
        self,
        models: List[_ModelData],
        var_list: List[str],
        col_w: int,
        label_w: int,
        panel_idx: int = 0,
    ) -> List[str]:
        lines: List[str] = []
        base_idx = sum(len(p.models) for p in self.panels[:panel_idx])
        for var in var_list:
            label = self.coef_labels.get(var, var)
            row = f"{label:<{label_w}}"
            for off, m in enumerate(models):
                row += f"{self._coef_cell(m, var, base_idx + off):>{col_w}}"
            lines.append(row)
            # SE row
            row2 = " " * label_w
            for off, m in enumerate(models):
                row2 += f"{self._se_cell(m, var, base_idx + off):>{col_w}}"
            lines.append(row2)
            # Extra SE rows from multi_se. Each label maps to one entry per
            # model across the WHOLE table, so we slice into the panel using
            # cumulative model offsets.
            for ext_idx, (_, per_model_list) in enumerate(self.multi_se):
                row3 = " " * label_w
                for off, m in enumerate(models):
                    per_model = per_model_list[base_idx + off]
                    row3 += f"{self._multi_se_cell(per_model, var, ext_idx, m, base_idx + off):>{col_w}}"
                lines.append(row3)
            lines.append("")  # blank between vars
        return lines

    def to_text(self) -> str:
        if self.transpose:
            return self._to_text_transposed()
        col_w = 14
        all_models = self._all_models_flat()
        all_vars_set: set = set()
        for p in self.panels:
            all_vars_set.update(v for m in p.models for v in m.params.index)
        label_names = [self.coef_labels.get(v, v) for v in all_vars_set]
        stat_names = [_STAT_DISPLAY.get(k, k) for k in self._stat_keys]
        add_row_names = list(self.add_rows.keys())
        max_label = max(
            (len(n) for n in label_names + stat_names + add_row_names),
            default=10,
        )
        label_w = max(max_label + 2, 18)
        total_w = label_w + col_w * len(all_models) + 2

        thick = "\u2501" * total_w
        thin = "\u2500" * total_w
        lines: List[str] = []

        if self.title:
            lines.append(f"  {self.title}")
            lines.append("")

        lines.append(thick)

        # Column spanners: a header row above the model-label row, with
        # each label centered over its block of columns.
        if self.column_spanners:
            span_row = " " * label_w
            for lbl, span in self.column_spanners:
                block_w = col_w * span
                span_row += f"{lbl:^{block_w}}"
            lines.append(span_row)
            # Thin underline beneath each spanner block (cmidrule analog)
            rule_row = " " * label_w
            for _, span in self.column_spanners:
                block_w = col_w * span
                rule_row += " " + "─" * (block_w - 2) + " "
            lines.append(rule_row)

        # Header: model labels
        hdr = " " * label_w
        for lbl in self.model_labels:
            hdr += f"{lbl:>{col_w}}"
        lines.append(hdr)

        # Dep-var row
        if self.dep_var_labels:
            dvr = " " * label_w
            for dv in self.dep_var_labels:
                dvr += f"{dv:>{col_w}}"
            lines.append(dvr)

        lines.append(thick)

        # Panels
        multi = len(self.panels) > 1
        for pi, panel in enumerate(self.panels):
            if multi and self.panel_labels and pi < len(self.panel_labels):
                lines.append(f"  {self.panel_labels[pi]}")
                lines.append(thin)

            var_list = self._resolve_vars(panel.models)
            lines.extend(self._text_panel(panel.models, var_list, col_w, label_w, panel_idx=pi))

            if multi and pi < len(self.panels) - 1:
                lines.append(thin)

        lines.append(thick)

        # Add rows (Controls, FE, etc.)
        for row_label, row_vals in self.add_rows.items():
            row = f"{row_label:<{label_w}}"
            for i, m in enumerate(all_models):
                val = row_vals[i] if i < len(row_vals) else ""
                row += f"{val:>{col_w}}"
            lines.append(row)

        if self.add_rows:
            lines.append(thick)

        # Stats
        for key in self._stat_keys:
            disp = _STAT_DISPLAY.get(key, key)
            row = f"{disp:<{label_w}}"
            for m in all_models:
                row += f"{self._stat_cell(m, key):>{col_w}}"
            lines.append(row)

        # Hypothesis-test rows (Wald F, Hansen-J, etc.). They sit below
        # the stats block and above the bottom rule so they read as part
        # of the diagnostic strip — closer to AER convention than mixing
        # them with the structural ``add_rows``.
        if self.tests_rows:
            for row_label, row_vals in self.tests_rows:
                row = f"{row_label:<{label_w}}"
                for i, m in enumerate(all_models):
                    val = row_vals[i] if i < len(row_vals) else ""
                    row += f"{val:>{col_w}}"
                lines.append(row)

        lines.append(thick)

        # Notes
        lines.append(f"{self._se_label()} in parentheses")
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            lines.append(f"{label} in {lo}…{hi}")
        if self._has_any_eform():
            lines.append(self._eform_note())
        if self.show_stars:
            lines.append(self._star_note())
        for note in self.notes:
            lines.append(note)

        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # Transposed renderers (single-panel pivot — rows ↔ models)
    # ═══════════════════════════════════════════════════════════════════════

    def _transposed_var_list(self) -> List[str]:
        """All variable names that the (single) panel renders, in order."""
        panel = self.panels[0]
        return self._resolve_vars(panel.models)

    def _to_text_transposed(self) -> str:
        col_w = 14
        panel = self.panels[0]
        models = panel.models
        var_list = self._transposed_var_list()
        # Each variable contributes one cell per model; we render the
        # estimate (with stars) and SE on a single row to keep the
        # transposed form compact, then a SE row for every model row.
        var_labels = [self.coef_labels.get(v, v) for v in var_list]
        # Width of leftmost column (model labels)
        label_w = max(
            (len(l) for l in self.model_labels),
            default=4,
        )
        label_w = max(label_w + 2, 8)

        total_w = label_w + col_w * len(var_list) + 4
        thick = "━" * total_w
        thin = "─" * total_w
        lines: List[str] = []
        if self.title:
            lines.append(f"  {self.title}")
            lines.append("")
        lines.append(thick)
        # Header: blank + variable labels
        hdr = " " * label_w
        for vl in var_labels:
            hdr += f"{vl:>{col_w}}"
        lines.append(hdr)
        lines.append(thick)
        # Body: one (estimate, SE) pair of rows per model
        for mi, m in enumerate(models):
            row_e = f"{self.model_labels[mi]:<{label_w}}"
            for var in var_list:
                row_e += f"{self._coef_cell(m, var, mi):>{col_w}}"
            lines.append(row_e)
            row_s = " " * label_w
            for var in var_list:
                row_s += f"{self._se_cell(m, var, mi):>{col_w}}"
            lines.append(row_s)
        lines.append(thick)
        # Stats: one row per stat — N, R², Adj. R², F (each cell is the
        # value for that model under the corresponding column? No — in
        # transposed form, stats need to live somewhere new. Convention
        # we adopt: append a *trailing* per-model column for each stat
        # would over-stuff the header; instead emit a final block where
        # each stat name is a *new column* on the right, holding the
        # per-model value.
        if self._stat_keys:
            # Re-render header with stats appended for visibility
            stat_disp = [_STAT_DISPLAY.get(k, k) for k in self._stat_keys]
            stat_hdr = " " * (label_w + col_w * len(var_list))
            for sd in stat_disp:
                stat_hdr += f"{sd:>{col_w}}"
            # Splice stat columns into existing rows
            new_lines: List[str] = []
            row_idx = 0
            for ln in lines:
                if ln == thick or ln == thin or ln == "" or ln.startswith("  "):
                    new_lines.append(ln)
                    continue
                # First data row after the second thick is the header;
                # the body rows alternate (est, se). Append stat values
                # only on est rows.
                pass
            # Simpler: rebuild header + body with stat columns
            # ────────────────────────────────────────────────────────
            # Reset lines and rebuild fresh
            lines = []
            if self.title:
                lines.append(f"  {self.title}")
                lines.append("")
            total_w_full = label_w + col_w * (len(var_list) + len(self._stat_keys))
            thick = "━" * total_w_full
            lines.append(thick)
            # Combined header
            hdr = " " * label_w
            for vl in var_labels:
                hdr += f"{vl:>{col_w}}"
            for sd in stat_disp:
                hdr += f"{sd:>{col_w}}"
            lines.append(hdr)
            lines.append(thick)
            # Body
            for mi, m in enumerate(models):
                row_e = f"{self.model_labels[mi]:<{label_w}}"
                for var in var_list:
                    row_e += f"{self._coef_cell(m, var, mi):>{col_w}}"
                for k in self._stat_keys:
                    row_e += f"{self._stat_cell(m, k):>{col_w}}"
                lines.append(row_e)
                row_s = " " * label_w
                for var in var_list:
                    row_s += f"{self._se_cell(m, var, mi):>{col_w}}"
                for _ in self._stat_keys:
                    row_s += f"{'':>{col_w}}"
                lines.append(row_s)
            lines.append(thick)
        # Notes
        lines.append(f"{self._se_label()} in parentheses")
        if self._has_any_eform():
            lines.append(self._eform_note())
        if self.show_stars:
            lines.append(self._star_note())
        for note in self.notes:
            lines.append(note)
        return "\n".join(lines)

    def _to_html_transposed(self) -> str:
        panel = self.panels[0]
        models = panel.models
        var_list = self._transposed_var_list()
        var_labels = [self.coef_labels.get(v, v) for v in var_list]
        stat_disp = [_STAT_DISPLAY.get(k, k) for k in self._stat_keys]
        ncols = 1 + len(var_list) + len(self._stat_keys)
        lines: List[str] = []
        lines.append(
            '<table class="regtable regtable-transposed" '
            'style="border-collapse:collapse; '
            'font-family:\'Times New Roman\', serif; font-size:13px;">'
        )
        if self.title:
            lines.append(
                f'<caption style="font-weight:bold; font-size:14px; '
                f'caption-side:top;">{self._esc_html(self.title)}</caption>'
            )
        lines.append("<thead>")
        lines.append("<tr>")
        lines.append(
            '<th style="border-top:3px solid black; border-bottom:1px solid black; '
            'padding:4px 8px;"></th>'
        )
        for vl in var_labels:
            lines.append(
                f'<th style="border-top:3px solid black; border-bottom:1px solid black; '
                f'padding:4px 12px;">{self._esc_html(vl)}</th>'
            )
        for sd in stat_disp:
            lines.append(
                f'<th style="border-top:3px solid black; border-bottom:1px solid black; '
                f'padding:4px 12px; font-style:italic;">{self._esc_html(sd)}</th>'
            )
        lines.append("</tr>")
        lines.append("</thead>")
        lines.append("<tbody>")
        for mi, m in enumerate(models):
            # Estimate row
            lines.append("<tr>")
            lines.append(
                f'<td style="text-align:left; padding:1px 8px;">'
                f'{self._esc_html(self.model_labels[mi])}</td>'
            )
            for var in var_list:
                lines.append(
                    f'<td style="text-align:center; padding:1px 12px;">'
                    f'{_html_escape(self._coef_cell(m, var, mi))}</td>'
                )
            for k in self._stat_keys:
                lines.append(
                    f'<td style="text-align:center; padding:1px 12px;">'
                    f'{self._stat_cell(m, k)}</td>'
                )
            lines.append("</tr>")
            # SE row
            lines.append("<tr>")
            lines.append("<td></td>")
            for var in var_list:
                lines.append(
                    f'<td style="text-align:center; padding:0 12px; '
                    f'color:#555; font-size:12px;">'
                    f'{_html_escape(self._se_cell(m, var, mi))}</td>'
                )
            for _ in self._stat_keys:
                lines.append("<td></td>")
            lines.append("</tr>")
        # Bottom rule
        lines.append(
            f'<tr><td colspan="{ncols}" '
            f'style="border-top:3px solid black; padding:0;"></td></tr>'
        )
        lines.append("</tbody>")
        # Notes
        lines.append("<tfoot>")
        note = f"{self._se_label()} in parentheses"
        lines.append(
            f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
            f'padding:4px 8px 0 8px;">{_html_escape(note)}</td></tr>'
        )
        if self.show_stars:
            lines.append(
                f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
                f'padding:0 8px;">{_html_escape(self._star_note())}</td></tr>'
            )
        lines.append("</tfoot>")
        lines.append("</table>")
        return "\n".join(lines)

    def _eform_note(self) -> str:
        """Footer note explaining the eform transformation."""
        if all(self.eform_flags):
            return ("Coefficients reported as exp(b); standard errors via "
                    "delta method (exp(b)·SE). Stars from p-values of the "
                    "untransformed estimates.")
        cols = [i + 1 for i, f in enumerate(self.eform_flags) if f]
        col_str = ", ".join(f"({c})" for c in cols)
        return (f"Columns {col_str} report exp(b) (delta-method SE); "
                f"other columns report b. Stars from p-values of the "
                f"untransformed estimates.")

    # ═══════════════════════════════════════════════════════════════════════
    # HTML (also _repr_html_)
    # ═══════════════════════════════════════════════════════════════════════

    def to_html(self) -> str:
        if self.transpose:
            return self._to_html_transposed()
        all_models = self._all_models_flat()
        ncols = len(all_models) + 1
        lines: List[str] = []
        lines.append(
            '<table class="regtable" style="border-collapse:collapse; '
            'font-family:\'Times New Roman\', serif; font-size:13px; min-width:500px;">'
        )

        if self.title:
            lines.append(
                f'<caption style="font-weight:bold; font-size:14px; '
                f'margin-bottom:8px; caption-side:top;">'
                f'{self._esc_html(self.title)}</caption>'
            )

        # Header
        lines.append("<thead>")

        # Column spanners (above the model-label row) when set
        if self.column_spanners:
            lines.append("<tr>")
            lines.append(
                '<th style="text-align:left; border-top:3px solid black; '
                'padding:4px 8px;"></th>'
            )
            for lbl, span in self.column_spanners:
                lines.append(
                    f'<th colspan="{span}" style="text-align:center; '
                    f'border-top:3px solid black; border-bottom:1px solid #999; '
                    f'padding:4px 12px;">{self._esc_html(lbl)}</th>'
                )
            lines.append("</tr>")

        lines.append("<tr>")
        # When spanners present, the top rule already lives on the spanner row
        top_rule = "" if self.column_spanners else "border-top:3px solid black; "
        lines.append(
            f'<th style="text-align:left; {top_rule}'
            'border-bottom:1px solid black; padding:4px 8px;"></th>'
        )
        for lbl in self.model_labels:
            lines.append(
                f'<th style="text-align:center; {top_rule}'
                f'border-bottom:1px solid black; padding:4px 12px;">'
                f'{self._esc_html(lbl)}</th>'
            )
        lines.append("</tr>")

        # Dep-var row
        if self.dep_var_labels:
            lines.append("<tr>")
            lines.append('<th style="text-align:left; padding:2px 8px;"></th>')
            for dv in self.dep_var_labels:
                lines.append(
                    f'<th style="text-align:center; padding:2px 12px; '
                    f'font-style:italic; font-weight:normal;">'
                    f'{self._esc_html(dv)}</th>'
                )
            lines.append("</tr>")

        lines.append("</thead>")
        lines.append("<tbody>")

        # Panels
        multi = len(self.panels) > 1
        model_idx = 0
        for pi, panel in enumerate(self.panels):
            if multi and self.panel_labels and pi < len(self.panel_labels):
                lines.append(
                    f'<tr><td colspan="{ncols}" style="text-align:left; '
                    f'font-weight:bold; padding:6px 8px 2px 8px; '
                    f'border-top:1px solid #999;">'
                    f'{self._esc_html(self.panel_labels[pi])}</td></tr>'
                )

            var_list = self._resolve_vars(panel.models)
            for var in var_list:
                label = self._esc_html(self.coef_labels.get(var, var))
                lines.append("<tr>")
                lines.append(
                    f'<td style="text-align:left; padding:1px 8px;">{label}</td>'
                )
                # Empty cells for models NOT in this panel
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            lines.append(
                                f'<td style="text-align:center; padding:1px 12px;">'
                                f'{_html_escape(self._coef_cell(m, var, flat_idx))}</td>'
                            )
                        else:
                            lines.append(
                                '<td style="text-align:center; padding:1px 12px;"></td>'
                            )
                        flat_idx += 1
                lines.append("</tr>")
                # SE row
                lines.append("<tr>")
                lines.append("<td></td>")
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            lines.append(
                                f'<td style="text-align:center; padding:0 12px; '
                                f'color:#555; font-size:12px;">'
                                f'{_html_escape(self._se_cell(m, var, flat_idx))}</td>'
                            )
                        else:
                            lines.append(
                                '<td style="text-align:center; padding:0 12px;"></td>'
                            )
                        flat_idx += 1
                lines.append("</tr>")
                # Extra SE rows from multi_se
                base_idx = sum(len(p.models) for p in self.panels[:pi])
                for ext_idx, (_, per_model_list) in enumerate(self.multi_se):
                    lines.append("<tr>")
                    lines.append("<td></td>")
                    for gi, p2 in enumerate(self.panels):
                        for off, m in enumerate(p2.models):
                            if gi == pi:
                                per_model = per_model_list[base_idx + off]
                                cell = self._multi_se_cell(
                                    per_model, var, ext_idx, m, base_idx + off
                                )
                                lines.append(
                                    f'<td style="text-align:center; padding:0 12px; '
                                    f'color:#777; font-size:12px;">'
                                    f'{_html_escape(cell)}</td>'
                                )
                            else:
                                lines.append(
                                    '<td style="text-align:center; padding:0 12px;"></td>'
                                )
                    lines.append("</tr>")

        # Separator
        lines.append(
            f'<tr><td colspan="{ncols}" '
            f'style="border-top:1px solid black; padding:0;"></td></tr>'
        )

        # Add rows
        for row_label, row_vals in self.add_rows.items():
            lines.append("<tr>")
            lines.append(
                f'<td style="text-align:left; padding:1px 8px;">'
                f'{self._esc_html(row_label)}</td>'
            )
            for i in range(len(all_models)):
                val = row_vals[i] if i < len(row_vals) else ""
                lines.append(
                    f'<td style="text-align:center; padding:1px 12px;">'
                    f'{self._esc_html(val)}</td>'
                )
            lines.append("</tr>")

        if self.add_rows:
            lines.append(
                f'<tr><td colspan="{ncols}" '
                f'style="border-top:1px solid #aaa; padding:0;"></td></tr>'
            )

        # Stats
        for key in self._stat_keys:
            disp = _html_escape(_STAT_DISPLAY.get(key, key))
            lines.append("<tr>")
            lines.append(
                f'<td style="text-align:left; padding:1px 8px;">{disp}</td>'
            )
            for m in all_models:
                lines.append(
                    f'<td style="text-align:center; padding:1px 12px;">'
                    f'{self._stat_cell(m, key)}</td>'
                )
            lines.append("</tr>")

        # Hypothesis-test rows
        for row_label, row_vals in self.tests_rows:
            lines.append("<tr>")
            lines.append(
                f'<td style="text-align:left; padding:1px 8px;">'
                f'{self._esc_html(row_label)}</td>'
            )
            for i in range(len(all_models)):
                val = row_vals[i] if i < len(row_vals) else ""
                lines.append(
                    f'<td style="text-align:center; padding:1px 12px;">'
                    f'{self._esc_html(val)}</td>'
                )
            lines.append("</tr>")

        # Bottom border
        lines.append(
            f'<tr><td colspan="{ncols}" '
            f'style="border-top:3px solid black; padding:0;"></td></tr>'
        )

        lines.append("</tbody>")

        # Notes
        lines.append("<tfoot>")
        note_text = f"{self._se_label()} in parentheses"
        lines.append(
            f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
            f'padding:4px 8px 0 8px;">{_html_escape(note_text)}</td></tr>'
        )
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            multi_note = f"{label} in {lo}…{hi}"
            lines.append(
                f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
                f'padding:0 8px;">{_html_escape(multi_note)}</td></tr>'
            )
        if self._has_any_eform():
            lines.append(
                f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
                f'padding:0 8px;">{_html_escape(self._eform_note())}</td></tr>'
            )
        if self.show_stars:
            lines.append(
                f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
                f'padding:0 8px;">{_html_escape(self._star_note())}</td></tr>'
            )
        for note in self.notes:
            lines.append(
                f'<tr><td colspan="{ncols}" style="text-align:left; font-size:11px; '
                f'padding:0 8px;">{_html_escape(note)}</td></tr>'
            )
        lines.append("</tfoot>")
        lines.append("</table>")
        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # LaTeX
    # ═══════════════════════════════════════════════════════════════════════

    def to_latex(self) -> str:
        all_models = self._all_models_flat()
        n_cols = len(all_models) + 1
        col_spec = "l" + "c" * len(all_models)
        lines: List[str] = []
        lines.append("\\begin{table}[htbp]")
        lines.append("\\centering")
        if self.title:
            lines.append(f"\\caption{{{self._esc_latex(self.title)}}}")
        lines.append(f"\\begin{{tabular}}{{{col_spec}}}")
        lines.append("\\hline\\hline")

        # Column spanners (above the model-label row) when set
        if self.column_spanners:
            cells_sp: List[str] = [""]
            cmidrules: List[str] = []
            cur_col = 2  # LaTeX column 1 is the row-label, models start at 2
            for lbl, span in self.column_spanners:
                cells_sp.append(
                    f"\\multicolumn{{{span}}}{{c}}{{{self._esc_latex(lbl)}}}"
                )
                cmidrules.append(
                    f"\\cmidrule(lr){{{cur_col}-{cur_col + span - 1}}}"
                )
                cur_col += span
            lines.append(" & ".join(cells_sp) + " \\\\")
            lines.append("".join(cmidrules))

        # Header
        hdr = " & ".join(
            [""] + [self._esc_latex(n) for n in self.model_labels]
        ) + " \\\\"
        lines.append(hdr)

        # Dep-var
        if self.dep_var_labels:
            dvr = " & ".join(
                [""] + [f"\\textit{{{self._esc_latex(dv)}}}" for dv in self.dep_var_labels]
            ) + " \\\\"
            lines.append(dvr)

        lines.append("\\hline")

        # Panels
        multi = len(self.panels) > 1
        for pi, panel in enumerate(self.panels):
            if multi and self.panel_labels and pi < len(self.panel_labels):
                lines.append(
                    f"\\multicolumn{{{n_cols}}}{{l}}"
                    f"{{\\textbf{{{self._esc_latex(self.panel_labels[pi])}}}}}"
                    " \\\\"
                )
                lines.append("\\hline")

            var_list = self._resolve_vars(panel.models)
            for var in var_list:
                label = self._esc_latex(self.coef_labels.get(var, var))
                cells: List[str] = []
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            cells.append(_latex_escape(self._coef_cell(m, var, flat_idx)))
                        else:
                            cells.append("")
                        flat_idx += 1
                lines.append(f"{label} & " + " & ".join(cells) + " \\\\")
                # SE row
                cells2: List[str] = []
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            cells2.append(_latex_escape(self._se_cell(m, var, flat_idx)))
                        else:
                            cells2.append("")
                        flat_idx += 1
                lines.append(" & " + " & ".join(cells2) + " \\\\")
                # Extra SE rows (multi_se)
                base_idx = sum(len(p.models) for p in self.panels[:pi])
                for ext_idx, (_, per_model_list) in enumerate(self.multi_se):
                    cells_ext: List[str] = []
                    for gi, p2 in enumerate(self.panels):
                        for off, m in enumerate(p2.models):
                            if gi == pi:
                                per_model = per_model_list[base_idx + off]
                                cells_ext.append(
                                    _latex_escape(self._multi_se_cell(
                                        per_model, var, ext_idx, m, base_idx + off
                                    ))
                                )
                            else:
                                cells_ext.append("")
                    lines.append(" & " + " & ".join(cells_ext) + " \\\\")

            if multi and pi < len(self.panels) - 1:
                lines.append("\\hline")

        lines.append("\\hline")

        # Add rows
        for row_label, row_vals in self.add_rows.items():
            cells_ar: List[str] = []
            for i in range(len(all_models)):
                val = row_vals[i] if i < len(row_vals) else ""
                cells_ar.append(self._esc_latex(val))
            lines.append(
                f"{self._esc_latex(row_label)} & " + " & ".join(cells_ar) + " \\\\"
            )

        if self.add_rows:
            lines.append("\\hline")

        # Stats
        for key in self._stat_keys:
            disp = _STAT_DISPLAY.get(key, key)
            if key == "R-squared":
                disp = "R$^2$"
            elif key == "Adj. R-squared":
                disp = "Adj. R$^2$"
            else:
                disp = _latex_escape(disp)
            cells_s = [self._stat_cell(m, key) for m in all_models]
            lines.append(f"{disp} & " + " & ".join(cells_s) + " \\\\")

        # Hypothesis-test rows
        for row_label, row_vals in self.tests_rows:
            cells_t: List[str] = []
            for i in range(len(all_models)):
                val = row_vals[i] if i < len(row_vals) else ""
                cells_t.append(self._esc_latex(val))
            lines.append(
                f"{self._esc_latex(row_label)} & " + " & ".join(cells_t) + " \\\\"
            )

        lines.append("\\hline\\hline")

        # Notes
        note_line = f"{self._se_label()} in parentheses"
        lines.append(
            f"\\multicolumn{{{n_cols}}}{{l}}"
            f"{{\\footnotesize {_latex_escape(note_line)}}} \\\\"
        )
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            multi_note = f"{label} in {lo}…{hi}"
            lines.append(
                f"\\multicolumn{{{n_cols}}}{{l}}"
                f"{{\\footnotesize {_latex_escape(multi_note)}}} \\\\"
            )
        if self._has_any_eform():
            lines.append(
                f"\\multicolumn{{{n_cols}}}{{l}}"
                f"{{\\footnotesize {_latex_escape(self._eform_note())}}} \\\\"
            )
        if self.show_stars:
            lines.append(
                f"\\multicolumn{{{n_cols}}}{{l}}"
                f"{{\\footnotesize {_latex_escape(self._star_note())}}} \\\\"
            )
        for note in self.notes:
            lines.append(
                f"\\multicolumn{{{n_cols}}}{{l}}"
                f"{{\\footnotesize {_latex_escape(note)}}} \\\\"
            )

        lines.append("\\end{tabular}")
        lines.append("\\end{table}")
        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # Markdown
    # ═══════════════════════════════════════════════════════════════════════

    def to_markdown(self, *, quarto: bool = False) -> str:
        """Render the table as Markdown.

        Parameters
        ----------
        quarto : bool, default False
            When ``True``, append a Quarto cross-reference caption block
            of the form ``: <caption> {#tbl-<label>}`` so the table can be
            referenced via ``@tbl-<label>`` in the manuscript. Requires
            ``quarto_label`` to have been set on the result (typically via
            ``regtable(..., quarto_label="main")``). Equivalent to calling
            :meth:`to_quarto`.
        """
        if quarto:
            return self.to_quarto()
        all_models = self._all_models_flat()
        lines: List[str] = []
        if self.title:
            lines.append(f"**{self.title}**")
            lines.append("")

        # Column spanners — Markdown has no native colspan, so we render
        # each spanner block as repeated label cells (one per column it
        # covers) inside a bold formatting band. Pandoc / GFM renderers
        # show this as a centered visual group.
        if self.column_spanners:
            spanner_cells: List[str] = []
            for lbl, span in self.column_spanners:
                spanner_cells.extend([f"**{lbl}**"] * span)
            lines.append(
                "| |"
                + "|".join(f" {c} " for c in spanner_cells)
                + "|"
            )

        # Header
        hdr = "| |" + "|".join(f" {n} " for n in self.model_labels) + "|"
        sep = "|---|" + "|".join("---:" for _ in self.model_labels) + "|"
        lines.append(hdr)
        lines.append(sep)

        multi = len(self.panels) > 1
        for pi, panel in enumerate(self.panels):
            if multi and self.panel_labels and pi < len(self.panel_labels):
                lines.append(
                    f"| **{self.panel_labels[pi]}** |"
                    + "|".join(" " for _ in self.model_labels)
                    + "|"
                )

            var_list = self._resolve_vars(panel.models)
            for var in var_list:
                label = self.coef_labels.get(var, var)
                cells: List[str] = []
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            cells.append(self._coef_cell(m, var, flat_idx))
                        else:
                            cells.append("")
                        flat_idx += 1
                lines.append(f"| {label} |" + "|".join(f" {c} " for c in cells) + "|")
                # SE row
                cells2: List[str] = []
                flat_idx = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        if gi == pi:
                            cells2.append(self._se_cell(m, var, flat_idx))
                        else:
                            cells2.append("")
                        flat_idx += 1
                lines.append("| |" + "|".join(f" {c} " for c in cells2) + "|")
                # Extra SE rows from multi_se
                base_idx = sum(len(p.models) for p in self.panels[:pi])
                for ext_idx, (_, per_model_list) in enumerate(self.multi_se):
                    cells3: List[str] = []
                    for gi, p2 in enumerate(self.panels):
                        for off, m in enumerate(p2.models):
                            if gi == pi:
                                per_model = per_model_list[base_idx + off]
                                cells3.append(self._multi_se_cell(
                                    per_model, var, ext_idx, m, base_idx + off
                                ))
                            else:
                                cells3.append("")
                    lines.append("| |" + "|".join(f" {c} " for c in cells3) + "|")

        # Add rows
        for row_label, row_vals in self.add_rows.items():
            cells_ar: List[str] = []
            for i in range(len(all_models)):
                val = row_vals[i] if i < len(row_vals) else ""
                cells_ar.append(val)
            lines.append(
                f"| {row_label} |" + "|".join(f" {c} " for c in cells_ar) + "|"
            )

        # Stats
        for key in self._stat_keys:
            disp = _STAT_DISPLAY.get(key, key)
            cells_s = [self._stat_cell(m, key) for m in all_models]
            lines.append(f"| {disp} |" + "|".join(f" {c} " for c in cells_s) + "|")

        # Hypothesis-test rows
        for row_label, row_vals in self.tests_rows:
            cells_t = [
                row_vals[i] if i < len(row_vals) else ""
                for i in range(len(all_models))
            ]
            lines.append(
                f"| {row_label} |" + "|".join(f" {c} " for c in cells_t) + "|"
            )

        lines.append("")
        lines.append(f"*{self._se_label()} in parentheses*")
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            lines.append(f"*{label} in {lo}…{hi}*")
        if self._has_any_eform():
            lines.append(f"*{self._eform_note()}*")
        if self.show_stars:
            lines.append(f"*{self._star_note()}*")
        for note in self.notes:
            lines.append(f"*{note}*")

        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # Quarto
    # ═══════════════════════════════════════════════════════════════════════

    def to_quarto(self) -> str:
        """Render as a Quarto-cross-referenceable Markdown table.

        Builds on :meth:`to_markdown` and appends a Quarto caption block
        of the form::

            : <caption> {#tbl-<label>}

        which lets the manuscript reference the table via
        ``@tbl-<label>``. The ``tbl-`` prefix is auto-prepended when the
        user passes a bare ``quarto_label="main"``.

        Behaviour
        ---------
        - ``quarto_label`` is required. Without it, ``ValueError`` is
          raised — Quarto cross-references need an id.
        - ``quarto_caption`` falls back to ``title`` when not provided.
          If neither is set, a generic ``"Regression results"`` is used
          and a warning is emitted.
        - The leading title line is dropped (the caption block replaces
          it) to avoid duplicating the heading.
        """
        if not self.quarto_label:
            raise ValueError(
                "to_quarto() requires quarto_label to be set. "
                "Pass quarto_label='main' (or similar) to regtable()."
            )

        raw_label = str(self.quarto_label).strip()
        label = raw_label if raw_label.startswith("tbl-") else f"tbl-{raw_label}"

        if self.quarto_caption:
            caption = str(self.quarto_caption)
        elif self.title:
            caption = str(self.title)
        else:
            warnings.warn(
                "to_quarto(): no quarto_caption or title provided; "
                "using default 'Regression results'. Quarto cross-refs "
                "render better with an explicit caption.",
                UserWarning,
                stacklevel=2,
            )
            caption = "Regression results"

        saved_title = self.title
        try:
            self.title = None
            body = self.to_markdown()
        finally:
            self.title = saved_title

        body = body.rstrip()
        return f"{body}\n\n: {caption} {{#{label}}}\n"

    # ═══════════════════════════════════════════════════════════════════════
    # DataFrame
    # ═══════════════════════════════════════════════════════════════════════

    def to_dataframe(self) -> pd.DataFrame:
        """Return the table as a pandas DataFrame."""
        all_models = self._all_models_flat()
        records: List[Dict[str, str]] = []

        multi = len(self.panels) > 1
        for pi, panel in enumerate(self.panels):
            if multi and self.panel_labels and pi < len(self.panel_labels):
                row_ph: Dict[str, str] = {"": self.panel_labels[pi]}
                for n in self.model_labels:
                    row_ph[n] = ""
                records.append(row_ph)

            var_list = self._resolve_vars(panel.models)
            for var in var_list:
                label = self.coef_labels.get(var, var)
                row: Dict[str, str] = {"": label}
                mi = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        col_name = self.model_labels[mi]
                        if gi == pi:
                            row[col_name] = self._coef_cell(m, var, mi)
                        else:
                            row[col_name] = ""
                        mi += 1
                records.append(row)
                # SE row
                row2: Dict[str, str] = {"": ""}
                mi = 0
                for gi, p2 in enumerate(self.panels):
                    for m in p2.models:
                        col_name = self.model_labels[mi]
                        if gi == pi:
                            row2[col_name] = self._se_cell(m, var, mi)
                        else:
                            row2[col_name] = ""
                        mi += 1
                records.append(row2)
                # Extra SE rows (multi_se)
                base_idx = sum(len(p.models) for p in self.panels[:pi])
                for ext_idx, (_, per_model_list) in enumerate(self.multi_se):
                    row3: Dict[str, str] = {"": ""}
                    mi = 0
                    for gi, p2 in enumerate(self.panels):
                        for off, m in enumerate(p2.models):
                            col_name = self.model_labels[mi]
                            if gi == pi:
                                per_model = per_model_list[base_idx + off]
                                row3[col_name] = self._multi_se_cell(
                                    per_model, var, ext_idx, m, base_idx + off
                                )
                            else:
                                row3[col_name] = ""
                            mi += 1
                    records.append(row3)

        # Add rows
        for row_label, row_vals in self.add_rows.items():
            row_ar: Dict[str, str] = {"": row_label}
            for i, lbl in enumerate(self.model_labels):
                row_ar[lbl] = row_vals[i] if i < len(row_vals) else ""
            records.append(row_ar)

        # Stats
        for key in self._stat_keys:
            disp = _STAT_DISPLAY.get(key, key)
            row_s: Dict[str, str] = {"": disp}
            for i, m in enumerate(all_models):
                row_s[self.model_labels[i]] = self._stat_cell(m, key)
            records.append(row_s)

        # Hypothesis-test rows
        for row_label, row_vals in self.tests_rows:
            row_t: Dict[str, str] = {"": row_label}
            for i, lbl in enumerate(self.model_labels):
                row_t[lbl] = row_vals[i] if i < len(row_vals) else ""
            records.append(row_t)

        df = pd.DataFrame(records)
        df = df.set_index("")
        df.index.name = None
        return df

    # ──────────────────────────────────────────────────────────────────────
    # Agent-native serialisation
    # ──────────────────────────────────────────────────────────────────────

    @staticmethod
    def _jsonable(value: Any) -> Any:
        """Coerce a scalar to a JSON-safe primitive (NaN/Inf → ``None``)."""
        if value is None:
            return None
        if isinstance(value, (bool, str)):
            return value
        if isinstance(value, (int, np.integer)):
            return int(value)
        if isinstance(value, (float, np.floating)):
            v = float(value)
            return v if np.isfinite(v) else None
        if isinstance(value, np.ndarray):
            return [RegtableResult._jsonable(x) for x in value.tolist()]
        try:
            if pd.isna(value):
                return None
        except (TypeError, ValueError):
            pass
        return str(value)

    def to_dict(
        self,
        *,
        renders: Union[bool, Sequence[str], None] = None,
    ) -> Dict[str, Any]:
        """Return a JSON-safe dict representation of the table.

        The package is agent-native (CLAUDE.md §1): a rendered regression
        table is a first-class artifact an LLM tool loop should be able to
        serialise, cache, and reason over without re-rendering.  The payload
        carries three layers:

        - **metadata** — ``model_labels``, ``dep_var_labels``,
          ``panel_labels``, ``title``, ``notes``, ``template``, ``se_type``,
          ``stars`` / ``star_levels``, ``requested_stats``, ``coef_labels``.
        - **table** — the *rendered* cell grid (the same strings
          :meth:`to_dataframe` produces: ``"2.067***"``, ``"(0.074)"`` …),
          as a list of ``{"term": ..., <model label>: <cell>}`` records.
        - **models** — the *numeric* truth per model: coefficient estimates,
          standard errors, t / p values, confidence bounds, summary stats and
          the dependent variable.  Use this layer for machine reasoning; use
          ``table`` for faithful re-display.

        Parameters
        ----------
        renders : bool | sequence of str, optional
            When truthy, also embed fully rendered strings under
            ``"renders"``.  ``True`` embeds ``latex`` / ``html`` / ``markdown``
            / ``text``; a sequence selects specific formats
            (e.g. ``renders=["latex"]``).  Default ``None`` keeps the payload
            compact.

        Returns
        -------
        dict
            JSON-safe; round-trips through ``json.dumps``.

        Examples
        --------
        >>> import statspai as sp
        >>> tbl = sp.regtable(m1, m2, template="aer")
        >>> payload = tbl.to_dict()
        >>> payload["models"][0]["coefficients"]["x"]["estimate"]  # doctest: +SKIP
        2.067
        """
        df = self.to_dataframe()
        table_rows: List[Dict[str, Any]] = []
        for idx, row in df.iterrows():
            rec: Dict[str, Any] = {
                "term": "" if idx is None else str(idx)
            }
            for col in df.columns:
                val = row[col]
                rec[str(col)] = "" if pd.isna(val) else str(val)
            table_rows.append(rec)

        models: List[Dict[str, Any]] = []
        flat = self._all_models_flat()
        for i, m in enumerate(flat):
            label = (self.model_labels[i]
                     if i < len(self.model_labels) else f"({i + 1})")
            coefs: Dict[str, Dict[str, Any]] = {}
            for term in list(m.params.index):
                t = str(term)
                coefs[t] = {
                    "estimate": self._jsonable(m.params.get(term)),
                    "std_error": self._jsonable(m.std_errors.get(term)),
                    "t_statistic": self._jsonable(
                        m.tvalues.get(term) if m.tvalues is not None else None),
                    "p_value": self._jsonable(
                        m.pvalues.get(term) if m.pvalues is not None else None),
                    "conf_low": self._jsonable(
                        m.conf_int_lower.get(term)
                        if m.conf_int_lower is not None else None),
                    "conf_high": self._jsonable(
                        m.conf_int_upper.get(term)
                        if m.conf_int_upper is not None else None),
                }
            models.append({
                "label": str(label),
                "depvar": self._jsonable(getattr(m, "depvar", None)),
                "coefficients": coefs,
                "stats": {str(k): self._jsonable(v)
                          for k, v in (getattr(m, "stats", {}) or {}).items()},
                "df_resid": self._jsonable(getattr(m, "df_resid", None)),
            })

        payload: Dict[str, Any] = {
            "kind": "regression_table",
            "n_models": self.n_models,
            "n_panels": len(self.panels),
            "model_labels": [str(x) for x in self.model_labels],
            "dep_var_labels": (
                [str(x) for x in self.dep_var_labels]
                if self.dep_var_labels else None),
            "panel_labels": (
                [str(x) for x in self.panel_labels]
                if self.panel_labels else None),
            "title": self.title,
            "notes": [str(x) for x in self.notes],
            "template": self.template,
            "se_type": self.se_type,
            "stars": bool(self.show_stars),
            "star_levels": [self._jsonable(x) for x in self.star_levels],
            "requested_stats": [str(x) for x in self.requested_stats],
            "coef_labels": {str(k): str(v)
                            for k, v in self.coef_labels.items()},
            "columns": ["term"] + [str(c) for c in df.columns],
            "table": table_rows,
            "models": models,
            # Render-controlling parameters so from_dict() can faithfully
            # reconstruct a RegtableResult that re-renders identically for the
            # common feature set (coefficients / SE / stats / labels / keep /
            # drop / order / fmt). Exotic features (multi_se, eform,
            # column_spanners, tests, apply_coef) are intentionally NOT in the
            # payload and do not survive a round-trip — see from_dict().
            "render_spec": {
                "fmt": self.fmt,
                "alpha": self._jsonable(self.alpha),
                "panel_sizes": [len(p.models) for p in self.panels],
                "add_rows": {str(k): [str(x) for x in v]
                             for k, v in self.add_rows.items()},
                "keep": list(self.keep) if self.keep else None,
                "drop": sorted(str(x) for x in self.drop) if self.drop
                        else None,
                "order": list(self.order) if self.order else None,
                "se_label": self._se_label_override,
            },
        }

        if renders:
            if renders is True:
                wanted = ["latex", "html", "markdown", "text"]
            else:
                wanted = [str(f) for f in renders]
            renderers = {
                "latex": self.to_latex,
                "html": self.to_html,
                "markdown": self.to_markdown,
                "text": self.to_text,
            }
            rendered: Dict[str, str] = {}
            for fmt in wanted:
                fn = renderers.get(fmt)
                if fn is None:
                    raise ValueError(
                        f"renders={fmt!r} is not one of "
                        f"{sorted(renderers)}"
                    )
                rendered[fmt] = fn()
            payload["renders"] = rendered

        return payload

    def to_json(
        self,
        *,
        indent: Optional[int] = None,
        renders: Union[bool, Sequence[str], None] = None,
    ) -> str:
        """Serialise :meth:`to_dict` via ``json.dumps``."""
        import json
        return json.dumps(
            self.to_dict(renders=renders), indent=indent, default=str
        )

    @classmethod
    def from_dict(cls, payload: Dict[str, Any]) -> "RegtableResult":
        """Reconstruct a :class:`RegtableResult` from a :meth:`to_dict` payload.

        The inverse of :meth:`to_dict`.  Rebuilds one normalised model per
        entry in the ``models`` layer (coefficient estimates, SE, t, p, CI,
        summary stats, dependent variable), re-splits them into panels per
        ``render_spec.panel_sizes``, and restores the render-controlling
        metadata (labels, ``fmt``, ``se_type``, stars / star levels, stats,
        ``keep`` / ``drop`` / ``order``, ``add_rows``, ``alpha``, template).
        For a table built without exotic options this round-trips exactly:

        >>> import statspai as sp
        >>> t = sp.regtable(m1, m2, template="aer")        # doctest: +SKIP
        >>> RegtableResult.from_dict(t.to_dict()).to_latex() == t.to_latex()
        True

        Notes
        -----
        Exotic features that are not part of the ``to_dict`` payload —
        stacked ``multi_se`` rows, ``eform`` transforms, ``column_spanners``,
        ``tests`` rows, and custom ``apply_coef`` transforms — are NOT
        preserved across the round-trip (they reconstruct as a plain table).
        The serialised ``table`` / ``renders`` layers already capture their
        rendered form if you only need to re-display, not re-compute.
        """
        if not isinstance(payload, dict) or \
                payload.get("kind") != "regression_table":
            raise ValueError(
                "from_dict() expects a RegtableResult.to_dict() payload "
                "(kind == 'regression_table')."
            )

        spec = payload.get("render_spec", {}) or {}

        def _series(coefs: Dict[str, Any], field: str,
                    terms: List[str]) -> pd.Series:
            return pd.Series(
                {t: coefs[t].get(field) for t in terms}, dtype=float
            )

        def _build_model(m: Dict[str, Any]) -> _ModelData:
            coefs = m.get("coefficients", {}) or {}
            terms = list(coefs.keys())
            return _ModelData(
                params=_series(coefs, "estimate", terms),
                std_errors=_series(coefs, "std_error", terms),
                tvalues=_series(coefs, "t_statistic", terms),
                pvalues=_series(coefs, "p_value", terms),
                conf_int_lower=_series(coefs, "conf_low", terms),
                conf_int_upper=_series(coefs, "conf_high", terms),
                stats=dict(m.get("stats", {}) or {}),
                depvar=m.get("depvar") or "",
                df_resid=m.get("df_resid"),
            )

        models = [_build_model(m) for m in payload.get("models", [])]
        sizes = spec.get("panel_sizes") or [len(models)]
        # Guard against a sizes/models mismatch (e.g. hand-edited payload):
        # fall back to a single panel rather than silently dropping models.
        if sum(int(s) for s in sizes) != len(models):
            sizes = [len(models)]
        panels: List[_PanelData] = []
        cursor = 0
        for sz in sizes:
            sz = int(sz)
            panels.append(_PanelData(models[cursor:cursor + sz]))
            cursor += sz

        star_levels = payload.get("star_levels")
        return cls(
            panels,
            panel_labels=payload.get("panel_labels"),
            model_labels=list(payload.get("model_labels", [])),
            dep_var_labels=payload.get("dep_var_labels"),
            coef_labels=dict(payload.get("coef_labels", {}) or {}),
            keep=spec.get("keep"),
            drop=spec.get("drop"),
            order=spec.get("order"),
            se_type=payload.get("se_type", "se"),
            stars=bool(payload.get("stars", True)),
            star_levels=(tuple(star_levels) if star_levels
                         else (0.10, 0.05, 0.01)),
            fmt=spec.get("fmt", "%.3f"),
            title=payload.get("title"),
            notes=list(payload.get("notes", []) or []),
            add_rows={k: list(v)
                      for k, v in (spec.get("add_rows", {}) or {}).items()},
            stats=list(payload.get("requested_stats", []) or []) or None,
            alpha=float(spec.get("alpha", 0.05) or 0.05),
            se_label=spec.get("se_label"),
            template=payload.get("template"),
        )

    # ═══════════════════════════════════════════════════════════════════════
    # Excel
    # ═══════════════════════════════════════════════════════════════════════

    def to_excel(self, filename: str) -> None:
        """Export table to Excel as a strict book-tab three-line table.

        Uses the shared ``_excel_style`` primitives so the visual output
        is byte-aligned with ``sumstats``, ``tab``, ``paper_tables``,
        ``collection``, ``modelsummary`` and ``outreg2``: thick top rule
        above the column header, thin mid rule between header and body,
        thick bottom rule below the last data row, Times New Roman
        throughout.
        """
        try:
            import openpyxl
            from openpyxl.styles import Alignment, Font
        except ImportError:
            warnings.warn(
                "openpyxl is required for Excel export. "
                "Install with: pip install openpyxl"
            )
            return

        from ._excel_style import (
            BODY_PT, HEADER_PT, NOTES_PT, TIMES,
            apply_booktab_borders, autofit_columns, write_title,
        )

        df = self.to_dataframe()
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Regression Table"

        header_font = Font(bold=True, name=TIMES, size=HEADER_PT)
        body_font = Font(name=TIMES, size=BODY_PT)
        notes_font = Font(italic=True, name=TIMES, size=NOTES_PT)
        center = Alignment(horizontal="center")
        left = Alignment(horizontal="left")

        n_cols = len(df.columns) + 1  # +1 for the row-label column
        row_idx = 1
        if self.title:
            row_idx = write_title(ws, row_idx, n_cols, self.title)

        # Column spanners (above the model-label header) when set —
        # rendered as merged cells centered over each column block.
        header_top_row = row_idx
        if self.column_spanners:
            cur_col = 2  # col 1 reserved for row-label
            for lbl, span in self.column_spanners:
                top_cell = ws.cell(row=row_idx, column=cur_col, value=str(lbl))
                top_cell.font = header_font
                top_cell.alignment = center
                if span > 1:
                    ws.merge_cells(
                        start_row=row_idx, start_column=cur_col,
                        end_row=row_idx, end_column=cur_col + span - 1,
                    )
                cur_col += span
            row_idx += 1

        # Header row
        header_label_row = row_idx
        c0 = ws.cell(row=row_idx, column=1, value="")
        c0.font = header_font
        for j, col in enumerate(df.columns, 2):
            cell = ws.cell(row=row_idx, column=j, value=col)
            cell.font = header_font
            cell.alignment = center
        header_bot_row = row_idx
        row_idx += 1

        # Body rows
        body_top_row = row_idx
        for idx, row_data in df.iterrows():
            c0 = ws.cell(row=row_idx, column=1, value=str(idx))
            c0.font = body_font
            c0.alignment = left
            for j, val in enumerate(row_data, 2):
                cell = ws.cell(row=row_idx, column=j, value=str(val))
                cell.font = body_font
                cell.alignment = center
            row_idx += 1
        body_bot_row = row_idx - 1

        apply_booktab_borders(
            ws,
            header_top_row=header_top_row,
            header_bot_row=header_bot_row,
            body_top_row=body_top_row,
            body_bot_row=body_bot_row,
            n_cols=n_cols,
        )

        # Notes — emit the same lines that to_text/to_html/to_latex/to_word
        # emit so users who pass multi_se / repro / notes do not lose them
        # when exporting to Excel.
        note_row = body_bot_row + 1
        ws.cell(
            row=note_row, column=1,
            value=f"{self._se_label()} in parentheses"
        ).font = notes_font
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            note_row += 1
            ws.cell(row=note_row, column=1,
                    value=f"{label} in {lo}…{hi}").font = notes_font
        if self.show_stars:
            note_row += 1
            ws.cell(row=note_row, column=1,
                    value=self._star_note()).font = notes_font
        for note in self.notes:
            note_row += 1
            ws.cell(row=note_row, column=1, value=note).font = notes_font

        autofit_columns(ws, n_cols, max_width=25)
        wb.save(filename)

    # ═══════════════════════════════════════════════════════════════════════
    # Word (docx)
    # ═══════════════════════════════════════════════════════════════════════

    def to_word(self, filename: str) -> None:
        """Export table to Word (.docx) file in AER/QJE book-tab style.

        The exported document follows economics-journal conventions:
        a heavy top rule, thin mid rule below the header, heavy bottom
        rule above notes, and **no** internal vertical borders. Body
        text is Times New Roman 10pt; the notes paragraph is 8pt italic.
        """
        try:
            from docx import Document
        except ImportError:
            warnings.warn(
                "python-docx is required for Word export. "
                "Install with: pip install python-docx"
            )
            return

        from ._aer_style import (
            apply_word_booktab_rules,
            style_word_table_typography,
            add_word_notes_paragraph,
        )

        doc = Document()
        if self.title:
            doc.add_heading(self.title, level=2)

        df = self.to_dataframe()
        # When column_spanners are set we add an extra header row above
        # the model-label row and merge cells across each column block.
        spanner_extra = 1 if self.column_spanners else 0
        n_rows = len(df) + 1 + spanner_extra
        n_cols = len(df.columns) + 1
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.autofit = True

        spanner_row_idx = 0 if self.column_spanners else None
        header_row_idx = 1 if self.column_spanners else 0
        body_start_idx = header_row_idx + 1

        # Spanner row (above the model-label row) when set
        if self.column_spanners:
            cur_col = 1  # col 0 reserved for row-label
            spanner_row = table.rows[spanner_row_idx]
            spanner_row.cells[0].text = ""
            for lbl, span in self.column_spanners:
                spanner_row.cells[cur_col].text = str(lbl)
                if span > 1:
                    end_col = cur_col + span - 1
                    spanner_row.cells[cur_col].merge(spanner_row.cells[end_col])
                cur_col += span

        # Populate header (model labels)
        header_row = table.rows[header_row_idx]
        header_row.cells[0].text = ""
        for j, col in enumerate(df.columns, 1):
            header_row.cells[j].text = str(col)
        # Populate body
        for i, (idx, row_data) in enumerate(df.iterrows(), body_start_idx):
            table.rows[i].cells[0].text = str(idx)
            for j, val in enumerate(row_data, 1):
                table.rows[i].cells[j].text = str(val)

        header_rows = (
            (spanner_row_idx, header_row_idx)
            if spanner_row_idx is not None
            else (header_row_idx,)
        )
        style_word_table_typography(table, header_rows=header_rows)
        apply_word_booktab_rules(
            table,
            header_top_idx=spanner_row_idx if spanner_row_idx is not None else header_row_idx,
            header_bot_idx=header_row_idx,
        )

        # Notes (italic, 8pt)
        note_lines = [f"{self._se_label()} in parentheses"]
        for ext_idx, (label, _) in enumerate(self.multi_se):
            lo, hi = _MULTI_SE_BRACKETS[ext_idx % len(_MULTI_SE_BRACKETS)]
            note_lines.append(f"{label} in {lo}…{hi}")
        if self._has_any_eform():
            note_lines.append(self._eform_note())
        if self.show_stars:
            note_lines.append(self._star_note())
        note_lines.extend(self.notes)
        add_word_notes_paragraph(doc, "\n".join(note_lines))

        doc.save(filename)

    def to_docx(self, filename: str) -> None:
        """Alias for :meth:`to_word` — mirrors Stata ``outreg2`` convention."""
        self.to_word(filename)

    # ═══════════════════════════════════════════════════════════════════════
    # Save (auto-detect from extension)
    # ═══════════════════════════════════════════════════════════════════════

    def save(self, filename: str) -> None:
        """Auto-detect format from file extension and save."""
        path = Path(filename)
        ext = path.suffix.lower()

        if ext in (".xlsx", ".xls"):
            self.to_excel(filename)
        elif ext == ".docx":
            self.to_word(filename)
        elif ext == ".tex":
            path.write_text(self.to_latex(), encoding="utf-8")
        elif ext in (".html", ".htm"):
            path.write_text(self.to_html(), encoding="utf-8")
        elif ext == ".md":
            path.write_text(self.to_markdown(), encoding="utf-8")
        elif ext == ".qmd":
            path.write_text(self.to_quarto(), encoding="utf-8")
        elif ext == ".csv":
            self.to_dataframe().to_csv(filename)
        elif ext == ".json":
            path.write_text(self.to_json(indent=2), encoding="utf-8")
        else:
            path.write_text(self.to_text(), encoding="utf-8")

    # ═══════════════════════════════════════════════════════════════════════
    # Dunder methods
    # ═══════════════════════════════════════════════════════════════════════

    def _render(self, fmt: str) -> str:
        return {
            "text": self.to_text,
            "latex": self.to_latex,
            "tex": self.to_latex,
            "html": self.to_html,
            "markdown": self.to_markdown,
            "md": self.to_markdown,
            "quarto": self.to_quarto,
            "qmd": self.to_quarto,
        }.get(fmt, self.to_text)()

    def __str__(self) -> str:
        return self._render(self._output)

    def __repr__(self) -> str:
        return self.__str__()

    def _repr_html_(self) -> str:
        return self.to_html()


# ═══════════════════════════════════════════════════════════════════════════
# _PanelData: lightweight container for a group of models
# ═══════════════════════════════════════════════════════════════════════════

class _PanelData:
    __slots__ = ("models",)

    def __init__(self, models: List[_ModelData]):
        self.models = models


# ═══════════════════════════════════════════════════════════════════════════
# regtable() — the main public API
# ═══════════════════════════════════════════════════════════════════════════

def regtable(
    *args,
    panel_labels: Optional[List[str]] = None,
    coef_labels: Optional[Dict[str, str]] = None,
    dep_var_labels: Optional[List[str]] = None,
    model_labels: Optional[List[str]] = None,
    keep: Optional[Sequence[str]] = None,
    drop: Optional[Sequence[str]] = None,
    order: Optional[Sequence[str]] = None,
    stats: Optional[Sequence[str]] = None,
    se_type: str = "se",
    stars: bool = True,
    star_levels: Optional[Tuple[float, ...]] = None,
    fmt: str = "%.3f",
    output: str = "text",
    filename: Optional[str] = None,
    title: Optional[str] = None,
    notes: Optional[List[str]] = None,
    add_rows: Optional[Dict[str, List[str]]] = None,
    alpha: float = 0.05,
    template: Optional[str] = None,
    diagnostics: Union[str, bool] = "auto",
    multi_se: Optional[Dict[str, Sequence[Any]]] = None,
    repro: Union[bool, Dict[str, Any], None] = None,
    quarto_label: Optional[str] = None,
    quarto_caption: Optional[str] = None,
    eform: Union[bool, Sequence[bool]] = False,
    column_spanners: Optional[Sequence[Tuple[str, int]]] = None,
    coef_map: Optional[Dict[str, str]] = None,
    consistency_check: bool = True,
    estimate: Optional[str] = None,
    statistic: Optional[str] = None,
    notation: Union[str, Tuple[str, ...]] = "stars",
    apply_coef: Optional[Any] = None,
    apply_coef_deriv: Optional[Any] = None,
    escape: bool = True,
    tests: Optional[Dict[str, Sequence[Any]]] = None,
    fixef_sizes: bool = False,
    vcov: Optional[str] = None,
    transpose: bool = False,
) -> RegtableResult:
    """
    Unified publication-quality regression table.

    Accepts model results as positional arguments. If the first argument
    is a list, each list is treated as a separate panel.

    Parameters
    ----------
    *args : model results or lists of model results
        ``EconometricResults``, ``CausalResult``, or any duck-typed object
        with ``params`` / ``std_errors`` attributes. Pass multiple lists
        to create a multi-panel table.
    panel_labels : list of str, optional
        Labels for each panel (e.g., ``["Panel A: Wages", "Panel B: Hours"]``).
    coef_labels : dict, optional
        Rename variables: ``{"education": "Years of Education"}``.
    dep_var_labels : list of str, optional
        Dependent variable labels shown below column headers.
    model_labels : list of str, optional
        Column header labels. Defaults to ``(1), (2), ...``.
    keep : list of str, optional
        Only show these variables.
    drop : list of str, optional
        Hide these variables.
    order : list of str, optional
        Reorder variables.
    stats : list of str, optional
        Summary statistics. Defaults to ``["N", "R2", "adj_R2", "F"]``.
    se_type : str, default ``"se"``
        What to show beneath coefficients: ``"se"``, ``"t"``, ``"p"``,
        or ``"ci"`` for confidence intervals.
    stars : bool, default True
        Append significance stars.
    star_levels : tuple, default ``(0.10, 0.05, 0.01)``
        Thresholds for ``*``, ``**``, ``***``.
    fmt : str, default ``"%.3f"``
        Format string for numeric values. Pass any C-style format
        (``"%.0f"``, ``"%.4f"``, ...) for fixed precision, or
        ``"auto"`` for magnitude-adaptive precision (recommended when
        a single table mixes dollar-magnitude coefficients like
        ``1521`` with elasticity-magnitude coefficients like ``0.288``
        — fixed ``"%.0f"`` would round the latter to ``0``).
    output : str, default ``"text"``
        Controls what ``str(result)`` / ``repr(result)`` / ``print(result)``
        returns — one of ``"text"``, ``"latex"``, ``"html"``, ``"markdown"``,
        ``"quarto"``, ``"word"``, ``"excel"``. In Jupyter, ``_repr_html_``
        always renders HTML regardless of this setting.
    filename : str, optional
        Save the table to this file path. The format is chosen from the
        **file extension** (``.tex``/``.html``/``.md``/``.qmd``/``.docx``/
        ``.xlsx``/``.csv``/``.json``), independently of ``output=``. The
        ``.json`` form writes the agent-native :meth:`RegtableResult.to_dict`
        payload. Pass a matching extension and ``output=`` to avoid surprises.
    quarto_label : str, optional
        Quarto cross-reference id. Pass ``"main"`` to make the table
        referenceable as ``@tbl-main`` from the manuscript prose. The
        ``tbl-`` prefix is auto-prepended when missing. Required for
        ``to_quarto()`` and ``output="quarto"``.
    quarto_caption : str, optional
        Caption rendered alongside the Quarto cross-ref id. Falls back
        to ``title`` when omitted; if both are absent, a generic
        ``"Regression results"`` is used and a warning is emitted.
    title : str, optional
        Table title / caption.
    notes : list of str, optional
        Additional notes beneath the table.
    add_rows : dict, optional
        Custom rows: ``{"Controls": ["No", "Yes", "Yes"]}``. User-provided
        rows take precedence over auto-extracted diagnostic rows with the
        same label.
    alpha : float, default 0.05
        Significance level used when ``se_type='ci'``. Displayed CI is
        ``(1 - alpha) * 100``%. With ``alpha=0.05`` (default) the bounds
        come from the model's stored 95% CI; for any other ``alpha`` the
        bounds are recomputed as ``b ± crit · se``, using the
        t-distribution when ``df_resid`` is known, else the standard
        normal.
    template : str, optional
        Journal preset name. One of ``"aer"``, ``"qje"``, ``"econometrica"``,
        ``"restat"``, ``"jf"``, ``"aeja"``, ``"jpe"``, ``"restud"``. When
        set, fills in defaults for ``star_levels``, the SE-row footer label
        (e.g. QJE → "Robust standard errors"), the default ``stats``
        selection (e.g. JF/AEJA include Adj. R²), and any extra notes —
        but every explicit kwarg you pass still wins. See
        :data:`statspai.output._journals.JOURNALS`.
    diagnostics : {'auto', 'off'} or bool, default ``'auto'``
        Auto-extract publication-quality diagnostic rows from the result
        objects:

        - **FE / Cluster indicators** — one row per distinct fixed effect
          variable (AER style: ``"Firm FE: Yes/No"``, ``"Year FE: Yes/No"``;
          interactions render as ``"Firm × Year FE"``), plus
          ``"Cluster SE: <var>"``. Falls back to a single
          ``"Fixed Effects: Yes/No"`` row when FE metadata is present but
          unparseable.
        - **IV** — first-stage F (Olea-Pflueger / KP), Hansen-J p.
        - **DiD** — pre-trend p-value, treated-group count.
        - **RD** — bandwidth, kernel, polynomial order.

        ``"auto"`` (and ``True``) emit only rows where at least one column
        produces a non-empty cell; ``False`` / ``"off"`` disables all
        auto-extraction. User-supplied ``add_rows`` always override.
    multi_se : dict, optional
        Stack additional SE specifications under the primary SE row.
        Keys are display labels (e.g. ``"Cluster SE"``, ``"Bootstrap SE"``)
        and values are sequences of :class:`pandas.Series` or dicts
        (one per model column) mapping coefficient names to SE values.
        Bracket styles cycle ``[]``/``{}``/``⟨⟩``/``||``. Footer notes
        record each label automatically.
    repro : bool or dict, optional
        Append a reproducibility metadata note (StatsPAI version, optional
        seed and data hash, timestamp) as the last footer line. ``True``
        emits the version + timestamp only. Pass a dict to record more:
        ``{"data": df, "seed": 42, "extra": "git@<sha>"}``.
    eform : bool or list of bool, default ``False``
        Report exponentiated coefficients — odds ratios for ``logit`` /
        ``probit``, incidence-rate ratios for ``poisson``, hazard ratios
        for Cox-style models. Standard errors use the delta method
        (``exp(b)·SE(b)``), CI bounds are ``(exp(lo), exp(hi))`` of the
        original endpoints, and t / p values are unchanged because
        ``H_0: b=0`` is equivalent to ``H_0: exp(b)=1``. Pass a per-model
        list (length matches ``n_models``) to mix transformed and
        untransformed columns (e.g. logit + OLS in the same table). A
        footer note transparently flags which columns are exponentiated.
    column_spanners : list of (label, span), optional
        Multi-row header above the model labels — each tuple groups
        ``span`` consecutive columns under ``label``. Spans must
        partition all model columns (sum equals ``n_models``). Renders
        as ``\\multicolumn{n}{c}{label}`` + ``\\cmidrule`` in LaTeX,
        ``colspan="n"`` in HTML, repeated bold cells in Markdown,
        and centered ASCII in text. Word and Excel exports inherit
        ``to_dataframe()``'s flat column model and currently omit the
        spanner row — use the LaTeX or HTML output for paper-grade
        spanners. Mirrors Stata ``mgroups()`` and R ``modelsummary``'s
        ``group`` argument. Example: ``column_spanners=[("OLS", 2),
        ("IV", 2)]`` over four models.
    coef_map : dict, optional
        Single-shot rename + reorder + drop. Mirrors R
        ``modelsummary``'s ``coef_map``: pass an ordered dict whose
        keys are coefficient names to **keep** (in display order) and
        values are the rendered labels. Variables not in ``coef_map``
        are dropped. Mutually exclusive with ``coef_labels`` /
        ``keep`` / ``drop`` / ``order`` — pass either the unified map
        or the legacy four-parameter spec.
    consistency_check : bool, default True
        When two or more columns are passed and their sample sizes
        differ, emit a ``UserWarning``. Reviewer red flag — disable by
        setting ``False`` (or annotate with ``notes=[...]``) when the
        N-mismatch is intentional (IV first stage on a subsample,
        RD bandwidth restriction, etc.).
    estimate : str, optional
        Custom format string for the **top** (coefficient) line in
        each cell. Mirrors R ``modelsummary``'s ``estimate=`` argument.
        Placeholders: ``{estimate}``, ``{stars}``, ``{std_error}``,
        ``{t_value}``, ``{p_value}``, ``{conf_low}``, ``{conf_high}``.
        Default ``"{estimate}{stars}"``. Pass e.g. ``"{stars}{estimate}"``
        for stars-first, or ``"{estimate} ({std_error}){stars}"`` for an
        inline single-line cell.
    statistic : str, optional
        Custom format string for the **bottom** (statistic) line in
        each cell. Same placeholders as ``estimate``. Default depends
        on ``se_type``: ``"({std_error})"`` for ``se``, ``"[{conf_low},
        {conf_high}]"`` for ``ci``, etc. Pass e.g. ``"t={t_value},
        p={p_value}"`` for working-paper-style cells.
    notation : ``"stars"`` | ``"symbols"`` | tuple of 3 strings
        Family of significance markers used when ``stars=True``.
        ``"stars"`` (default) → ``("*", "**", "***")``;
        ``"symbols"`` → ``("†", "‡", "§")`` (AER / JPE alternative
        when stars conflict with footnote markers); a 3-tuple of
        custom strings is accepted, ordered low → high.
    apply_coef : callable, optional
        Apply an arbitrary transformation ``f(b)`` to each rendered
        coefficient. Generalises ``eform`` (which is shorthand for
        ``apply_coef=np.exp``). Useful for percentage transforms
        (``apply_coef=lambda b: 100*b``), log scales, or signed
        sqrt for distortion measures. Pair with ``apply_coef_deriv``
        for delta-method SE rescaling.
        Mutually exclusive with ``eform``.
    apply_coef_deriv : callable, optional
        Derivative ``f'(b)`` of the ``apply_coef`` callable. When
        provided, SEs are rescaled as ``|f'(b)| · SE(b)``. When omitted,
        SEs stay on the original scale and a footer warns the reader.
    escape : bool, default True
        Auto-escape user-supplied label strings (``coef_labels``,
        ``model_labels``, ``panel_labels``, ``dep_var_labels``,
        ``column_spanners`` labels, ``add_rows`` labels and values,
        ``title``) for the active output format (LaTeX / HTML).
        Pass ``escape=False`` when those strings already contain raw
        markup you want to preserve verbatim — e.g. math-mode
        coefficient names like ``"$\\beta_1$"``, or HTML tags like
        ``"<i>β</i>"``. Cell content (numeric estimates, computed
        stats) is unaffected; it never contains user-controlled
        metacharacters.
    tests : dict, optional
        Render hypothesis-test rows in the diagnostic strip below the
        stats block. Keys are display labels ("F-test x1=0",
        "Hansen J p-value", "Wald χ²"); values are sequences whose
        length equals ``n_models``. Each per-model entry can be:

        - ``(statistic, pvalue)`` tuple → ``"<stat>***"`` (stars from p)
        - bare ``pvalue`` float        → ``"<p>***"``
        - ``None`` / ``NaN``           → empty cell
        - any string                   → passed through as-is

        Stars honour the configured ``notation`` family for cross-table
        consistency. Closes the gap to Stata ``estadd scalar`` /
        ``test`` integration where reviewers expect Wald / Sargan /
        Hansen-J / first-stage F right under the main results block.
    fixef_sizes : bool, default False
        Auto-emit "# Firm: 1,234" / "# Year: 30" rows showing the
        number of distinct levels per fixed effect. Reads
        ``model_info['n_fe_levels']`` from each result — currently
        populated by ``count.py`` (Poisson/NegBin) and the pyfixest
        adapter; other estimators silently no-op. Mirrors R fixest's
        ``etable(..., fixef_sizes=TRUE)``.
    vcov : str, optional
        Recompute the SE / t / p / 95% CI columns at print time using
        a different variance estimator. Currently supports OLS-style
        results that store ``data_info['X']`` and
        ``data_info['residuals']``:

        - ``"HC0"``                 — White heteroskedasticity-robust
        - ``"HC1"`` / ``"robust"``  — Stata's ``robust`` (HC0 × n/(n-k))
        - ``"HC2"``                 — leverage-weighted
        - ``"HC3"``                 — leverage-squared (recommended for
          small samples; Long-Ervin)

        Columns whose underlying result lacks the X/residuals fields
        emit a ``UserWarning`` and retain their fit-time SEs, so a
        heterogeneous mix of OLS + non-OLS does not blow up — the
        warning lists the affected columns so the user can audit.
    transpose : bool, default False
        Render with axes swapped: rows become models, columns become
        variables. Single-panel only; multi-panel input or
        ``multi_se=`` is rejected with ``NotImplementedError`` to keep
        the layout pivot semantics tight. Currently supports text and
        HTML renderers.

    Returns
    -------
    RegtableResult
        Object with ``.to_text()``, ``.to_latex()``, ``.to_html()``,
        ``.to_markdown()``, ``.to_excel(filename)``, ``.to_word(filename)``,
        ``.to_dataframe()``, ``.to_dict()`` / ``.to_json()`` (agent-native),
        ``.save(filename)`` methods.
        Renders as rich HTML in Jupyter notebooks via ``_repr_html_()``.

    Examples
    --------
    >>> import statspai as sp
    >>> m1 = sp.regress("y ~ x1", data=df)
    >>> m2 = sp.regress("y ~ x1 + x2", data=df)
    >>> sp.regtable(m1, m2)
    >>> sp.regtable(m1, m2, output="latex", filename="table1.tex")
    >>> sp.regtable([m1, m2], [m3, m4],
    ...     panel_labels=["Panel A: OLS", "Panel B: IV"])
    >>>
    >>> # Logit odds ratios
    >>> sp.regtable(sp.logit("y ~ x", data=df), eform=True)
    >>>
    >>> # IV three-block table with column spanners
    >>> sp.regtable(
    ...     ols1, ols2, iv1, iv2,
    ...     column_spanners=[("OLS", 2), ("IV", 2)],
    ...     stats=["N", "R2", "depvar_mean", "depvar_sd"],
    ... )
    >>>
    >>> # Unified coef_map (rename + order + drop in one shot)
    >>> sp.regtable(m1, m2, coef_map={
    ...     "x2": "Education",
    ...     "x1": "Experience",
    ...     "Intercept": "Constant",
    ... })
    """
    if not args:
        raise ValueError("At least one model result is required.")

    _VALID_OUTPUTS = {
        "text", "latex", "tex", "html", "markdown", "md",
        "quarto", "qmd", "word", "excel",
    }
    if output not in _VALID_OUTPUTS:
        raise ValueError(
            f"output={output!r} is invalid. Must be one of: "
            f"{sorted(_VALID_OUTPUTS)}"
        )

    # --- Resolve journal template (sets defaults; explicit kwargs win) ---
    se_label_override: Optional[str] = None
    template_notes: List[str] = []
    if template is not None:
        preset = get_template(template)
        if star_levels is None:
            star_levels = tuple(preset["star_levels"])
        if stats is None:
            stats = list(preset["stats"])
        if se_type == "se":
            se_label_override = preset.get("se_label")
        # Footer notes from the template are appended *after* user notes,
        # except we skip the boilerplate "stars" / "SE in parentheses"
        # lines because the renderer emits those itself.
        for line in preset.get("notes_default", ()):
            low = line.lower()
            if "in parenthes" in low or "p<0" in low:
                continue
            template_notes.append(line)

    if star_levels is None:
        star_levels = (0.10, 0.05, 0.01)

    # --- coef_map shortcut (mirrors R modelsummary's three-in-one) ----
    # When set, it simultaneously renames + reorders + drops via a single
    # ordered dict. Conflicts with the legacy keep/drop/order/coef_labels
    # parameters are rejected up front because resolving them is ambiguous
    # and silent precedence would surprise users.
    if coef_map is not None:
        if coef_labels is not None:
            raise ValueError(
                "Pass either coef_map or coef_labels, not both. coef_map is "
                "the unified shortcut (rename + order + drop); coef_labels "
                "only renames."
            )
        if keep is not None or drop is not None or order is not None:
            raise ValueError(
                "coef_map already defines the keep / order behaviour "
                "(via its insertion order and key set). Drop the explicit "
                "keep/drop/order arguments when using coef_map."
            )
        coef_labels = dict(coef_map)
        keep = list(coef_map.keys())
        order = list(coef_map.keys())

    # --- Detect panel structure ---
    # If first arg is a list, treat each positional arg as a panel
    if isinstance(args[0], list):
        raw_panels = list(args)
        flat_results = [r for raw in raw_panels for r in raw]
    else:
        raw_panels = [list(args)]
        flat_results = list(args)

    # Extract model data per panel
    panels: List[_PanelData] = []
    total_models = 0
    for raw in raw_panels:
        model_data_list = [_extract_model_data(r) for r in raw]
        panels.append(_PanelData(model_data_list))
        total_models += len(model_data_list)

    # --- vcov= : recompute SE / t / p / CI for OLS-style results ------
    # The recompute happens *after* extraction so that journal templates
    # and other pre-extraction logic stay untouched. Non-OLS results
    # whose data_info lacks X / residuals are silently skipped (with a
    # UserWarning) so heterogeneous tables don't blow up — the user can
    # audit which columns retained their fit-time SEs.
    if vcov is not None:
        valid = {"HC0", "HC1", "HC2", "HC3", "ROBUST"}
        if vcov.upper() not in valid:
            raise ValueError(
                f"vcov={vcov!r} not supported. Valid choices: "
                f"{sorted(valid)} (Stata's 'robust' is an alias for HC1)."
            )
        _apply_vcov_to_panels(panels, flat_results, vcov)

    # --- transpose validation (must run BEFORE we lose multi-panel info) ----
    if transpose:
        if isinstance(args[0], list) and len(args) > 1:
            raise NotImplementedError(
                "transpose=True is single-panel only. Multi-panel "
                "input is rejected to keep the layout pivot semantics "
                "tight; render each panel as its own transposed table."
            )
        if multi_se:
            raise NotImplementedError(
                "transpose=True is incompatible with multi_se: extra "
                "SE specs are bracket-rendered below the primary SE in "
                "the conventional layout, and there is no equally clean "
                "convention for the transposed view."
            )

    # --- Resolve eform flags (one bool per flat model position) -------
    if isinstance(eform, bool):
        eform_flags = [eform] * total_models
    else:
        eform_seq = list(eform)
        if len(eform_seq) != total_models:
            raise ValueError(
                f"eform list has {len(eform_seq)} entries but there are "
                f"{total_models} models."
            )
        eform_flags = [bool(f) for f in eform_seq]

    # eform and apply_coef are mutually exclusive — both transform the
    # point estimate, and silently combining them would hide whichever
    # the user listed second.
    if apply_coef is not None and any(eform_flags):
        raise ValueError(
            "eform and apply_coef both transform the coefficient. "
            "Pick one — eform=True is the canonical exp(b) shortcut "
            "for logit/poisson/cox; apply_coef accepts an arbitrary "
            "callable for percentage / log / signed-sqrt transforms."
        )

    # Validate template placeholders early — surface a KeyError at the
    # regtable() call site rather than buried inside the renderer.
    if estimate is not None or statistic is not None:
        _allowed_keys = {
            "estimate", "std_error", "t_value", "p_value",
            "conf_low", "conf_high", "stars",
        }
        for tmpl_name, tmpl in (("estimate", estimate), ("statistic", statistic)):
            if tmpl is None:
                continue
            try:
                tmpl.format_map({k: "" for k in _allowed_keys})
            except KeyError as exc:
                raise KeyError(
                    f"{tmpl_name}={tmpl!r} references unknown placeholder "
                    f"{exc.args[0]!r}. Allowed placeholders: "
                    f"{sorted(_allowed_keys)}."
                ) from exc

    # --- Consistency checks: warn on N-mismatch -----------------------
    # Mixed sample sizes across columns is a Reviewer red flag. We don't
    # *block* (sometimes mixing is intentional — IV first stage on a
    # subsample, RD bandwidth restriction); we warn so the user puts an
    # explicit footnote.
    if consistency_check and total_models >= 2:
        ns = []
        for p in panels:
            for m in p.models:
                n_val = m.stats.get("N")
                if n_val is not None:
                    try:
                        ns.append(int(n_val))
                    except (TypeError, ValueError):
                        pass
        if len(ns) >= 2 and (max(ns) - min(ns) > 0):
            warnings.warn(
                f"Sample sizes differ across columns (range "
                f"{min(ns):,}–{max(ns):,}). If this is intentional (e.g. "
                f"IV first stage on a subsample), add a footnote via "
                f"notes=[...]; otherwise re-fit on a common sample.",
                UserWarning,
                stacklevel=2,
            )

    # Default model labels
    if model_labels is None:
        model_labels = [f"({i + 1})" for i in range(total_models)]
    elif len(model_labels) != total_models:
        raise ValueError(
            f"model_labels has {len(model_labels)} entries but "
            f"there are {total_models} models."
        )

    # Validate dep_var_labels length
    if dep_var_labels is not None and len(dep_var_labels) != total_models:
        raise ValueError(
            f"dep_var_labels has {len(dep_var_labels)} entries but "
            f"there are {total_models} models."
        )

    # --- Auto-extract diagnostic rows ----------------------------------
    if diagnostics in (False, "off"):
        auto_rows: Dict[str, List[str]] = {}
    else:
        auto_rows = dict(extract_diagnostic_rows(flat_results))

    # Merge: user's add_rows wins on collisions, auto rows fill gaps.
    # Backwards-compat: if user supplies the legacy "Fixed Effects" row,
    # suppress the auto-emitted per-variable FE rows ("Firm FE", "Year FE",
    # …) so old tables don't suddenly show a single row + auto rows stacked.
    user_rows = dict(add_rows) if add_rows else {}
    if "Fixed Effects" in user_rows:
        auto_rows = {k: v for k, v in auto_rows.items() if not k.endswith(" FE")}
    merged_add_rows: Dict[str, List[str]] = {}
    for label, vals in auto_rows.items():
        if label not in user_rows:
            merged_add_rows[label] = list(vals)
    for label, vals in user_rows.items():
        merged_add_rows[label] = list(vals)

    # --- fixef_sizes: emit "# <FE>: count" rows ------------------------
    # Reads ``model_info['n_fe_levels']`` (a dict of FE-name → integer
    # count) per model — currently populated by count.py / nbreg /
    # ppmlhdfe; other estimators are silent no-ops. Mirrors R fixest's
    # ``fixef_sizes = TRUE``.
    if fixef_sizes:
        fe_size_rows = _build_fe_size_rows(flat_results)
        for label, vals in fe_size_rows.items():
            if label not in merged_add_rows:
                merged_add_rows[label] = vals

    # --- tests=: render hypothesis-test footer rows --------------------
    tests_rows_norm: List[Tuple[str, List[str]]] = []
    if tests:
        tests_rows_norm = _resolve_tests(
            tests, total_models, fmt=fmt, stars=stars,
            star_levels=star_levels, notation=notation,
        )

    # --- Resolve multi_se -----------------------------------------------
    multi_se_norm = _resolve_multi_se(multi_se, total_models)

    # --- Resolve reproducibility note -----------------------------------
    final_notes = list(notes) if notes else []
    final_notes.extend(template_notes)
    if repro:
        repro_kwargs = dict(repro) if isinstance(repro, dict) else {}
        repro_note = build_repro_note(**repro_kwargs)
        if repro_note:
            final_notes.append(repro_note)

    result = RegtableResult(
        panels=panels,
        panel_labels=panel_labels,
        model_labels=model_labels,
        dep_var_labels=dep_var_labels,
        coef_labels=coef_labels,
        keep=list(keep) if keep else None,
        drop=list(drop) if drop else None,
        order=list(order) if order else None,
        se_type=se_type,
        stars=stars,
        star_levels=tuple(star_levels),
        fmt=fmt,
        title=title,
        notes=final_notes,
        add_rows=merged_add_rows,
        stats=list(stats) if stats else None,
        output=output,
        alpha=alpha,
        multi_se=multi_se_norm,
        se_label=se_label_override,
        template=template,
        quarto_label=quarto_label,
        quarto_caption=quarto_caption,
        eform_flags=eform_flags,
        column_spanners=list(column_spanners) if column_spanners else None,
        estimate_template=estimate,
        statistic_template=statistic,
        notation=notation,
        apply_coef=apply_coef,
        apply_coef_deriv=apply_coef_deriv,
        escape=escape,
        tests_rows=tests_rows_norm,
        transpose=transpose,
    )

    # --- Output handling ---
    # Do NOT auto-print: Jupyter renders via _repr_html_, REPL via __repr__.
    # Scripts that want the rendered text should `print(regtable(...))`.
    if filename:
        result.save(filename)
    elif output in ("word", "excel"):
        warnings.warn(
            f"output='{output}' requires a filename. "
            f"Use filename='table.{'docx' if output == 'word' else 'xlsx'}'"
        )

    return result


# ═══════════════════════════════════════════════════════════════════════════
# ``MeanComparisonResult`` / ``mean_comparison`` moved to
# ``mean_comparison.py`` (was 510 lines wedged at the tail of this file).
# Re-export here so existing
# ``from statspai.output.regression_table import MeanComparisonResult``
# imports keep working unchanged.
# ═══════════════════════════════════════════════════════════════════════════

from .mean_comparison import MeanComparisonResult, mean_comparison  # noqa: E402,F401
